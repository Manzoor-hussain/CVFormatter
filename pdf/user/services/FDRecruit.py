import os
import openai
import docx
import docx2txt
from .keys import api_key
from pprint import pprint
import json
import re
import textwrap
import PyPDF2
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def get_completion(prompt, model="gpt-3.5-turbo"):
    messages = [{"role": "user", "content": prompt}]
    response = openai.ChatCompletion.create(
        model=model,
        messages=messages,
        temperature=0, # this is the degree of randomness of the model's output
    )
    return response.choices[0].message["content"]
 

def fd_recruit_converter(path, path_out, path_save):
    
    formatted = path_out
    
    try:
        with open(path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            unformated_text = ""
            for i in range (len(pdf_reader.pages)):
                first_page = pdf_reader.pages[i]
                unformated_text += first_page.extract_text() + " "
            print('Its PDF')
    except:
        try:
            unformated_text = docx2txt.process(path)
            print('Its Docx')
        except:
            print('WE DONT SUPPORT THIS TYPE OF FILE')

    # formatted document
    formatted_text = docx2txt.process(formatted)
    
    print("----------------------------------------------------------------")
    print("                          Unformatted Text                            ")
    print("----------------------------------------------------------------")
    print(unformatted_text)
    print("Process has started...")
    
    openai.api_key = api_key
    test_text = """

    Extract data from this text:

    \"""" + unformatted_text + """\"

    in following JSON format:
    {
    "Name" : "value",
    "Resides" : "value",
    "Education" : [
        {"Institute Name" : "Name Of institute",
        "Degree Name": "Name of degree",
        "Duration" : "Studying duration in institute",
        "Study Details" : ["Study Detail1", "Study Detail2", ...],
        },
        {"Institute Name" : "Name Of institute",
        "Degree Name": "Name of degree",
        "Duration" : "Studying duration in institute",
        "Study Details" : ["Study Detail1", "Study Detail2", ...],
        },
        ...
        ],
    "Profile" : "value",

    "Career History" : [
        {"Company Name" : "Name of company",
        "Job Title" : "Title of job",
        "Duration" : "Working Duration in Company",
        "Responsibilities" : ["Responsibility 1", "Responsibility 2", ...],
        },
        {"Company Name" : "Name of company",
        "Job Title" : "Title of job",
        "Duration" : "Working Duration in Company",
        "Responsibilities" : ["Responsibility 1", "Responsibility 2", ...],
        },
        ...
        ],

    "Courses and Trainings" : ["Course and Training 1", "Course and Training 2", ...],
    "Key Skills" : ["Key Skill1", "Key Skill2", ...],
    "Languages" : ["language1", "language2", ...],
    "Interests" : ["Interest1", "Interest2", ...],

    }
        You must keep the following points in considration while extracting data from text:
        1. Do NOT split, rephrase or summarize list of Responsibilities. Extract each Responsibility as a complete sentence from text.
        2. Make it sure to keep the response in JSON format.
        3. If value not found then leave it empty/blank.
        4. Do not include Mobile number, Email and Home address.
        5. Summary/Personal Statement should be as it is. Do not change or rephrase it.
    """
    result = get_completion(test_text)
    
    
    print("----------------------------------------------------------------")
    print("                          Result                            ")
    print("----------------------------------------------------------------")
    print(result)
    
    dc = dict(json.loads(re.sub(r'\[\"\"\]',r'[]',re.sub(r'\"[Un]nknown\"|\"[Nn]one\"|\"[Nn]ull\"|\"[Nn]ot [Mm]entioned\"',r'""',re.sub(r',[ \n]*\]',r']',re.sub(r',[ \n]*\}',r'}',result.replace('...','')))))))
    
    print("----------------------------------------------------------------")
    print("                          Dictionary                            ")
    print("----------------------------------------------------------------")
    print(dc)
    
    
    doc = docx.Document(formatted)


    for table in doc.tables:
            for row in table.rows:
                for i,cell in enumerate(row.cells):

    #                 doc.paragraphs[i].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                    try:
                        if cell.text.strip(' :\n').lower() == 'name':
                            if dc['Name'].strip() and dc['Name'].lower().replace(' ','') != 'value':
                                row.cells[i+1].text = dc['Name']
                    except:
                        pass

                    try:
                        if cell.text.strip(' :\n').lower() == 'resides':
                            if dc['Resides'].strip() and dc['Resides'].lower().replace(' ','') != 'value':
                                row.cells[i+1].text = dc['Resides']
                    except:
                        pass

                    try:
                        if cell.text.strip(' :\n').lower() == 'education':
                            for j in dc['Education']:
                                institute_name = ""
                                duration = ""
                                degree_name = ""
                                if j['Degree Name'].strip() and j['Degree Name'].lower().replace(' ','') != 'nameofdegree':
                                    degree_name = j['Degree Name'].strip()
                                    
                                    if j['Institute Name'].strip() and j['Institute Name'].lower().replace(' ','') != 'nameofinstitute':
                                        institute_name = j['Institute Name'].strip()
                                    if j['Duration'].strip() and j['Duration'].lower().replace(' ','') != 'studyingdurationininstitute':
                                        duration = j['Duration'].strip()
                                    

                                    if duration:
                                        row.cells[i+1].text += duration + '\n'
                                    else:
                                        row.cells[i+1].text += "Duration not mentioned" + '\n'

                                    if institute_name:
                                        row.cells[i+1].text += institute_name + '\n'
                                    else:
                                        row.cells[i+1].text += "Institute not mentioned" + '\n'

                                    if degree_name:
                                        row.cells[i+1].text += degree_name + '\n\n'
                                    else:
                                        row.cells[i+1].text += "Degree not mentioned" + '\n\n'
                                    
                    except:
                        pass



    for i,p in enumerate(doc.paragraphs):
    
        if p.text.strip(' :\n').lower() == 'profile':
            try:
                if dc['Profile'].strip() and dc['Profile'].lower().replace(' ','') != 'value':
                    doc.paragraphs[i+2].add_run(dc['Profile'].strip())
    #             name_paragraph.runs[0].bold = True
            except:
                pass

        if p.text.strip(' :\n').lower() == 'career history':
            try:
                for j in dc['Career History']:
                    try:
                        company_name = ""
                        duration = ""
                        job_title = ""
                        if (j['Company Name'].strip() and j['Company Name'].lower().replace(' ','') != 'nameofcompany') or (j['Job Title'].strip() and j['Job Title'].lower().replace(' ','') != 'titleofjob'):                    
                            if (j['Company Name'].strip() and j['Company Name'].lower().replace(' ','') != 'nameofcompany'):
                                company_name = j['Company Name'].strip()
                            if (j['Duration'].strip() and j['Duration'].lower().replace(' ','') != 'workingdurationincompany'):
                                duration = j['Duration'].strip()
                            if (j['Job Title'].strip() and j['Job Title'].lower().replace(' ','') != 'titleofjob'):
                                job_title = j['Job Title'].strip()

                            if company_name:
                                doc.paragraphs[i+2].add_run(company_name + ' ').bold = True
                            else:
                                doc.paragraphs[i+2].add_run("Company not mentioned" + ' ').bold = True

                            if duration:
                                doc.paragraphs[i+2].add_run('(' + duration + ')' + '\n').bold = True
                            else:
                                doc.paragraphs[i+2].add_run('(' + "Duration not mentioned" + ')' + '\n').bold = True
                            if job_title:
                                doc.paragraphs[i+2].add_run(job_title + '\n\n').bold = False
                            else:
                                doc.paragraphs[i+2].add_run("Job Title not mentioned" + '\n\n').bold = False
                                
                            try:
                                if j["Responsibilities"][0].lower().replace(' ','') != "responsibility1":
                                    for k in j['Responsibilities']:
                                        if k.strip():
                                            doc.paragraphs[i+2].add_run('  • ' + k.strip() + '\n')
                                    doc.paragraphs[i+2].add_run("\n\n")
                            except:
                                pass
                    except:
                        pass
            except:
                pass



        if p.text.strip(' :\n').lower() == 'courses and trainings':
            try:
                if dc['Courses and Trainings'][0].lower().replace(' ','') != 'courseandtraining1':
                    for j in dc['Courses and Trainings']:
                        if j.strip():
                            language_run = doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n')

            except:
                pass

        if p.text.strip(' :\n').lower() == 'key skills':
            try:
                if dc['Key Skills'][0].lower().replace(' ','') != 'keyskill1':
                    for j in dc['Key Skills']:
                        if j.strip():
                            language_run = doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n')

            except:
                pass

        if p.text.strip(' :\n').lower() == 'languages':
            try:
                if dc['Languages'][0].lower().replace(' ','') != 'langauge1':
                    for j in dc['Languages']:
                        if j.strip():
                            language_run = doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n')

            except:
                pass

        if p.text.strip(' :\n').lower() == 'interests':
            try:
                if dc['Interests'][0].lower().replace(' ','') != 'interest1':
                    for j in dc['Interests']:
                        if j.strip():
                            doc.paragraphs[i+1].add_run("\n")
                            doc.paragraphs[i+1].add_run('  • ' + j.strip())
            except:
                pass

    doc.save(path_save)
    print("Conversion completed !!")
