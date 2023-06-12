import os
import openai
import docx
import docx2txt
from .keys import api_key
from pprint import pprint
import json
import re
import textwrap
import PyPDF2
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def get_completion(prompt, model="gpt-3.5-turbo"):
    messages = [{"role": "user", "content": prompt}]
    response = openai.ChatCompletion.create(
        model=model,
        messages=messages,
        temperature=0, 
    )
    return response.choices[0].message["content"]


def e_mobility_converter(path, pathout, path_save):
    formatted = pathout
    file_path = path

    try:
        with open(path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            unformated_text = ""
            for i in range (len(pdf_reader.pages)):
                first_page = pdf_reader.pages[i]
                unformated_text += first_page.extract_text() + " "
            print('Its PDF')
    except:
        try:
            unformated_text = docx2txt.process(path)
            print('Its Docx')
        except:
            print('WE DONT SUPPORT THIS TYPE OF FILE')
    
    
    print("----------------------------------------------------------------")
    print("                          Unformatted Text                            ")
    print("----------------------------------------------------------------")
    print(unformated_text)
    
    os.environ["OPEN_API_KEY"] = api_key
    openai.api_key = api_key

    print("Process has Started...")
    test_text = """

    Extract data from this text:

    \"""" + unformated_text + """\"

    in following JSON format:
    {
    "Name" : "value",
    "Current Company" : "value",
    "Position applied" : "value",
    "Location" : "value",
    "Notice period" : "value",
    "Reason for Leaving" : "value",
    "System Used" : "value",
    "Dealbreakers" : "value",
    "Candidate Summary" : "value",
    "Experience" : [
        {"Job Title" : "Title of job",
        "Company Name" : "Name of Company",
        "Duration" : "Working Duration in Company",
        "Responsibilities" : ["Responsibility1", "Responsibility2", ...]
        },
        {"Job Title" : "Title of job",
        "Company Name" : "Name of Company",
        "Duration" : "Working Duration in Company",
        "Responsibilities" : ["Responsibility1", "Responsibility2", ...]
        },
        ...
        ],
    "Education" : [
        {"Institute Name" : "Name Of institute",
        "Degree Name": "Name of degree",
        "Duration" : "Studying duration in institute",
        },
        {"Institute Name" : "Name Of institute",
        "Degree Name": "Name of degree",
        "Duration" : "Studying duration in institute",
        },
        ...
        ],
    "Publications" : ["Publication1","Publication2",...],
    "Projects" : ["Project1","Project2",...],
    "Qualifications" : ["Qualification1", "Qualification2", ...],
    "Certifications" : ["Certification1","Certification2",...],
    "Achievements" : ["Achievement1","Achievement2",...],
    "Skills" : ["Skill1", "Skill2", ...],
    "Languages" : ["Language1", "Language2", ...],
    "Interests" : ["interest1", "interest2", ...]
    }
    
    You must keep the following points in considration while extracting data from text:
        1. Do NOT split, rephrase or summarize list of Responsibilities. Extract each Responsibility as a complete sentence from text.
        2. Make it sure to keep the response in JSON format.
        3. If value not found then leave it empty/blank.
        4. Do not include Mobile number, Email and Home address.
        5. Summary/Personal Statement should be as it is. Do not change or rephrase it.
    """

    result = get_completion(test_text)

    print("----------------------------------------------------------------")
    print("                          Result                            ")
    print("----------------------------------------------------------------")
    print(result)
    
    dc = dict(json.loads(re.sub(r'\[\"\"\]',r'[]',re.sub(r'\"[Un]nknown\"|\"[Nn]one\"|\"[Nn]ull\"|\"[Nn]ot [Mm]entioned\"',r'""',re.sub(r',[ \n]*\]',r']',re.sub(r',[ \n]*\}',r'}',result.replace('...','')))))))
    
    print("----------------------------------------------------------------")
    print("                          Dictionary                            ")
    print("----------------------------------------------------------------")
    print(dc)


    doc = docx.Document(formatted)
    for table in doc.tables:
        for row in table.rows:
            for i,cell in enumerate(row.cells):
                try:
                    if cell.text.strip(' :\n').lower().replace(' ','') == 'name':
                        if dc['Name'] and dc['Name'].lower().replace(' ','') != 'value':
                            row.cells[i+1].text = dc['Name']
                except:
                    pass

                try:
                    if cell.text.strip(' :\n').lower().replace(' ','') == 'currentcompany':
                        if dc['Current Company'] and dc['Current Company'].lower().replace(' ','') != 'value':
                            row.cells[i+1].text = dc['Current Company']
                except:
                    pass

                try:
                    if cell.text.strip(' :\n').lower().replace(' ','') == 'positionapplied':
                        if dc['Position Applied'] and dc['Position Applied'].lower().replace(' ','') != 'value':
                            row.cells[i+1].text = dc['Position applied']
                except:
                    pass              

                try:
                    if cell.text.strip(' :\n').lower().replace(' ','') == 'location':
                        if dc['Location'] and dc['Location'].lower().replace(' ','') != 'value':
                            row.cells[i+1].text = dc['Location']
                except:
                    pass                  

                try:
                    if cell.text.strip(' :\n').lower().replace(' ','') == 'noticeperiod':
                        if dc['Notice Period'] and dc['Notice Period'].lower().replace(' ','') != 'value':
                            row.cells[i+1].text = ['Notice period']
                except:
                    pass                

                try:
                    if cell.text.strip(' :\n').lower().replace(' ','') == 'reasonforleaving':
                        if dc['Reason for Leaving'] and dc['Reason for Leaving'].lower().replace(' ','') != 'value':
                            row.cells[i+1].text = dc['Reason for Leaving']
                except:
                    pass                

                try:
                    if cell.text.strip(' :\n').lower().replace(' ','') == 'systemused':
                        if dc['System Used'] and dc['System Used'].lower().replace(' ','') != 'value':
                            row.cells[i+1].text = dc['System Used']
                except:
                    pass

                try:
                    if cell.text.strip(' :\n').lower().replace(' ','') == 'dealbreakers':
                        if dc['Dealbreakers'] and dc['Dealbreakers'].lower().replace(' ','') != 'value':    
                            row.cells[i+1].text = dc['Dealbreakers']
                except:
                    pass

    for i,p in enumerate(doc.paragraphs):

        if p.text.strip(' :\n').lower() == 'candidate summary':
            try:
                if dc['Candidate Summary'] and dc['Candidate Summary'].lower().replace(' ','') != 'value':
                    summary = doc.paragraphs[i+1] 
                    summary.text = str(dc['Candidate Summary'])
                    doc.paragraphs[i+1].add_run('    •   ' + summary)
    #                 doc.paragraphs[i+1].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            except:
                pass

        if p.text.strip(' :\n').lower() == 'experience':
            try:
                for j in dc['Experience']:
                    job_title = j['Job Title'].strip()
                    company_name = j['Company Name'].strip()
                    duration = j['Duration'].strip()
                    
                    if (j['Company Name'] and j['Company Name'].lower().replace(' ','') != 'nameofcompany') or (j['Job Title'] and j['Job Title'].lower().replace(' ','') != 'titleofjob'):
                        if j['Job Title'] and j['Job Title'].lower().replace(' ','') != 'titleofjob':
                            doc.paragraphs[i+1].add_run(job_title + '\n').bold = True
                        if j['Company Name'] and j['Company Name'].lower().replace(' ','') != 'nameofcompany':  
                            doc.paragraphs[i+1].add_run(company_name + '\n').bold = True
                        if j['Duration'] and j['Duration'].lower().replace(' ','') != 'workingdurationincompany':    
                            doc.paragraphs[i+1].add_run('(' + duration + ')' + '\n').bold = True
                        
                        doc.paragraphs[i+1].add_run('\n')
                        if j["Responsibilities"] and j["Responsibilities"][0].lower().replace(' ','') != "responsibility1":
                            for k in j['Responsibilities']:
                                if k.strip():
                                    doc.paragraphs[i+1].add_run('  • ' + k.strip() + '\n')
                            doc.paragraphs[i+1].add_run('\n\n')
            except:
                pass

        if p.text.strip(' :\n').lower() == 'education':
            try:
                for j in dc['Education']:
                    institute_name = j['Institute Name'].strip()
                    duration = j['Duration'].strip()
                    degree_name = j['Degree Name'].strip()
                    
                    if j['Degree Name'].strip() and j['Degree Name'].lower().replace(' ','') != 'nameofdegree': 
                        if j['Institute Name'].strip() and j['Institute Name'].lower().replace(' ','') != 'nameofinstitute':
                            doc.paragraphs[i+1].add_run(institute_name + ' ').bold = True
                        
                        doc.paragraphs[i+1].add_run(degree_name + '\n').bold = False
                        
                        if j['Duration'].strip() and j['Duration'].lower().replace(' ','') != 'studyingdurationininstitute':
                            doc.paragraphs[i+1].add_run('(' + duration + ')' + '\n\n').bold = True
                        else:
                            doc.paragraphs[i+1].add_run('(' + "Not mentioned" + ')' + '\n\n').bold = True
                        
            except:
                pass

        if p.text.strip(' :\n').lower() == 'publications':
            try:
                if dc['Publications'][0] and dc['Publications'][0].lower().strip() != 'publication1':
                    for j in dc['Publications']:
                        if j.strip():
                            doc.paragraphs[i+1].add_run('    •   ' + j.strip() + '\n')
            except:
                pass

        if p.text.strip(' :\n').lower() == 'projects':
            try:
                if dc['Projects'][0] and dc['Projects'][0].lower().strip() != 'project1':
                    for j in dc['Projects']:
                        if j.strip():
                            doc.paragraphs[i+1].add_run('    •   ' + j.strip() + '\n')
            except:
                pass

        if p.text.strip(' :\n').lower() == 'qualifications':
            try:
                if dc['Qualifications'][0] and dc['Qualifications'][0].lower().strip() != 'qualification1':
                    for j in dc['Qualifications']:
                        if j.strip():
                            doc.paragraphs[i+1].add_run('    •   ' + j.strip() + '\n')
            except:
                pass

        if p.text.strip(' :\n').lower() == 'certifications':
            try:
                if dc['Certifications'][0] and dc['Certifications'][0].lower().strip() != 'certification1':
                    for j in dc['Certifications']:
                        if j.strip():
                            doc.paragraphs[i+1].add_run('    •   ' + j.strip() + '\n')
            except:
                pass

        if p.text.strip(' :\n').lower() == 'achievements':
            try:
                if dc['Achievements'][0] and dc['Achievements'][0].lower().strip() != 'achievement1':
                    for j in dc['Achievements']:
                        if j.strip():
                            doc.paragraphs[i+1].add_run('    •   ' + j.strip() + '\n')
            except:
                pass

        if p.text.strip(' :\n').lower() == 'skills':
            try:
                if dc['Skills'][0] and dc['Skills'][0].lower().strip() != 'skill1':
                    for j in dc['Skills']:
                        if j.strip():
                            doc.paragraphs[i+1].add_run('    •   ' + j.strip() + '\n')
            except:
                pass

        if p.text.strip(' :\n').lower() == 'languages':
            try:
                if dc['Languages'][0] and dc['Languages'][0].lower().strip() != 'language1':
                    for j in dc['Languages']:
                        if j.strip():
                            doc.paragraphs[i+1].add_run('    •   ' + j.strip() + '\n')
            except:
                pass


        if p.text.strip(' :\n').lower() == 'interests':
            try:
                if dc['Interests'][0] and dc['Interests'][0].lower().strip() != 'interest1':
                    for j in dc['Interests']:
                        if j.strip():
                            doc.paragraphs[i+1].add_run('    •   ' + j.strip() + '\n')
            except:
                pass


    doc.save(path_save)
    print("Process has Completed...")
