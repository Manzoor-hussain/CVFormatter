import os
import openai
import docx
import docx2txt
from .keys import api_key
# from keys import api_key
from pprint import pprint
import json
import re
import textwrap
import PyPDF2
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def get_completion(prompt, model="gpt-3.5-turbo"):
    messages = [{"role": "user", "content": prompt}]
    response = openai.ChatCompletion.create(
        model=model,
        messages=messages,
        temperature=0, # this is the degree of randomness of the model's output
    )
    return response.choices[0].message["content"]



def fmcg_converter(path, path_out, path_save):
    
    formatted = path_out 
    
    try:
        with open(path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            unformated_text = ""
            for i in range (len(pdf_reader.pages)):
                first_page = pdf_reader.pages[i]
                unformated_text += first_page.extract_text() + " "
            print('Its PDF')
    except:
        try:
            unformated_text = docx2txt.process(path)
            print('Its Docx')
        except:
            print('WE DONT SUPPORT THIS TYPE OF FILE')
    
    formatted_text = docx2txt.process(formatted)
    
    print("----------------------------------------------------------------")
    print("                          Unformatted Text                            ")
    print("----------------------------------------------------------------")
    print(unformated_text)
    
    print("Process has started...")
    
    # Prompt
    openai.api_key = api_key

    test_text = """

    Extract data from this text:

    \"""" + unformated_text + """\"

    in following JSON format:
    {
    "Current Employer" : "value",
    "Job title" : "value",
    "Location" : "value",
    "Salary Sought" : "value",
    "Notice Period" : "value",

    "Name" : "value",
    "Profile" : "value",
    "Education" : [
        {"Institute Name" : "Name Of institute",
        "Degree Name": "Name of degree",
        "Duration" : "Studying duration in institute",
        },
        {"Institute Name" : "Name Of institute",
        "Degree Name": "Name of degree",
        "Duration" : "Studying duration in institute",
        },
        ...
        ],
    "Professional Qualifications" : ["Qualification1", "Qualification2", ...],
    "Skills" : ["Skill1", "Skill2", ...],
    "IT Skills" : ["IT Skill1", "IT Skill2", ...],
    "Activities" : ["Activity1", "Activity2", ...],
    "Interests" : ["interest1", "interest2", ...],
    "Languages" : ["Language1", "Language2", ...],
    "Career History" : [
        {"Company Name" : "Name of company",
        "Job Title" : "Title of job",
        "Duration" : "Working Duration in Company",
        "Responsibilities" : ["Responsibility 1", "Responsibility 2", ...],
        },
        {"Company Name" : "Name of company",
        "Job Title" : "Title of job",
        "Duration" : "Working Duration in Company",
        "Responsibilities" : ["Responsibility 1", "Responsibility 2", ...],
        },
        ...
        ]
    }
        1. Do NOT split, rephrase or summarize list of Responsibilities. Extract each Responsibility as a complete sentence from text.
        2. Make it sure to keep the response in JSON format.
        3. If value not found then leave it empty/blank.
        4. Do not include Mobile number, Email and Home address.
        5. Summary/Personal Statement should be as it is. Do not change or rephrase it.
        
    """

    result = get_completion(test_text)
    
    print("----------------------------------------------------------------")
    print("                          Result                            ")
    print("----------------------------------------------------------------")
    print(result)
    
    dc = dict(json.loads(re.sub(r'\[\"\"\]',r'[]',re.sub(r'\"[Un]nknown\"|\"[Nn]one\"|\"[Nn]ull\"|\"[Nn]ot [Mm]entioned\"',r'""',re.sub(r',[ \n]*\]',r']',re.sub(r',[ \n]*\}',r'}',result.replace('...','')))))))
    
    print("----------------------------------------------------------------")
    print("                          Dictionary                            ")
    print("----------------------------------------------------------------")
    print(dc)
 
    
    doc = docx.Document(formatted)

    for table in doc.tables:
        for row in table.rows:
            for i,cell in enumerate(row.cells):
                try:
                    if cell.text.strip(' :\n').lower() == 'current employer':
                        if dc['Current Employer'].strip() and dc['Current Employer'].lower().replace(' ','') != 'value':
                            row.cells[i+1].text = '\n' + dc['Current Employer']
                except:
                    pass
                try:
                    if cell.text.strip(' :\n').lower() == 'job title':
                        if dc['Job Title'].strip() and dc['Job Title'].lower().replace(' ','') != 'value':
                            row.cells[i+1].text = '\n' + dc['Job Title']
                except:
                    pass                
                try:
                    if cell.text.strip(' :\n').lower() == 'location':
                        if dc['Location'].strip() and dc['Location'].lower().replace(' ','') != 'value':
                            row.cells[i+1].text = '\n' + dc['Location']
                except:
                    pass                
                try:
                    if cell.text.strip(' :\n').lower() == 'salary sought':
                        if dc['Salary Sought'].strip() and dc['Salary Sought'].lower().replace(' ','') != 'value':
                            row.cells[i+1].text = '\n' + dc['Salary Sought']
                except:
                    pass                
                try:
                    if cell.text.strip(' :\n').lower() == 'notice period':
                        if dc['Notice Period'].strip() and dc['Notice Period'].lower().replace(' ','') != 'value':
                            row.cells[i+1].text = '\n' + dc['Notice Period']
                except:
                    pass                

    font_size = 16
    for i,p in enumerate(doc.paragraphs):
        
        if p.text.strip(' :\n').lower() == 'name':
            try:
                if dc['Name'].strip() and dc['Name'].lower().replace(' ','') != 'value':
                    name_paragraph = doc.paragraphs[i]
                    name_paragraph.text = str(dc['Name'])
                    name_paragraph.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
                    name_paragraph.runs[0].bold = True
                    name_paragraph.runs[0].font.size = Pt(font_size)
            except:
                pass


        if p.text.strip(' :\n').lower() == 'profile':
            try:
                if dc['Profile'].strip() and dc['Profile'].lower().replace(' ','') != 'value':
                    doc.paragraphs[i+2].text = str(dc['Profile'])
                    doc.paragraphs[i].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            except:
                pass


        if p.text.strip(' :\n').lower() == 'education':
            for j in dc['Education']:
                try:
                    institute_name = ""
                    duration = ""
                    degree_name = ""
                    if j['Degree Name'].strip() and j['Degree Name'].lower().replace(' ','') != 'nameofdegree':
                        degree_name = j['Degree Name'].strip()
                        if j['Institute Name'].strip() and j['Institute Name'].lower().replace(' ','') != 'nameofinstitute':
                            institute_name = j['Institute Name'].strip()
                        if j['Duration'].strip() and j['Duration'].lower().replace(' ','') != 'studyingdurationininstitute':
                            duration = j['Duration'].strip()

                        if duration:
                            doc.paragraphs[i+2].add_run(institute_name + ' ').bold = True
                        else:
                            doc.paragraphs[i+2].add_run("Institute not mentioned ").bold = True
                        if duration:
                            doc.paragraphs[i+2].add_run('(' + duration + ')' + '\n').bold = True
                        else:
                            doc.paragraphs[i+2].add_run('(' + "Duration not mentioned" + ')' + '\n').bold = True

                        doc.paragraphs[i+2].add_run(degree_name + '\n\n').bold = False
                except:
                    pass
            

        if p.text.strip(' :\n').lower() == 'professional qualifications':
            try:
                if dc['Professional Qualifications'][0].lower().replace(' ','') != 'qualification1':
                    for j in dc['Professional Qualifications']:
                        if k.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n')
            except:
                pass


        if p.text.strip(' :\n').lower() == 'skills':
            try:
                if dc['Skills'][0].lower().replace(' ','') != 'skill1':
                    for j in dc['Skills']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n')
            except:
                pass


        if p.text.strip(' :\n').lower() == 'it skills':
            try:
                if dc['IT Skills'][0].lower().replace(' ','') != 'itskill1':
                    for j in dc['IT Skills']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n')
            except:
                pass        



        if p.text.strip(' :\n').lower() == 'activities':
            try:
                if dc['Activities'][0].lower().replace(' ','') != 'activity1':
                    for j in dc['Activities']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n')
            except:
                pass

        if p.text.strip(' :\n').lower() == 'interests':
            try:
                if dc['Interests'][0].lower().replace(' ','') != 'interest1':
                    for j in dc['Interests']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n')
            except:
                pass

        if p.text.strip(' :\n').lower() == 'languages':
            try:
                if dc['Languages'][0].lower().replace(' ','') != 'language1':
                    for j in dc['Languages']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n')
            except:
                pass

        if p.text.strip(' :\n').lower() == 'career history':
            try:
                for j in dc['Career History']:
                    try:
                        company_name = ""
                        duration = ""
                        job_title = ""
                        if (j['Company Name'].strip() and j['Company Name'].lower().replace(' ','') != 'nameofcompany') or (j['Job Title'].strip() and j['Job Title'].lower().replace(' ','') != 'titleofjob'):
                            if (j['Company Name'].strip() and j['Company Name'].lower().replace(' ','') != 'nameofcompany'):
                                company_name = j['Company Name'].strip()
                            if (j['Duration'].strip() and j['Duration'].lower().replace(' ','') != 'workingdurationincompany'):
                                duration = j['Duration'].strip()
                            if (j['Job Title'].strip() and j['Job Title'].lower().replace(' ','') != 'titleofjob'):
                                job_title = j['Job Title'].strip()

                            if company_name:
                                doc.paragraphs[i+2].add_run(company_name + ' ').bold = True
                            else:
                                doc.paragraphs[i+2].add_run("Company not mentioned" + ' ').bold = True

                            if duration:
                                doc.paragraphs[i+2].add_run('(' + duration + ')' + '\n').bold = True
                            else:
                                doc.paragraphs[i+2].add_run('(' + "Duration not mentioned" + ')' + '\n').bold = True
                            if job_title:
                                doc.paragraphs[i+2].add_run(job_title + '\n\n').bold = False
                            else:
                                doc.paragraphs[i+2].add_run("Job Title not mentioned" + '\n\n').bold = False

                            try:
                                if j["Responsibilities"][0].lower().replace(' ','') != "responsibility1":
                                    for k in j['Responsibilities']:
                                        if k.strip():
                                            doc.paragraphs[i+2].add_run('  • ' + k.strip() + '\n')
                                    doc.paragraphs[i+2].add_run("\n\n")
                            except:
                                pass
                    except:
                        pass
            except:
                pass


    doc.save(path_save)
    print("Conversion has completed !!")
