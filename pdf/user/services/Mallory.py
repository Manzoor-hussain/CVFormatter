import os
import openai
import docx
import docx2txt
import re
import json
from .keys import api_key
from docx.enum.text import WD_UNDERLINE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
import PyPDF2
from docx.shared import Pt
from docx.shared import RGBColor


def get_completion(prompt, model="gpt-3.5-turbo"):
    messages = [{"role": "user", "content": prompt}]
    response = openai.ChatCompletion.create(
        model=model,
        messages=messages,
        temperature=0, # this is the degree of randomness of the model's output
    )
    return response.choices[0].message["content"]

def mallory_converter(path ,formatted_path,save_path):
    formatted= formatted_path
#     un_formatted=os.getcwd() + path
    
#     doc = docx.Document(un_formatted)
#     formated_text = docx2txt.process(formatted)
#     unformated_text = docx2txt.process(un_formatted)
    def read_text_from_docx(path):
        doc = docx.Document(path)
        text = [paragraph.text for paragraph in doc.paragraphs]
        return '\n'.join(text)

    def read_text_from_pdf(path):
        with open(path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = []
            for page in pdf_reader.pages:
                text.append(page.extract_text())
            return '\n'.join(text)

    if path.endswith('.docx'):
        unformated_text = read_text_from_docx(path)
    elif path.endswith('.pdf'):
        unformated_text = read_text_from_pdf(path)
    else:
        unformated_text = 'Unsupported file format'
        

    os.environ["OPEN_API_KEY"] = api_key
    openai.api_key = api_key

    
    print ("Process has Started...")
    test_text = """

    Ectract data from this text:

    \"""" + unformated_text + """\"

    in following JSON format:
{
"Name":"candidate name",
"Profile" : "value",
"Employment History" : [
    {"Company Name" : "Name of company",
     "Duration" : "Working duration in company",
     "Designation" : "Specific designation in that Company",
     "Responsibilities" : ["Responsibility 1", "Responsibility 2", ...],
    },
    
    {"Company Name" : "Name of company",
     "Duration" : "Working duration in company",
     "Designation" : "Specific designation in that Company",
     "Responsibilities" : ["Responsibility 1", "Responsibility 2", ...],
    },
    ...
    ]
"Education" : [
    {"Institute Name":"Name of that institute",
     "Duration":"duration of that degree",
     "Degree":"Name of that degree",
    },
    
    {"Institute Name":"Name of that institute",
     "Duration":"duration of that degree",
     "Degree":"Name of that degree",
    },
   ...],
"Trainings" : ["trainings1", "trainings2", ...],
"Skills" : ["skills1", "skills2", ...],
"Qualifications" : ["qualifications1", "qualifications2", ...],
"Language":["language1","language2",...],
"Interests":["interests1","interests2",...],

}

   You must keep the following points in considration while extracting data from text:
    1. Do NOT split, rephrase or summarize list of Responsibilities. Extract each Responsibility as a complete sentence from text.
    2. Make it sure to keep the response in JSON format.
    3. If value not found then leave it empty/blank.
    4. Do not include Mobile number, Email and Home address.
    5. Summary/Personal Statement should be complete without being rephrased.


    """


    result = get_completion(test_text)
#     print(result)

    dc = dict(json.loads(re.sub(r'\[\"\"\]',r'[]',re.sub(r'\"[Un]nknown\"|\"[Nn]one\"|\"[Nn]ull\"|\"[Nn]ot [Mm]entioned\"',r'""',re.sub(r',[ \n]*\]',r']',re.sub(r',[ \n]*\}',r'}',result.replace('...','')))))))
#     print(dc)

    doc = docx.Document(formatted)

    for i,p in enumerate(doc.paragraphs):
        try:
            if p.text.strip().lower() == 'name':
                doc.paragraphs[i].text = ""
                if dc['Name'].lower().replace(' ','')!='candidatename':
                    run = doc.paragraphs[i].add_run(dc['Name'].strip().title())
                    run.bold = True
                    run.font.color.rgb = RGBColor(72,0,0)  # Red color
                    run.font.size = Pt(23.5)
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'profile':
                if dc['Profile'].lower().replace(' ','') != 'value':
                    run1=doc.paragraphs[i+2].add_run(dc['Profile'].strip())
                    run1.bold=False
                    run1=doc.paragraphs[i+2].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        except:
                pass

        if p.text.strip(' :\n').lower() == 'education':
            for j in dc['Education']:
                if j['Degree'].strip() and j['Degree'].lower().replace(' ','') != "nameofthatdegree":
                    run3 = doc.paragraphs[i+2].add_run(j['Degree'].strip()+"\n")
                    run3.bold=False
                    run3.font.color.rgb = RGBColor(145, 148, 146)  # Red color

                    if j['Institute Name'].strip():
                        run1 = doc.paragraphs[i+2].add_run(j['Institute Name'].strip()+"\n")
                        run1.font.color.rgb = RGBColor(145, 148, 146)  # Red color
                    else:
                         doc.paragraphs[i+2].add_run("Institute not mentioned").bold=False
                    if j['Duration'].strip():
                        run2 = doc.paragraphs[i+2].add_run(j['Duration'].strip()+"\n\n")
                        run2.bold=False
                        run2.font.color.rgb = RGBColor(145, 148, 146)  # Red color
                    else:
                        doc.paragraphs[i+2].add_run("Duration not mentioned" + "\n").bold=False
          
                               
    #             run4.italic=True

    #             doc.paragraphs[i+2].add_run(j['Duration'].strip()+"\n").bold=False


    #             doc.paragraphs[i+2].add_run(j['Thesis'].strip() + "\n\n").bold=False

    #               + j['Institute'].strip() + "–" + j['Duration'].strip()).bold = False
    #               doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False



        try:
            if p.text.strip(' :\n').lower() == 'certifications':
                if dc['Certifications'][0].lower().replace(' ','') != 'certifications1':
                    for j in dc['Certifications']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'qualifications':
                if dc['Qualifications'][0].lower().replace(' ','') != 'qualifications1':
                    for j in dc['Qualifications']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'language':
                if dc['Language'][0].lower().replace(' ','') != 'language1':
                    for j in dc['Language']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'interests':
                if dc['Interests'][0].lower().replace(' ','') != 'interests1':
                    for j in dc['Interests']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'trainings':
                if dc['Training'][0].lower().replace(' ','') != 'training1':
                    for j in dc['Trainings']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'skills':
                if dc['Skills'][0].lower().replace(' ','')!= 'skills1':
                    for j in dc['Skills']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'employment history':
                for j in dc['Employment History']:                  
                    if j['Designation'].strip() and j['Designation'].lower().replace(' ','') !='specificdesignationinthatcompany' or (j['Company Name'].strip() and j['Company Name'].lower().replace(' ','') !='nameofcompany'):                        
                        if j['Designation'].strip():
                            run3=doc.paragraphs[i+2].add_run(j['Designation'].strip() + "\n")
                            run3.font.color.rgb = RGBColor(145, 148, 146)  # Red color
                            run3.bold=False
                        else:
                            doc.paragraphs[i+2].add_run("Designation not mentioned\n").bold=False
                        if j['Company Name'].strip():                                                   
                            run1=doc.paragraphs[i+2].add_run(j['Company Name'].strip()+"\n")
                            run1.font.color.rgb = RGBColor(145, 148, 146)  # Red color
                        else:
                            doc.paragraphs[i+2].add_run("Company Name not mentioned\n").bold=False
                        if j['Duration'].strip():                          
                            run2=doc.paragraphs[i+2].add_run(j['Duration'].strip() + "\n\n")
                            run2.font.color.rgb = RGBColor(145, 148, 146)  # Red color
                        else:
                            doc.paragraphs[i+2].add_run("Duration not mentioned\n\n").bold=False
                                
                        if j['Responsibilities'] and j['Responsibilities'][0].lower().replace(' ','') != 'responsibility1':
                            for k in j['Responsibilities']:
                                doc.paragraphs[i+2].add_run('  • ' + k.strip() + '\n').bold = False
                            doc.paragraphs[i+2].add_run('\n\n')
        except:
            pass

    doc.save(save_path)
    print("\n")
    print("---------------------------------------------------------------------------------------------------------------------")
    print("\n")
    print("Process has Completed...")
