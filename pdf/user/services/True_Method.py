import os
import openai
import docx
import docx2txt
from .keys import api_key
from pprint import pprint
import json
import re
import textwrap
import PyPDF2
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt


def get_completion(prompt, model="gpt-3.5-turbo"):
    messages = [{"role": "user", "content": prompt}]
    response = openai.ChatCompletion.create(
        model=model,
        messages=messages,
        temperature=0, # this is the degree of randomness of the model's output
    )
    return response.choices[0].message["content"]
    
    
def true_method_converter(path, path_out, path_save):
    
    formatted = path_out
    
    try:
        with open(path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            unformated_text = ""
            for i in range (len(pdf_reader.pages)):
                first_page = pdf_reader.pages[i]
                unformated_text += first_page.extract_text() + " "
            print('Its PDF')
    except:
        try:
            unformated_text = docx2txt.process(path)
            print('Its Docx')
        except:
            print('WE DONT SUPPORT THIS TYPE OF FILE')

    # formatted document
    formatted_text = docx2txt.process(formatted)
    print("----------------------------------------------------------------")
    print("                          Unformatted Text                            ")
    print("----------------------------------------------------------------")
    print(unformated_text)
    print("Process has started...")
    
    openai.api_key = api_key
    test_text = """

    Extract data from this text:

    \"""" + unformated_text + """\"

    in following JSON format:
    {
    "Name" : "value",
    "Nationality" : "value",
    "Profile" : "value",
    "Key Skills" : ["Key Skill1", "Key Skill2", ...],

    "Work Experience" : [
        {"Company Name" : "Name of company",
        "Job Title" : "Title of job",
        "Duration" : "Working Duration in Company",
        "Responsibilities" : ["Responsibility 1", "Responsibility 2", ...],
        },
        {"Company Name" : "Name of company",
        "Job Title" : "Title of job",
        "Duration" : "Working Duration in Company",
        "Responsibilities" : ["Responsibility 1", "Responsibility 2", ...],
        },
        ...
        ],

    "Languages" : ["Language1", "Language2", ...],
    "Education" : [
            {"Institute Name" : "Name Of institute",
            "Degree Name": "Name of degree",
            "Duration" : "Studying duration in institute",
            },
            {"Institute Name" : "Name Of institute",
            "Degree Name": "Name of degree",
            "Duration" : "Studying duration in institute",
            },
            ...
            ],
    "Professional Certifications" : ["Professional Certification1", "Professional Certification2", ...],
    "Interests" : ["Interest1", "Interest2", ...],

    }
    You must keep the following points in considration while extracting data from text:
        1. Do NOT split, rephrase or summarize list of Responsibilities. Extract each Responsibility as a complete sentence from text.
        2. Make it sure to keep the response in JSON format.
        3. If value not found then leave it empty/blank.
        4. Do not include Mobile number, Email and Home address.        
        5. Summary/Personal Statement should be as it is. Do not change or rephrase it.

    """
    result = get_completion(test_text)
    
    print("----------------------------------------------------------------")
    print("                          Result                            ")
    print("----------------------------------------------------------------")
    print(result)
    
    dc = dict(json.loads(re.sub(r'\[\"\"\]',r'[]',re.sub(r'\"[Un]nknown\"|\"[Nn]one\"|\"[Nn]ull\"|\"[Nn]ot [Mm]entioned\"',r'""',re.sub(r',[ \n]*\]',r']',re.sub(r',[ \n]*\}',r'}',result.replace('...','')))))))
    
    print("----------------------------------------------------------------")
    print("                          Dictionary                            ")
    print("----------------------------------------------------------------")
    print(dc)
     
    
    doc = docx.Document(formatted)

    font_size = 12

    for i,p in enumerate(doc.paragraphs):

    #         doc.paragraphs[i].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        if p.text.strip(' :\n').lower() == 'name':
            try:
                if dc['Name'] and dc['Name'].lower().replace(' ','') != 'value':
                    name_paragraph = doc.paragraphs[i+2]
                    name_paragraph.text = str(dc['Name'] + '\n')
    #             name_paragraph.runs[0].bold = True
            except:
                pass

        if p.text.strip(' :\n').lower() == 'nationality':
            try:
                if dc['Nationality'] and dc['Nationality'].lower().replace(' ','') != 'value':
                    name_paragraph = doc.paragraphs[i+2]
                    name_paragraph.text = str(dc['Nationality'] + '\n')
                    print("\n")
    #             name_paragraph.runs[0].bold = True
            except:
                pass


        if p.text.strip(' :\n').lower() == 'profile':
            try:
                if dc['Profile'] and dc['Profile'].lower().replace(' ','') != 'value':
                    doc.paragraphs[i+2].text = str(dc['Profile'] + '\n')
#                     doc.paragraphs[i].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

            except:
                pass

        if p.text.strip(' :\n').lower() == 'key skills':
            try:
                if dc['Key Skills'][0] and dc['Key Skills'][0].lower().strip() != 'keyskill1':
                    for j in dc['Key Skills']:
                        skill_run = doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n')
                        skill_run.font.size = Pt(font_size)
            except:
                pass


        if p.text.strip(' :\n').lower() == 'work experience':
            try:
                for j in dc['Work Experience']:
                    company_name = j['Company Name'].strip()
                    duration = j['Duration'].strip()
                    job_title = j['Job Title'].strip()
                    
                    if (j['Company Name'] and j['Company Name'].lower().replace(' ','') != 'nameofcompany') or (j['Job Title'] and j['Job Title'].lower().replace(' ','') != 'titleofjob'):
                        if j['Company Name'] and j['Company Name'].lower().replace(' ','') != 'nameofcompany':  
                            doc.paragraphs[i+2].add_run(company_name + '\n').bold = True
                        
                        if j['Duration'] and j['Duration'].lower().replace(' ','') != 'workingdurationincompany':    
                            doc.paragraphs[i+2].add_run('(' + duration + ')' + '\n').bold = True
                        else:
                            doc.paragraphs[i+2].add_run('(' + "Duration not mentioned" + ')' + '\n').bold = True
                        
                        if j['Job Title'] and j['Job Title'].lower().replace(' ','') != 'titleofjob':
                            job_title_run = doc.paragraphs[i+2].add_run(job_title + '\n\n')
                            job_title_run.bold = False

                        if j["Responsibilities"] and j["Responsibilities"][0].lower().replace(' ','') != "responsibility1":
                            for k in j['Responsibilities']:
                                doc.paragraphs[i+2].add_run('  • ' + k.strip() + '\n')
                            doc.paragraphs[i+2].add_run('\n')
            except:
                pass


        if p.text.strip(' :\n').lower() == 'languages':
            try:
                if dc['Languages'] and dc['Languages'][0].lower().strip() != 'language1':
                    for j in dc['Languages']:
                        language_run = doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n')
                        language_run.font.size = Pt(font_size)
            except:
                pass
            

        if p.text.strip(' :\n').lower() == 'education':
            try:
                for j in dc['Education']:
                    institute_name = j['Institute Name'].strip()
                    duration = j['Duration'].strip()
                    degree_name = j['Degree Name'].strip()
                    
                    if j['Degree Name'].strip() and j['Degree Name'].lower().replace(' ','') != 'nameofdegree': 
                        if j['Institute Name'].strip() and j['Institute Name'].lower().replace(' ','') != 'nameofinstitute':
                            doc.paragraphs[i+2].add_run(institute_name + ' ').bold = True
                        if j['Duration'].strip() and j['Duration'].lower().replace(' ','') != 'studyingdurationininstitute':
                            doc.paragraphs[i+2].add_run('(' + duration + ')' + '\n').bold = True
                        else:
                            doc.paragraphs[i+2].add_run('(' + "Not mentioned" + ')' + '\n').bold = True
                            
                        doc.paragraphs[i+2].add_run(degree_name + '\n\n').bold = False
            except:
                pass

        if p.text.strip(' :\n').lower() == 'professional certifications':
            try:
                if dc['Professional Certifications'] and dc['Professional Certifications'][0].lower().strip() != 'professionalcertification1':
                    for j in dc['Professional Certifications']:
                        language_run = doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n')
                        language_run.font.size = Pt(font_size)
            except:
                pass

        if p.text.strip(' :\n').lower() == 'interests':
            try:
                if dc['Interests'] and dc['Interests'][0].lower().strip() != 'interest1':
                    for j in dc['Interests']:
                        doc.paragraphs[i+1].add_run("\n")
                        doc.paragraphs[i+1].add_run('  • ' + j.strip())
            except:
                pass

    doc.save(path_save)
    print("Conversion completed !!")
    
