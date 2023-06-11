import os
import openai
import docx
import docx2txt
import re
import json
from .keys import api_key
from docx.enum.text import WD_UNDERLINE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
import PyPDF2


def get_completion(prompt, model="gpt-3.5-turbo"):
    messages = [{"role": "user", "content": prompt}]
    response = openai.ChatCompletion.create(
        model=model,
        messages=messages,
        temperature=0, # this is the degree of randomness of the model's output
    )
    return response.choices[0].message["content"]

def william_blake_converter(path, formatted_path,save_path):
    formatted = formatted_path

    try:
        with open(path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            unformated_text = ""
            for i in range (len(pdf_reader.pages)):
                first_page = pdf_reader.pages[i]
                unformated_text += first_page.extract_text() + " "
            print('Its PDF')
    except:
        try:
            unformated_text = docx2txt.process(path)
            print('Its Docx')
        except:
            print('WE DONT SUPPORT THIS TYPE OF FILE')
    
    
    print("----------------------------------------------------------------")
    print("                          Unformatted Text                            ")
    print("----------------------------------------------------------------")
    print(unformated_text)

    os.environ["OPEN_API_KEY"] = api_key
    openai.api_key = api_key

    
    print ("Process has Started...")
    test_text = """

    Ectract data from this text:

    \"""" + unformated_text + """\"

    in following JSON format:
    {
    "Name":"candidate name",
    "Summary" : "value",

    "Work Experience" : [
        {"Duration" : "Working duration in company",
         "Company Name" : "Name of company",
         "Location":"location of that company",
         "Designation" : "Specific designation in that Company",
         "Responsibilities" : ["Responsibility 1", "Responsibility 2", ...]
        },

        {"Duration" : "Working duration in company",
         "Company Name" : "Name of company",
         "Location":"location of that company",
         "Designation" : "Specific designation in that Company",
         "Responsibilities" : ["Responsibility 1", "Responsibility 2", ...]
        },
        ...
        ],
    "Education" : [
        {"Duration":"duration of that degree",
         "Institute Name":"Name of that institute",
         "Location":"Location of that institute",
         "Degree":"Name of that degree"
        },

        {"Duration":"duration of that degree",
         "Institute Name":"Name of that institute",
         "Location":"Location of that institute",
         "Degree":"Name of that degree"
        },
        ...
        ],
    "Skills" : ["skill1", "skill2", ...],
    "Qualifications" : ["qualifications1", "qualifications2", ...],
    "Certifications":["certifications1","certifications2",...],
    "Language":["language1","language2",...],
    "Interests":["interests1","interests2",...]

    }

    You must keep the following points in considration while extracting data from text:
    1. Do NOT split, rephrase or summarize list of Responsibilities. Extract each Responsibility as a complete sentence from text.
    2. Make it sure to keep the response in JSON format.
    3. If value not found then leave it empty/blank.
    4. Do not include Mobile number, Email and Home address.
    5. Summary/Personal Statement should be as it is. Do not change or rephrase it.
    """


    result = get_completion(test_text)
    print(result)

    print("----------------------------------------------------------------")
    print("                          Result                            ")
    print("----------------------------------------------------------------")
    print(result)
    
    dc = dict(json.loads(re.sub(r'\[\"\"\]',r'[]',re.sub(r'\"[Un]nknown\"|\"[Nn]one\"|\"[Nn]ull\"|\"[Nn]ot [Mm]entioned\"',r'""',re.sub(r',[ \n]*\]',r']',re.sub(r',[ \n]*\}',r'}',result.replace('...','')))))))
    
    print("----------------------------------------------------------------")
    print("                          Dictionary                            ")
    print("----------------------------------------------------------------")
    print(dc)

    doc = docx.Document(formatted)

    for i,p in enumerate(doc.paragraphs):
        try:
            if p.text.strip().lower() == 'name':
                if dc['Name'].strip() and dc['Name'].lower().replace(' ','') != 'candidatename':
                    doc.paragraphs[i].text = ""
                    run = doc.paragraphs[i].add_run(dc['Name'].strip().title())
                    run.bold = True
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'summary':
                if dc['Summary'].strip() and dc['Summary'].lower().replace(' ','') != 'value':
                    doc.paragraphs[i].text = ""
                    run1=doc.paragraphs[i].add_run(dc['Summary'].strip().title())
                    run1.bold=False
                    run1=doc.paragraphs[i].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        except:
            pass

        if p.text.strip(' :\n').lower() == 'education':
            try:
                for j in dc['Education']:
                    try:
                        if j['Degree'].strip() and j['Degree'].lower().replace(' ','') != 'nameofthatdegree':
                            if j['Duration'].strip() and j['Duration'].lower().replace(' ','') != 'durationofthatdegree':
                                run1 = doc.paragraphs[i+2].add_run(j['Duration'].strip())
                                run1.bold=True
                            else:
                                run1 = doc.paragraphs[i+2].add_run('Duration not mentioned')
                                run1.bold=True
                            if j['Institute Name'].strip() and j['Institute Name'].lower().replace(' ','') != 'nameofthatinstitute':
                                run2 = doc.paragraphs[i+2].add_run("\t\t"+j['Institute Name'].strip()+"\n").bold=True
                            else:
                                run2 = doc.paragraphs[i+2].add_run("\t\t"+'Institute name not mentioned'+"\n").bold=True
                            run4 = doc.paragraphs[i+2].add_run("\t\t\t\t"+j['Degree'].strip()+"\n\n")
                            run4.bold=False
                    except:
                        pass
            except:
                pass
            
        try:
            if p.text.strip(' :\n').lower() == 'certifications':
                if dc['Certifications'][0].lower().strip() != 'certifications1':
                    for j in dc['Certifications']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'language':
                if dc['Language'][0].lower().strip() != 'language1':
                    for j in dc['Language']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'interests':
                if dc['Interests'][0].lower().strip() != 'interest1':
                    for j in dc['Interests']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'training':
                if dc['Trainings'][0].lower().strip() != 'training1':
                    for j in dc['Training']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'skills':
                if dc['Skills'][0].lower().strip() != 'skill1':
                    for j in dc['Skills']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'work experience':
                for j in dc['Work Experience']:
                    try:
                        if (j['Company Name'].strip() and j['Company Name'].lower().replace(' ','') != 'nameofcompany') or (j['Job Title'].strip() and j['Job Title'].lower().replace(' ','') != 'specificdesignationinthatcompany'):
                            if (j['Duration'].strip() and j['Duration'].lower().replace(' ','') != 'workingdurationincompany'):
                                run1=doc.paragraphs[i+2].add_run(j['Duration'].strip()).bold=True
                            else:
                                run1=doc.paragraphs[i+2].add_run('Duration not mentioned').bold=True
                            if (j['Company Name'].strip() and j['Company Name'].lower().replace(' ','') != 'nameofcompany'):
                                run2=doc.paragraphs[i+2].add_run("\t\t\t"+j['Company Name'].strip()+"\n").bold=True
                            else:
                                run2=doc.paragraphs[i+2].add_run("\t\t\t"+'Company not mentioned'+"\n").bold=True
                            if (j['Designation'].strip() and j['Designation'].lower().replace(' ','') != 'specificdesignationinthatcompany'):
                                run3=doc.paragraphs[i+2].add_run("\t\t\t\t\t"+j['Designation'].strip() + "\n\n")
                                run3.bold=True
                            else:
                                run3=doc.paragraphs[i+2].add_run("\t\t\t\t\t"+'Designation not mentioned' + "\n\n")
                                run3.bold=True
                            try:
                                if j["Responsibilities"][0].lower().replace(' ','') != "responsibility1":
                                    for k in j['Responsibilities']:
                                        if k.strip():
                                            doc.paragraphs[i+2].add_run('  - ' + k.strip() + '\n')
                                    doc.paragraphs[i+2].add_run("\n\n")
                            except:
                                pass
                    except:
                        pass
        except:
            pass

    doc.save(save_path)

    print("Process has Completed...")
# path = ''
# EdEx_Converter(path)
