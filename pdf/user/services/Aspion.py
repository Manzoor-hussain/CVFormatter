import os
import openai
import docx
from pprint import pprint
import docx2txt
import PyPDF2
import re
import json
from .keys import api_key

def get_completion(prompt, model="gpt-3.5-turbo"):
    messages = [{"role": "user", "content": prompt}]
    response = openai.ChatCompletion.create(
        model=model,
        messages=messages,
        temperature=0, # this is the degree of randomness of the model's output
    )
    return response.choices[0].message["content"]


def aspion_converter(path,pathout,path_save):
    formatted = pathout
    un_formatted = path

    formated_text = docx2txt.process(formatted)

    try:
        with open(path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            unformated_text = ""
            for i in range (len(pdf_reader.pages)):
                first_page = pdf_reader.pages[i]
                unformated_text += first_page.extract_text() + " "
            print('Its PDF')
    except:
        try:
            unformated_text = docx2txt.process(path)
            print('Its Docx')
        except:
            print('WE DONT SUPPORT THIS TYPE OF FILE')


    openai.api_key = api_key
    
    print("Process has Started...")
    test_text = """Extract data from this text:

    \"""" + unformated_text + """\"

    in following JSON format:
    {
    "Profile Summary" : "value",
    "Education" : [
        {"Institute" : "Name Of institute",
        "Duration": "Studying duration in institute",
        "Location" : "Location of institute",
        "Degree title" : "Name of degree completed from that institute"
        },
        {"Institute" : "Name Of institute",
        "Duration": "Studying duration in institute",
        "Location" : "Location of institute",
        "Degree title" : "Name of degree completed from that institute"
        },
        ...
        ],
    "Achievements" : ["achievement1", "achievement2", ...],
    "Languages" : ["language1", "language2", ...],
    "Interests" : ["interest1", "interest2", ...],
    "Trainings" : ["training1", "training2", ...],
    "Skills" : ["skill1", "skill2", ...],
    "Experience" : [
        {"Designation" : "Specific designation in that Company",
        "Name of Company" : "Company Name here",
        "Company locality" : "Location of company",
        "Duration" : "Time period in whic he the person has worked in that company",
        "Responsibilities" : ["Responsibility 1", "Responsibility 2", ...],
        },
        {"Designation" : "Specific designation in that Company",
        "Name of Company" : "Company Name here",
        "Company locality" : "Location of company",
        "Duration" : "Time period in whic he the person has worked in that company",
        "Responsibilities" : ["Responsibility 1", "Responsibility 2", ...],
        },
        ...
        ]
    }

    You must keep the following points in considration while extracting data from text:
        1. Do NOT split, rephrase or summarize list of Responsibilities. Extract each Responsibility as a complete sentence from text.
        2. Make it sure to keep the response in JSON format.
        3. If value not found then leave it empty/blank.
        4. Do not include Mobile number, Email and Home address.
    """


    result = get_completion(test_text)
    print('\n\n\n')
    print(result)
    dc = dict(json.loads(re.sub(r'\[\"\"\]',r'[]',re.sub(r'\"[Un]nknown\"|\"[Nn]one\"|\"[Nn]ull\"|\"[Nn]ot [Mm]entioned\"',r'""',re.sub(r',[ \n]*\]',r']',re.sub(r',[ \n]*\}',r'}',result.replace('...','')))))))
    print('\n\n\n')
    print(dc)
    doc = docx.Document(formatted)

    for i,p in enumerate(doc.paragraphs):    

        try:
            if p.text.strip(' :\n').lower() == 'summary':
                if dc['Profile Summary'].lower().replace(' ','') != 'value':
                    doc.paragraphs[i+2].add_run(dc['Profile Summary'].strip()).bold = False
                    doc.paragraphs[i+2].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        except:
            pass

        try:     
            if p.text.strip(' :\n').lower() == 'education':
                for j in dc['Education']:
                    if j['Degree title'].strip():
                        if j['Degree title'].strip() and j['Degree title'].lower().replace(' ','') != "nameofdegreecompletedfromthatinstitute":
                            if j['Institute'].strip():
                                doc.paragraphs[i+2].add_run(j['Institute'].strip()).bold = True
                            else:
                                doc.paragraphs[i+2].add_run('Institute not mentioned').bold = True    

                            if j['Location'].strip():
                                doc.paragraphs[i+2].add_run(', ' + j['Location'].strip() + '\n').bold = True
                            else:
                                doc.paragraphs[i+2].add_run('\n').bold = True

                            doc.paragraphs[i+2].add_run(j['Degree title'].strip() + '\n').bold = True

                            if j['Duration'].strip():
                                doc.paragraphs[i+2].add_run(j['Duration'].strip() + '\n\n').bold = True
                            else:
                                doc.paragraphs[i+2].add_run('Duration not mentioned\n\n').bold = True

        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'achievements':
                if dc['Achievements'][0].lower().replace(' ','') != 'achievement1':
                    for j in dc['Achievements']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass


        try:
            if p.text.strip(' :\n').lower() == 'languages':
                if dc['Languages'][0].lower().replace(' ','') != 'language1':
                    for j in dc['Languages']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'interests':
                if dc['Interests'][0].lower().replace(' ','') != 'interest1':
                    for j in dc['Interests']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'training':
                if dc['Trainings'][0].lower().replace(' ','') != 'training1':
                    for j in dc['Trainings']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'skills':
                if dc['Skills'][0].lower().replace(' ','') != 'skill1':
                    for j in dc['Skills']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'career history':
                for j in dc['Experience']:
                    if j['Designation'].strip():
                        if (j['Designation'].strip() and j['Designation'].lower().replace(' ','') != "specificdesignationinthatcompany") or (j['Name of Company'].strip() and j['Name of Company'].lower().replace(' ','') != "companynamehere"):
                            if j['Designation'].strip():
                                doc.paragraphs[i+2].add_run(j['Designation'].strip() + '\n').bold = True
                            else:
                                doc.paragraphs[i+2].add_run('Designation not mentioned' + '\n').bold = True

                            if j['Name of Company'].strip():
                                doc.paragraphs[i+2].add_run(j['Name of Company'].strip()).bold = True
                            else:
                                doc.paragraphs[i+2].add_run('Company Name not mentioned').bold = True                    

                            if j['Company locality'].strip():
                                doc.paragraphs[i+2].add_run(', ' + j['Company locality'].strip() + '\n').bold = True
                            else:
                                doc.paragraphs[i+2].add_run('\n').bold = True

                            if j['Duration'].strip():
                                doc.paragraphs[i+2].add_run(j['Duration'].strip() + '\n\n').bold = True
                            else:
                                doc.paragraphs[i+2].add_run('Duration not mentioned' + '\n\n').bold = True                        
                            try:
                                if j["Responsibilities"] and j["Responsibilities"][0].lower().replace(' ','') != "responsibility1":
                                    doc.paragraphs[i+2].add_run("Responsibilities:" + '\n').bold = True
                                    for k in j['Responsibilities']:
                                        if k.strip():
                                            doc.paragraphs[i+2].add_run('  • ' + k.strip() + '\n').bold = False
                                    doc.paragraphs[i+2].add_run('\n')
                            except:
                                pass
        except:
            pass

    doc.save(path_save)
    print("Process has Completed...")
