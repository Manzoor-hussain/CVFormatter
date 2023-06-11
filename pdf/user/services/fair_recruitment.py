import os
import openai
import docx
import docx2txt
from .keys import api_key
from pprint import pprint
import json
import re
import textwrap
import PyPDF2
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def get_completion(prompt, model="gpt-3.5-turbo"):
    messages = [{"role": "user", "content": prompt}]
    response = openai.ChatCompletion.create(
        model=model,
        messages=messages,
        temperature=0, 
    )
    return response.choices[0].message["content"]

def fair_recruitment_converter(path, pathout, path_save):
    formatted = pathout
    file_path = path

    try:
        with open(path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            unformated_text = ""
            for i in range (len(pdf_reader.pages)):
                first_page = pdf_reader.pages[i]
                unformated_text += first_page.extract_text() + " "
            print('Its PDF')
    except:
        try:
            unformated_text = docx2txt.process(path)
            print('Its Docx')
        except:
            print('WE DONT SUPPORT THIS TYPE OF FILE')



    os.environ["OPEN_API_KEY"] = api_key
    openai.api_key = api_key
    
    print("Process has Started...")
    test_text = """

    Extract data from this text:

    \"""" + unformated_text + """\"

    in following JSON format:
    {
    "Name" : "value",
    "Summary" : "value",
    "Education" : [
        {"Institute Name" : "Name Of institute",
        "Degree Name": "Name of degree",
        "Duration" : "Studying duration in institute",
        },
        {"Institute Name" : "Name Of institute",
        "Degree Name": "Name of degree",
        "Duration" : "Studying duration in institute",
        },
        ...
        ],
    "Work and Team Experience" : [
        {"Company Name" : "Name of Company",
        "Duration" : "Working Duration in Company",
        "Job Title" : "Title of job",
        "Responsibilities" : ["Responsibility1", "Responsibility2", ...]
        },
        {"Company Name" : "Name of Company",
        "Duration" : "Working Duration in Company",
        "Job Title" : "Title of job",
        "Responsibilities" : ["Responsibilities1","Responsibilities2", ...]
        },
        ...
        ],
    "Achievements" : ["Achievement1","Achievement2",...],
    "Qualifications" : ["Qualification1", "Qualification2", ...],
    "Skills" : ["Skill1", "Skill2", ...],
    "Languages" : ["Language1", "Language2", ...],
    "Interests" : ["interest1", "interest2", ...]
    }

    You must keep the following points in considration while extracting data from text:
        1. Do NOT split, rephrase or summarize list of Responsibilities. Extract each Responsibility as a complete sentence from text.
        2. Make it sure to keep the response in JSON format.
        3. If value not found then leave it empty/blank.
        4. Do not include Mobile number, Email and Home address.
        5. Summary/Personal Statement should be complete without being rephrased.

    """

    result = get_completion(test_text)
    dc = dict(json.loads(re.sub(r'\[\"\"\]',r'[]',re.sub(r'\"[Un]nknown\"|\"[Nn]one\"|\"[Nn]ull\"|\"[Nn]ot [Mm]entioned\"',r'""',re.sub(r',[ \n]*\]',r']',re.sub(r',[ \n]*\}',r'}',result.replace('...','')))))))

    doc = docx.Document(formatted)

    for i,p in enumerate(doc.paragraphs):


        if p.text.strip(' :\n').lower() == 'name':
            try:
                name_paragraph = doc.paragraphs[i]
                if dc['Name'].lower().replace(' ','') != 'value':
                    name_paragraph.text = str(dc['Name'])
                    name_paragraph.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
            except:
                pass


        if p.text.strip(' :\n').lower() == 'summary':
            try:
                summary = doc.paragraphs[i] 
                if dc['Summary'].lower().replace(' ','') != 'value':
                    summary.text = str(dc['Summary'])
                    doc.paragraphs[i].add_run(summary + '\n')
                    doc.paragraphs[i].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            except:
                pass


        if p.text.strip(' :\n').lower() == 'education':
            try:
                for j in dc['Education']:
                    if j['Degree Name'].strip() and j['Degree Name'].lower().replace(' ','') != "nameofdegree":
                        if j['Institute Name'].strip(): 
                            doc.paragraphs[i+2].add_run(j['Institute Name'].strip() + '\n').bold = True
                        else:
                            doc.paragraphs[i+2].add_run("Institute Name not mentioned"+'\n')                           
                        if j['Degree Name'].strip():     
                            doc.paragraphs[i+2].add_run(j['Degree Name'].strip() + '\n').bold = True
                        else:
                             doc.paragraphs[i+2].add_run("Degree Name not mentioned"+'\n')                                         
                        if j['Duration'].strip():    
                            doc.paragraphs[i+2].add_run(j['Duration'].strip() + '\n\n').bold = False
                        else:
                            doc.paragraphs[i+2].add_run("Duration not mentioned"+'\n')                           

            except:
                pass

        if p.text.strip(' :\n').lower() == 'work and team experience':
            try:
                for j in dc['Work and Team Experience']:
                    if j['Job Title'].strip() and j['Job Title'].lower().replace(' ','') !='titleofjob' or (j['Company Name'].strip() and j['Company Name'].lower().replace(' ','') !='nameofcompany'):
                        if j['Company Name'].strip():
                            doc.paragraphs[i+2].add_run(j['Company Name'].strip() + '\n').bold = True
                        else:
                            doc.paragraphs[i+2].add_run("Company Name not mentioned"+'\n')                           
                         
                        if j['Duration'].strip():
                            doc.paragraphs[i+2].add_run(j['Duration'].strip() + '\n').bold = True
                        else:
                            doc.paragraphs[i+2].add_run("Duration not mentioned"+'\n')                           
                        if j['Job Title'].strip():
                            doc.paragraphs[i+2].add_run(j['Job Title'].strip() + '\n').bold = True
                            doc.paragraphs[i+2].add_run('\n')
                        else:
                            doc.paragraphs[i+2].add_run("Job Title not mentioned"+'\n') 
                            
                        if j['Responsibilities'] and j['Responsibilities'][0].lower().replace(' ','') != 'responsibility1':    
                            for k in j['Responsibilities']:
                                doc.paragraphs[i+2].add_run('     •   '+ k.strip() + '\n').bold = False
                            doc.paragraphs[i+2].add_run('\n')
            except:
                pass

        if p.text.strip(' :\n').lower() == 'achievements':
            try:
                if dc['Achievements'][0].lower().replace(' ','') != 'achievements1':
                    for j in dc['Achievements']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('     •   ' + j.strip() + '\n')
            except:
                pass

        if p.text.strip(' :\n').lower() == 'qualifications':
            try:
                if dc['Qualifications'][0].lower().replace(' ','') != 'qualifications1':
                    for j in dc['Qualifications']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('     •   ' + j.strip() + '\n')
            except:
                pass

        if p.text.strip(' :\n').lower() == 'skills':
            try:
                if dc['Skills'][0].lower().replace(' ','') != 'skills1':
                    for j in dc['Skills']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('     •   ' + j.strip() + '\n')
            except:
                pass

        if p.text.strip(' :\n').lower() == 'languages':
            try:
                if dc['Languages'][0].lower().replace(' ','') != 'languages1':
                    for j in dc['Languages']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('     •   '  + j.strip() + '\n')
            except:
                pass

        if p.text.strip(' :\n').lower() == 'interests':
            try:
                if dc['Interests'][0].lower().replace(' ','') != 'interests1':
                    for j in dc['Interests']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('     •   ' + j.strip() + '\n')
            except:
                pass


    doc.save(path_save)
    print("Procees has Completed...")
