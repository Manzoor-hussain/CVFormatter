import os
import openai
import docx
import docx2txt
from .keys import api_key
from pprint import pprint
import json
import re
import textwrap
import PyPDF2
import pdfplumber
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def get_completion(prompt, model="gpt-3.5-turbo"):
    messages = [{"role": "user", "content": prompt}]
    response = openai.ChatCompletion.create(
        model=model,
        messages=messages,
        temperature=0, # this is the degree of randomness of the model's output
    )
    return response.choices[0].message["content"]

# Functions to check whether the unformatted file is a docx or pdf
def read_text_from_docx(file_path):
    doc = docx.Document(file_path)
    text = [paragraph.text for paragraph in doc.paragraphs]
    return '\n'.join(text)

def read_text_from_pdf(file_path):
    with open(file_path, 'rb') as file:
        pdf_reader = PyPDF2.PdfReader(file)
        text = []
        for page in pdf_reader.pages:
            text.append(page.extract_text())
        return '\n'.join(text)

    
def joss_search_converter(path_in, path_out, path_save):
    
    formatted = path_out
    
    # unformatted document
    if path_in.endswith('.docx'):
        unformatted_text = read_text_from_docx(path_in)
    elif path_in.endswith('.pdf'):
        unformatted_text = read_text_from_pdf(path_in)
    else:
        error = 'Format not supported.'
        print(error)
    
    formatted_text = docx2txt.process(formatted)
    
    print("----------------------------------------------------------------")
    print("                          Unformatted Text                            ")
    print("----------------------------------------------------------------")
    print(unformatted_text)
    
    print("Process has started...")
    
    # Prompt
    openai.api_key = api_key
    
    test_text = """

    Extract data from this text:

    \"""" + unformatted_text + """\"

    in following JSON format:
    {
    "Name" : "value",
    "Notice Period" : "value",
    "Holiday Dates" : "value",
    "Candidate Overview" : "value",

    "Summary" : "value",
    "Experience" : [
        {"Company Name" : "Name of company",
        "Job Title" : "Title of job",
        "Duration" : "Working Duration in Company in (Mon YYYY - Mon YYYY) Format.",
        "Responsibilities" : ["Responsibility 1", "Responsibility 2", ...],
        },
        {"Company Name" : "Name of company",
        "Job Title" : "Title of job",
        "Duration" : "Working Duration in Company in (Mon YYYY - Mon YYYY) Format.",
        "Responsibilities" : ["Responsibility 1", "Responsibility 2", ...],
        },
        ...
        ]
    }
    "Education" : [
        {"Institute Name" : "Name Of institute",
        "Degree Name": "Name of degree",
        "Duration" : "Studying duration in institute in (Mon YYYY - Mon YYYY) Format.",
        },
        {"Institute Name" : "Name Of institute",
        "Degree Name": "Name of degree",
        "Duration" : "Studying duration in institute in (Mon YYYY - Mon YYYY) Format.",
        },
        ...
        ],
    "Courses" : ["Course1", "Course2", ...],
    "Previous Assignments" : ["Previous Assignment1", "Previous Assignment2", ...],
    "Professional Qualifications" : ["Qualification1", "Qualification2", ...],
    "Areas of Expertise" : ["Area of Expertise1", "Area of Expertise2", ...],
    "Key Skills" : ["Key Skill1", "Key Skill2", ...],
    "Computer Skills": ["Computer Skill1", "Computer Skill2"],
    "Languages" : ["Language1", "Language2", ...],
    "Interests" : ["interest1", "interest2", ...],

    You must keep the following points in considration while extracting data from text:
        1. Do NOT split, rephrase or summarize list of Responsibilities. Extract each Responsibility as a complete sentence from text.
        2. Make it sure to keep the response in JSON format.
        3. If value not found then leave it empty/blank.
        4. Do not include Mobile number, Email and Home address.
        5. Summary/Personal Statement should be as it is. Do not change or rephrase it.
    """

    result = get_completion(test_text)
    
    print("----------------------------------------------------------------")
    print("                          Result                            ")
    print("----------------------------------------------------------------")
    print(result)
    
    dc = dict(json.loads(re.sub(r'\[\"\"\]',r'[]',re.sub(r'\"[Un]nknown\"|\"[Nn]one\"|\"[Nn]ull\"|\"[Nn]ot [Mm]entioned\"',r'""',re.sub(r',[ \n]*\]',r']',re.sub(r',[ \n]*\}',r'}',result.replace('...','')))))))
    
    print("----------------------------------------------------------------")
    print("                          Dictionary                            ")
    print("----------------------------------------------------------------")
    print(dc)


    doc = docx.Document(formatted)
    
    font_size = 14

    for table in doc.tables:
        for row in table.rows:
            for i,cell in enumerate(row.cells):
                
#                 doc.paragraphs[i].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                
                try:
                    if cell.text.strip(' :\n').lower() == 'notice period':
                        if dc['Notice Period'] and dc['Notice Period'].lower().replace(' ','') != 'value':
                            row.cells[i+1].text = dc['Notice Period']                            
                except:
                    pass
                try:
                    if cell.text.strip(' :\n').lower() == 'holiday dates':
                        if dc['Holiday Dates'] and dc['Holiday Dates'].lower().replace(' ','') != 'value':
                            row.cells[i+1].text = dc['Holiday Dates']
                        else:
                            pass
                except:
                    pass                
                try:
                    if cell.text.strip(' :\n').lower() == 'candidate overview':
                        if dc['Candidate Overview'] and dc['Candidate Overview'].lower().replace(' ','') != 'value':
                            row.cells[i+1].text = dc['Candidate Overview']
                except:
                    pass                


    for i,p in enumerate(doc.paragraphs):


        if p.text.strip(' :\n').lower() == 'name':
            try:
                if dc['Name'] and dc['Name'].lower().replace(' ','') != 'value':
                    name_paragraph = doc.paragraphs[i]
                    name_paragraph.text = str(dc['Name'])
                    name_paragraph.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
                    name_paragraph.runs[0].bold = True
                    name_paragraph.runs[0].font.size = Pt(font_size)
               
            except:
                pass


        if p.text.strip(' :\n').lower() == 'summary':
            try:
                if dc['Summary'] and dc['Summary'].lower().replace(' ','') != 'value':
                    doc.paragraphs[i+2].text = str(dc['Summary'])
#                     doc.paragraphs[i+2].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            except:
                pass


        if p.text.strip(' :\n').lower() == 'experience':
            try:
                for j in dc['Experience']:
                    company_name = j['Company Name'].strip()
                    duration = j['Duration'].strip()
                    job_title = j['Job Title'].strip()
                    
                    if (j['Company Name'] and j['Company Name'].lower().replace(' ','') != 'name of company') or (j['Job Title'] and j['Job Title'].lower().replace(' ','') != 'title of job'):
                        if j['Company Name'] and j['Company Name'].lower().replace(' ','') != 'name of company':  
                            doc.paragraphs[i+2].add_run(company_name + ' ').bold = True
                        if j['Duration'] and j['Duration'].lower().replace(' ','') != 'working duration in company in (mon yyyy - mon yyyy) format.':    
                            doc.paragraphs[i+2].add_run('(' + duration + ')' + '\n').bold = True
                        if j['Job Title'] and j['Job Title'].lower().replace(' ','') != 'title of job':
                            doc.paragraphs[i+2].add_run(job_title + '\n\n').bold = False

                        if j["Responsibilities"] and j["Responsibilities"][0].lower().replace(' ','') != "responsibility1":
                            for k in j['Responsibilities']:
                                doc.paragraphs[i+2].add_run('  • ' + k.strip() + '\n')
                            doc.paragraphs[i+2].add_run('\n')
            except:
                pass



        if p.text.strip(' :\n').lower() == 'education':
            try:
                for j in dc['Education']:
                    institute_name = j['Institute Name'].strip()
                    duration = j['Duration'].strip()
                    degree_name = j['Degree Name'].strip()
                    
                    if j['Degree Name'].strip() and j['Degree Name'].lower().replace(' ','') != 'name of degree': 
                        if j['Institute Name'].strip() and j['Institute Name'].lower().replace(' ','') != 'name of institute':
                            doc.paragraphs[i+2].add_run(institute_name + ' ').bold = True
                        if j['Duration'].strip() and j['Duration'].lower().replace(' ','') != 'studying duration in institute in (mon yyyy - mon yyyy) format.':
                            doc.paragraphs[i+2].add_run('(' + duration + ')' + '\n').bold = True
                        else:
                            doc.paragraphs[i+2].add_run('(' + "Not mentioned" + ')' + '\n').bold = True
                        
                        if j['Degree Name'].strip() and j['Degree Name'].lower().replace(' ','') != 'name of degree':
                            doc.paragraphs[i+2].add_run(degree_name + '\n\n').bold = False

            except:
                pass



        if p.text.strip(' :\n').lower() == 'courses':
            try:
                if dc['Courses'][0] and dc['Courses'][0].lower().strip() != 'course1':
                    for j in dc['Courses']:
                        doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n')
    #                     doc.paragraphs[i+2].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            except:
                pass


        if p.text.strip(' :\n').lower() == 'previous assignments':
            try:
                if dc['Previous Assignments'][0] and dc['Previous Assignments'][0].lower().strip() != 'previous assignment1':
                    for j in dc['Previous Assignments']:
                        doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n')
    #                     doc.paragraphs[i+2].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            except:
                pass        



        if p.text.strip(' :\n').lower() == 'professional qualifications':
            try:
                if dc['Professional Qualifications'][0] and dc['Professional Qualifications'][0].lower().strip() != 'qualification1':
                    for j in dc['Professional Qualifications']:
                        doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n')
    #                     doc.paragraphs[i+2].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            except:
                pass

        if p.text.strip(' :\n').lower() == 'area of expertise':
            try:
                if dc['Area of Expertise'][0] and dc['Area of Expertise'][0].lower().strip() != 'area of expertise1':
                    for j in dc['Area of Expertise']:
                        doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n')
#                     doc.paragraphs[i+2].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            except:
                pass

        if p.text.strip(' :\n').lower() == 'key skills':
            try:
                if dc['Key Skills'][0] and dc['Key Skills'][0].lower().strip() != 'key skill1':
                    for j in dc['Key Skills']:
                        doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n')
    #                     doc.paragraphs[i+2].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            except:
                pass
        if p.text.strip(' :\n').lower() == 'computer skills':
            try:
                if dc['Computer Skills'][0] and dc['Computer Skills'][0].lower().strip() != 'computer skill1':
                    for j in dc['Computer Skills']:
                        doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n')
    #                     doc.paragraphs[i+2].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            except:
                pass

        if p.text.strip(' :\n').lower() == 'languages':
            try:
                if dc['Languages'][0] and dc['Languages'][0].lower().strip() != 'language1':
                    for j in dc['Languages']:
                        doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n')
            except:
                pass

        if p.text.strip(' :\n').lower() == 'interests':
            try:
                if dc['Interests'][0] and dc['Interests'][0].lower().strip() != 'interest1':
                    for j in dc['Interests']:
                        doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n')
    #                     doc.paragraphs[i+2].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            except:
                pass



    doc.save(path_save)
    print("Conversion has completed !!")
    
    
    
