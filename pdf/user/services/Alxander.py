import os
import openai
import docx
import docx2txt
import re
import json
from .keys import api_key
from docx.enum.text import WD_UNDERLINE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import PyPDF2
import pdfplumber
from math import ceil 

def get_completion(prompt, model="gpt-3.5-turbo"):
    messages = [{"role": "user", "content": prompt}]
    response = openai.ChatCompletion.create(
        model=model,
        messages=messages,
        temperature=0, # this is the degree of randomness of the model's output
    )
    return response.choices[0].message["content"]

def alexander_steele_converter(path,pathoutput,save_path):
    formatted = pathoutput

    try:
        with open(path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            unformated_text = ""
            for i in range (len(pdf_reader.pages)):
                first_page = pdf_reader.pages[i]
                unformated_text += first_page.extract_text() + " "
            print('Its PDF')
    except:
        try:
            unformated_text = docx2txt.process(path)
            print('Its Docx')
        except:
            print('WE DONT SUPPORT THIS TYPE OF FILE')
    
    os.environ["OPEN_API_KEY"] = api_key
    openai.api_key = api_key
    # llm=OpenAI(temperature=0, max_tokens=1500,openai_api_key=api_key)
    fields_labels = "Name, PROFILE, EDUCATION, IT LITERACY, CERTIFICATES, PROJECTS, PROFESSIONAL QUALIFICATIONS, SOFTWARES, languages, Interests, TRAININGS, skills, WORK EXPERIENCE"

    
    print ("Process has Started...")
    test_text = """
    Extract data from this text:
    \"""" + unformated_text + """\"
    in following JSON format:
    {
    "Profile" : "value",
    "Experience/Employment History" : [
        {"Company Name" : "Name of company",
        "Duration" : "Working Duration in Company",
        "Designation" : "Specific designation in that Company",
        "Responsibilities" : ["Responsibility 1", "Responsibility 2", ...],
        },
        {"Company Name" : "Name of company",
        "Duration" : "Working Duration in Company",
        "Designation" : "Specific designation in that Company",
        "Responsibilities" : ["Responsibility 1", "Responsibility 2", ...],
        },
        ...
        ],
    "Education" : [
        {"Institute Name" : "Name Of institute",
        "Duration" : "Studying duration in institute",
        "Degree title" : "Name or title of Degree",
        },
        {"Institute Name" : "Name Of institute",
        "Duration" : "Studying duration in institute",
        "Degree title" : "Name or title of Degree",
        },
        ...
        ],
    "Professional Training/Courses" : ["Training1", "Training2", ...],
    "Achievements" : ["achievement1", "achievement2", ...],
    "Languages" : ["language1", "language2", ...],
    "Interests" : ["interest1", "interest2", ...],
    "Computer Skills" : ["computerskill1", "computerskill2", ...],
    "Skills" : ["skill1", "skill2", ...],
    }
    You must keep the following points in considration while extracting data from text:
        1. Do NOT split, rephrase or summarize list of Responsibilities. Extract each Responsibility as a complete sentence from text.
        2. Make it sure to keep the response in JSON format.
        3. If value not found then leave it empty/blank.
        4. Do not include Mobile number, Email and Home address. 
        5. Summary/Personal Statement should be complete without being rephrased.

    """

    result = get_completion(test_text)

    print(result)
    dc = dict(json.loads(re.sub(r'\[\"\"\]',r'[]',re.sub(r'\"[Un]nknown\"|\"[Nn]one\"|\"[Nn]ull\"|\"[Nn]ot [Mm]entioned\"',r'""',re.sub(r',[ \n]*\]',r']',re.sub(r',[ \n]*\}',r'}',result.replace('...','')))))))
    print(dc)

    doc = docx.Document(formatted)

    for i,p in enumerate(doc.paragraphs):
        try:
            if p.text.strip(' :\n').lower() == 'profile':
                if dc['Profile'].lower().replace(' ','') != 'value':
                    doc.paragraphs[i+2].add_run(dc['Profile'].strip()).bold = False
                    doc.paragraphs[i+2].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        except:
            pass

        try:        
            if p.text.strip(' :\n').lower() == 'education':
                for j in dc['Education']:
                    if j['Degree title'].strip() and j['Degree title'].lower().replace(' ','') != "nameortitleofdegree":
                        if j["Duration"].strip(): 
                            doc.paragraphs[i+2].add_run('  • ' + j["Duration"].strip())
                        else:
                            doc.paragraphs[i+2].add_run('  • ' + "Duration not mentioned")                           
                        if j["Institute Name"].strip():   
                            doc.paragraphs[i+2].add_run(' – ' + j["Institute Name"].strip())
                        else:
                            doc.paragraphs[i+2].add_run(" _ Institute not mentioned")
                            
                        doc.paragraphs[i+2].add_run(', ' + j["Degree title"].strip() + '\n\n')                                           
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'training':
                if dc['Professional Training/Courses'][0].lower().replace(' ','') != 'training1':
                    for j in dc['Professional Training/Courses']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass


        try:
            if p.text.strip(' :\n').lower() == 'achievements':
                if dc['Achievementss'][0].lower().replace(' ','') != 'achievements1':
                    for j in dc['Achievements']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'computer skills':
                if dc['Computer Skills'][0].lower().replace(' ','') != 'computerskills1':
                    for j in dc['Computer Skills']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass


        try:
            if p.text.strip(' :\n').lower() == 'languages':
                if dc['Languages'][0].lower().replace(' ','') != 'languages1':
                    for j in dc['Languages']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'interests':
                if dc['Interests'][0].lower().replace(' ','') != 'interests1':
                    for j in dc['Interests']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'trainings':
                if dc['Trainings'][0].lower().replace(' ','') != 'trainings1':
                    for j in dc['Trainings']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'skills':
                if dc['Skills'][0].lower().replace(' ','') != 'skills1':
                    for j in dc['Skills']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass


        if p.text.strip(' :\n').lower() == 'experience':
            for j in dc['Experience/Employment History']:               
                if j['Designation'].strip() and j['Designation'].lower().replace(' ','') !='specificdesignationinthatcompany' or (j['Company Name'].strip() and j['Company Name'].lower().replace(' ','') !='nameofcompany'):
                    a = ceil(len(j["Duration"]) * 1.75)
                    b = " " * a  
                    if j["Duration"].strip():                                             
                        run = doc.paragraphs[i+2].add_run(j["Duration"].strip())
                        run.font.bold = False  # Unbold the first company name
                    else:
                         run = doc.paragraphs[i+2].add_run("Duration not mentioned").bold=False
                    if j['Company Name'].strip():
                        run = doc.paragraphs[i+2].add_run('\t\t'+j['Company Name'] + '\n')
                        run.font.bold = True
                    else:
                        run = doc.paragraphs[i+2].add_run('\t\t' + "Company Name not mentioned"+"\n").bold=True
                    if j["Designation"].strip():   
                        doc.paragraphs[i+2].add_run(b + "\t\t" + j["Designation"].strip() + "\n\n").bold=True
                    else:
                        doc.paragraphs[i + 2].add_run(b + "\t\t" + "Designation not mentioned"+"\n\n").bold=True
                    if j['Responsibilities'] and j['Responsibilities'][0].lower().replace(' ','') != 'responsibility1':
                        for k in j["Responsibilities"]:
                            if k.strip():
                                doc.paragraphs[i+2].add_run('  • ' + k.strip() + "\n")
                        doc.paragraphs[i+2].add_run('\n\n')

    #     except:
    #         pass

    doc.save(save_path)
    print("\n")
    print("---------------------------------------------------------------------------------------------------------------------")
    print("\n")
    print("Process has Completed...")
