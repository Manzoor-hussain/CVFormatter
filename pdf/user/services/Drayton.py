import os
import openai
import docx
import docx2txt
from .keys import api_key
from pprint import pprint
import json
import re
import textwrap
import PyPDF2
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt


def get_completion(prompt, model="gpt-3.5-turbo"):
    messages = [{"role": "user", "content": prompt}]
    response = openai.ChatCompletion.create(
        model=model,
        messages=messages,
        temperature=0, # this is the degree of randomness of the model's output
    )
    return response.choices[0].message["content"]


# Functions to check whether the unformatted file is a docx or pdf
def read_text_from_docx(file_path):
    doc = docx.Document(file_path)
    text = [paragraph.text for paragraph in doc.paragraphs]
    return '\n'.join(text)

def read_text_from_pdf(file_path):
    with open(file_path, 'rb') as file:
        pdf_reader = PyPDF2.PdfReader(file)
        text = []
        for page in pdf_reader.pages:
            text.append(page.extract_text())
        return '\n'.join(text)
    

def drayton_converter(path_in, path_out, path_save):
    
    formatted = path_out
    
    # unformatted document
    if path_in.endswith('.docx'):
        unformatted_text = read_text_from_docx(path_in)
    elif path_in.endswith('.pdf'):
        unformatted_text = read_text_from_pdf(path_in)
    else:
        error = 'Format not supported.'
        print(error)

    # formatted document
    formatted_text = docx2txt.process(formatted)
    
    print("----------------------------------------------------------------")
    print("                          Unformatted Text                            ")
    print("----------------------------------------------------------------")
    print(unformatted_text)
    
    
    print("Process has started...")
        
    # Prompt
    openai.api_key = api_key
    test_text = """

    Extract data from this text:

    \"""" + unformatted_text + """\"

    in following JSON format:
    {
    "Name" : "value",
    "Location" : "value",
    "Academic Qualification" : "value in string",
    "Personal Profile" : "value",

    "Career Experience" : [
        {"Company Name" : "Name of company",
        "Job Title" : "Title of job",
        "Duration" : "Working Duration in Company",
        "Responsibilities" : ["Responsibility 1", "Responsibility 2", ...],
        },
        {"Company Name" : "Name of company",
        "Job Title" : "Title of job",
        "Duration" : "Working Duration in Company",
        "Responsibilities" : ["Responsibility 1", "Responsibility 2", ...],
        },
        ...
        ],

    "Skills" : ["Skill1", "Skill2", ...],
    "Interest and Hobbies" : ["interest and hobbies1", "interest and hobbies2", ...],
    "Languages" : ["Language1", "Language2", ...],

    }
    
    You must keep the following points in considration while extracting data from text:
        1. Do NOT split, rephrase or summarize list of Responsibilities. Extract each Responsibility as a complete sentence from text.
        2. Make it sure to keep the response in JSON format.
        3. If value not found then leave it empty/blank.
        4. Do not include Mobile number, Email and Home address. 
        5. Summary/Personal Statement should be as it is. Do not change or rephrase it.
    """
    result = get_completion(test_text)
    
    print("----------------------------------------------------------------")
    print("                          Result                            ")
    print("----------------------------------------------------------------")
    print(result)
    
    dc = dict(json.loads(re.sub(r'\[\"\"\]',r'[]',re.sub(r'\"[Un]nknown\"|\"[Nn]one\"|\"[Nn]ull\"|\"[Nn]ot [Mm]entioned\"',r'""',re.sub(r',[ \n]*\]',r']',re.sub(r',[ \n]*\}',r'}',result.replace('...','')))))))
    
    print("----------------------------------------------------------------")
    print("                          Dictionary                            ")
    print("----------------------------------------------------------------")
    print(dc)
    
    doc = docx.Document(formatted)
    
    for table in doc.tables:
        for row in table.rows:
            for i,cell in enumerate(row.cells):
                try:
                    if cell.text.strip(' :\n').lower() == 'name':
                        if dc['Name'] and dc['Name'].lower().replace(' ','') != 'value':
                            row.cells[i+1].text = dc['Name']
                except:
                    pass
                try:
                    if cell.text.strip(' :\n').lower() == 'location':
                        if dc['Location'] and dc['Location'].lower().replace(' ','') != 'value':
                            row.cells[i+1].text = dc['Location']
                except:
                    pass                
                try:
                    if cell.text.strip(' :\n').lower() == 'academic qualification':
                        if dc['Academic Qualification'] and dc['Academic Qualification'].lower().replace(' ','') != 'value':
                            row.cells[i+1].text = dc['Academic Qualification']
                        
                except:
                    pass                
                try:
                    if cell.text.strip(' :\n').lower() == 'personal profile':
                        if dc['Personal Profile'] and dc['Personal Profile'].lower().replace(' ','') != 'value':
                            personal_profile = row.cells[i+1]
                            personal_profile.text = dc['Personal Profile']
                           
                            paragraph = personal_profile.paragraphs[0]
#                             paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                        
                except:
                    pass                

    font_size = 12

    for i,p in enumerate(doc.paragraphs):

    #         doc.paragraphs[i].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        if p.text.strip(' :\n').lower() == 'career experience':
            try:
                
                for j in dc['Career Experience']:
                    company_name = j['Company Name'].strip()
                    duration = j['Duration'].strip()
                    job_title = j['Job Title'].strip()
                    
                    if (j['Company Name'] and j['Company Name'].lower().replace(' ','') != 'name of company') or (j['Job Title'] and j['Job Title'].lower().replace(' ','') != 'title of job'):
                        if j['Company Name'] and j['Company Name'].lower().replace(' ','') != 'name of company':  
                            company_run = doc.paragraphs[i+2].add_run(company_name + ' ')
                            company_run.bold = True
                            company_run.font.size = Pt(font_size)
                        if j['Duration'] and j['Duration'].lower().replace(' ','') != 'working duration in company':
                            duration_run = doc.paragraphs[i+2].add_run('(' + duration + ')' + '\n')
                            duration_run.bold = True
                            duration_run.font.size = Pt(font_size)
                        if j['Job Title'] and j['Job Title'].lower().replace(' ','') != 'title of job':
                            job_title_run = doc.paragraphs[i+2].add_run(job_title + '\n\n')
                            job_title_run.bold = False
                            job_title_run.font.size = Pt(font_size)

                        if j["Responsibilities"] and j["Responsibilities"][0].lower().replace(' ','') != "responsibility1":
                            for k in j['Responsibilities']:
                                respo = doc.paragraphs[i+2].add_run('  • ' + k.strip() + '\n')
                                respo.font.size = Pt(font_size)
                            doc.paragraphs[i+2].add_run("\n\n")
                    
            except:
                pass

        if p.text.strip(' :\n').lower() == 'skills':
            try:
                if dc['Skills'][0] and dc['Skills'][0].lower().strip() != 'skill1':
                    for j in dc['Skills']:
                        skill_run = doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n')
                        skill_run.font.size = Pt(font_size)
            except:
                pass

        if p.text.strip(' :\n').lower() == 'interest and hobbies':
            try:
                if dc['Interest and Hobbies'][0] and dc['Interest and Hobbies'][0].lower().strip() != 'interest and hobbies1':
                    for j in dc['Interest and Hobbies']:
                        i_h_run = doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n')
                        i_h_run.font.size = Pt(font_size)
            except:
                pass

        if p.text.strip(' :\n').lower() == 'language':
            try:
                if dc['Languages'][0] and dc['Languages'][0].lower().strip() != 'language1':
                    for j in dc['Languages']:
                        language_run = doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n')
                        language_run.font.size = Pt(font_size)
            except:
                pass
    
    doc.save(path_save)
    print("Conversion completed !!")
    
