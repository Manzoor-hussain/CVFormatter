import os
import openai
import docx
import docx2txt
from .keys import api_key
from pprint import pprint
import json
import re
import textwrap
import PyPDF2
import pdfplumber
from docx.shared import Pt
from math import ceil
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def get_completion(prompt, model="gpt-3.5-turbo"):
    messages = [{"role": "user", "content": prompt}]
    response = openai.ChatCompletion.create(
        model=model,
        messages=messages,
        temperature=0, # this is the degree of randomness of the model's output
    )
    return response.choices[0].message["content"]


def leo_partner_converter(path, path_out, path_save):
    formatted = path_out
    
    
    try:
        with open(path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            unformated_text = ""
            for i in range (len(pdf_reader.pages)):
                first_page = pdf_reader.pages[i]
                unformated_text += first_page.extract_text() + " "
            print('Its PDF')
    except:
        try:
            unformated_text = docx2txt.process(path)
            print('Its Docx')
        except:
            print('WE DONT SUPPORT THIS TYPE OF FILE')
    
    formatted_text = docx2txt.process(formatted)
    
    print("----------------------------------------------------------------")
    print("                          Unformatted Text                            ")
    print("----------------------------------------------------------------")
    print(unformated_text)
    
    print("Process has started...")
    
        # Prompt
    os.environ["OPEN_API_KEY"] = api_key
    openai.api_key = api_key
    
    test_text = """

    Ectract data from this text:

    \"""" + unformated_text + """\"

    in following JSON format:
    {
    "Name":"candidate name",
    "Profile" : "value",

    "Professional Experiennce" : [
        {"Company Name" : "Name of company",
        "Location":"Location of that company",
        "Designation" : "Specific designation in that Company",
        "Duration" : "Working duration in company",
        "Responsibilities" : ["Responsibility 1", "Responsibility 2", ...]
        },
        
        {"Company Name" : "Name of company",
        "Location":"Location of that company",
        "Designation" : "Specific designation in that Company",
        "Duration" : "Working duration in company",
        "Responsibilities" : ["Responsibility 1", "Responsibility 2", ...]
        },
        ...
        ],
    "Education" : [
        {"Institute Name":"Name of that institute",
        "Location":"Location of that institute ",
        "Degree":"Name of that degree",
        "Duration":"duration of that degree"
        },
        {"Institute Name":"Name of that institute",
        "Location":"Location of that institute ",
        "Degree":"Name of that degree",
        "Duration":"duration of that degree"
        },
        ...
        ],
    "Skills" : ["skill1", "skill2", ...],
    "Relevant Qualifications" : ["relevant Qualifications1", "relevant qualifications2", ...],
    "Certification":["certification1","certification2",...],
    "Languages":["languages1","languages2",...],
    "Interests":["interests1","interests2",...]
    }

    You must keep the following points in considration while extracting data from text:
        1. Do NOT split, rephrase or summarize list of Responsibilities. Extract each Responsibility as a complete sentence from text.
        2. Make it sure to keep the response in JSON format.
        3. If value not found then leave it empty/blank.
        4. Do not include Mobile number, Email and Home address.
        5. Summary/Personal Statement should be as it is. Do not change or rephrase it.

        """


    # Prompt result
    result = get_completion(test_text)
    
    print("----------------------------------------------------------------")
    print("                          Result                            ")
    print("----------------------------------------------------------------")
    print(result)
    
    dc = dict(json.loads(re.sub(r'\[\"\"\]',r'[]',re.sub(r'\"[Un]nknown\"|\"[Nn]one\"|\"[Nn]ull\"|\"[Nn]ot [Mm]entioned\"',r'""',re.sub(r',[ \n]*\]',r']',re.sub(r',[ \n]*\}',r'}',result.replace('...','')))))))
    
    print("----------------------------------------------------------------")
    print("                          Dictionary                            ")
    print("----------------------------------------------------------------")
    print(dc)
    
    
    doc = docx.Document(formatted)

    for i,p in enumerate(doc.paragraphs):
        try:
            if p.text.strip().lower() == 'name':
                if dc['Name'] and dc['Name'].lower().replace(' ','') != 'value':
                    doc.paragraphs[i].text = ""
                    run = doc.paragraphs[i].add_run(dc['Name'].strip().title())
                    run.bold = True
        except:
            pass
        
        try:
            if p.text.strip(' :\n').lower() == 'profile':
                if dc['Profile'] and dc['Profile'].lower().replace(' ','') != 'value':
                    doc.paragraphs[i+2].add_run(dc['Profile'].strip()).bold = False
#                     doc.paragraphs[i+2].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'education':
                for j in dc['Education']:
                    if j['Degree'].strip() and j['Degree'].lower().replace(' ','') != 'nameofthatdegree':
                        j["Institute Name"] = j["Institute Name"].title()
                        a = len(j['Institute Name'])

                        b = len(j['Degree'])

                        c = "\t\t\t"
                        spc = " "
                        spec = " "

                        if a > b:
                            d = ceil((a - b) * 1.8)


                            spc += " " * d
                        else:
                            a < b
                            e = ceil((b - a)*1.8)

                            spec += " " * e
                        if j['Institute Name'].strip() and j['Institute Name'].lower().replace(' ','') != 'nameofthatinstitute':    
                            run1 = doc.paragraphs[i+2].add_run(j['Institute Name'].strip()+ spec +c + j['Location'].strip()+"\n").bold=True
                        
                        run3 = doc.paragraphs[i+2].add_run(j['Degree'].strip()+ spc + c + j['Duration'].strip()+"\n\n")
                        run3.bold=False
                        run3.italic=True
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'certification':
                if dc['Certification'][0] and dc['Certification'][0].lower().strip() != 'certification1':
                    for j in dc['Certification']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'languages':
                if dc['Languages'][0] and dc['Languages'][0].lower().strip() != 'langauge1':
                    for j in dc['Languages']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'interests':
                if dc['Interests'][0] and dc['Interests'][0].lower().strip() != 'interest1':
                    for j in dc['Interests']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'training':
                if dc['Training'][0] and dc['Training'][0].lower().strip() != 'training1':
                    for j in dc['Training']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'skills':
                if dc['Skills'][0] and dc['Skills'][0].lower().strip() != 'skill1':
                    for j in dc['Skills']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        if p.text.strip(' :\n').lower() == 'professional experience':
            for j in dc['Professional Experience']:
                j['Company Name'] = j["Company Name"].title()
                a = len(j['Company Name'])
              
                b = len(j['Designation'])
              
                c = "\t\t\t"
                spc = " "
                spec = " "

                if a > b:
                    d = ceil((a - b) * 2)
                  
                    spc += " " * d
                else:
                    a < b
                    e = ceil((b - a)*1.75)
                   
                    spec += " " * e
                
                if (j['Company Name'] and j['Company Name'].lower().replace(' ','') != 'nameofcompany') or (j['Designation'] and j['Designation'].lower().replace(' ','') != 'specificdesignationinthatcompany'):
                    
                    if j['Company Name'] and j['Company Name'].lower().replace(' ','') != 'nameofcompany':
                        run1=doc.paragraphs[i+2].add_run(j['Company Name'].strip()+ spec + c + j['Location'].strip()+"\n").bold=True
                    else:
                        doc.paragraphs[i+2].add_run('(' + "Company Name not mentioned" + ')' + '\n').bold = True
                    
                    if j['Designation'] and j['Designation'].lower().replace(' ','') != 'specificdesignationinthatcompany':
                        run = doc.paragraphs[i + 2].add_run(j['Designation'].strip()+ spc + c )
                        run.font.bold = True
                    else:
                        doc.paragraphs[i+2].add_run('(' + "Designation not mentioned" + ')' + '\n').bold = True
                    
                    if j['Duration'] and j['Duration'].lower().replace(' ','') != 'workingdurationincompany':
                        run = doc.paragraphs[i + 2].add_run(j['Duration'].strip()+"\n\n")
                        run.font.bold = False
                    else:
                        doc.paragraphs[i+2].add_run('(' + "Duration not mentioned" + ')' + '\n').bold = True

                    if j["Responsibilities"] and j["Responsibilities"][0].lower().replace(' ','') != "responsibility 1":          
                        for k in j['Responsibilities']:
                            if k.strip():
                                doc.paragraphs[i+2].add_run('  • ' + k.strip() + '\n').bold = False
                        doc.paragraphs[i+2].add_run('\n\n')
                    else:
                        doc.paragraphs[i+2].add_run('No responsibility is mentioned' + '\n\n').bold = False



    doc.save(path_save)
    
    print("Process has Completed...")
