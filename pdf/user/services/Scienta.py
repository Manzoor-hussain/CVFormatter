import os
import openai
import docx
import docx2txt
import re
import json
from .keys import api_key
from docx.enum.text import WD_UNDERLINE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
import PyPDF2


def get_completion(prompt, model="gpt-3.5-turbo"):
    messages = [{"role": "user", "content": prompt}]
    response = openai.ChatCompletion.create(
        model=model,
        messages=messages,
        temperature=0, # this is the degree of randomness of the model's output
    )
    return response.choices[0].message["content"]

def scienta_converter(path,formatted_path,save_path):
    formatted= formatted_path
    
    try:
        with open(path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            unformated_text = ""
            for i in range (len(pdf_reader.pages)):
                first_page = pdf_reader.pages[i]
                unformated_text += first_page.extract_text() + " "
            print('Its PDF')
    except:
        try:
            unformated_text = docx2txt.process(path)
            print('Its Docx')
        except:
            print('WE DONT SUPPORT THIS TYPE OF FILE')
        

    os.environ["OPEN_API_KEY"] = api_key
    openai.api_key = api_key

    
    print ("Process has Started...")
    test_text = """

    Ectract data from this text:

    \"""" + unformated_text + """\"

    in following JSON format:
{
"Name":"value"
"Profile" : "value",

"Professional Experience" : [
    {"Company Name" : "Name of company",
    "Company location": "Location of that company",
    "Duration" : "Working duration in company",
    "Designation" : "Specific designation in that Company",
    "Responsibilities" : ["Responsibility 1", "Responsibility 2", ...],
    },
    
    {"Company Name" : "Name of company",
    "Company location": "Location of that company",
    "Duration" : "Working duration in company",
    "Designation" : "Specific designation in that Company",
    "Responsibilities" : ["Responsibility 1", "Responsibility 2", ...],
    },
    ...
    ]
"Education" : [
    {"Institute" : "Name Of institute",
    "Duration" : "Studying duration in institute",
    "Degree": "Name of degree",
    },
    {"Institute" : "Name Of institute",
    "Duration" : "Studying duration in institute",
    "Degree": "Name of degree",
    },
    ...
    ],
"Training" : ["training1", "training2", ...],
"Skills" : ["skill1", "skill2", ...],
"Certificates" : ["certificate1", "certificate2", ...],
"Languages" : ["language1", "language2", ...],
"Interests" : ["interest1", "interest2", ...],
}

   You must keep the following points in considration while extracting data from text:
    1. Do NOT split, rephrase or summarize list of Responsibilities. Extract each Responsibility as a complete sentence from text.
    2. Make it sure to keep the response in JSON format.
    3. If value not found then leave it empty/blank.
    4. Do not include Mobile number, Email and Home address.
    5. Summary/Personal Statement should be complete without being rephrased.

    """


    result = get_completion(test_text)

    dc = dict(json.loads(re.sub(r'\[\"\"\]',r'[]',re.sub(r'\"[Un]nknown\"|\"[Nn]one\"|\"[Nn]ull\"|\"[Nn]ot [Mm]entioned\"',r'""',re.sub(r',[ \n]*\]',r']',re.sub(r',[ \n]*\}',r'}',result.replace('...','')))))))
    doc = docx.Document(formatted)

    for i,p in enumerate(doc.paragraphs):
        try:
             if p.text.strip().lower() == 'name':
                    doc.paragraphs[i].text = ""
                    if dc['Name'].lower().replace(' ','') != 'value':                
                        run = doc.paragraphs[i].add_run(dc['Name'].strip().title())
                        run.bold = True
                        run.font.size = Pt(16.5)
        except:
            pass
        try:
            if p.text.strip(' :\n').lower() == 'profile':
                if dc['Profile'].lower().replace(' ','') != 'value':
                    doc.paragraphs[i+2].add_run(dc['Profile'].strip()).bold = False
                    doc.paragraphs[i+2].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        except:
            pass

        if p.text.strip(' :\n').lower() == 'education':
            for j in dc['Education']:
                if j['Degree'].strip() and j['Degree'].lower().replace(' ','') != "nameofthatdegree":
                    if j['Degree'].strip():
                        doc.paragraphs[i+2].add_run('  • '+j['Degree'].strip()).bold=False
                    if j['Duration'].strip():
                         doc.paragraphs[i+2].add_run("                 "+j['Duration'].strip() + "\n").bold=False
                    else:
                        doc.paragraphs[i+2].add_run("Duration not mentioned").bold=False                  
                    if j['Institute'].strip():   
                        doc.paragraphs[i+2].add_run('  • '+j['Institute'].strip() + "\n\n").bold=False
                    else:
                        doc.paragraphs[i+2].add_run("Institute not mentioned").bold=False
                        
    #             doc.paragraphs[i+2].add_run(j['Thesis'].strip() + "\n\n").bold=False

    #               + j['Institute'].strip() + "–" + j['Duration'].strip()).bold = False
    #               doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False



        try:
            if p.text.strip(' :\n').lower() == 'certificates':
                if dc['Certificates'][0].lower().replace(' ','') != 'certificate1':
                    for j in dc['Certificates']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'languages':
                if dc['Languages'][0].lower().replace(' ','') != 'language1':
                    for j in dc['Languages']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'interests':
                if dc['Interests'][0].lower().replace(' ','') != 'interest1':
                    for j in dc['Interests']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'training':
                if dc['Training'][0].lower().replace(' ','') != 'training1':
                    for j in dc['Training']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'skills':
                if dc['Skills'][0].lower().replace(' ','') != 'skill1':
                    for j in dc['Skills']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'professional experience':
                for j in dc['Professional Experience']:
                    if (j['Designation'].strip() and j['Designation'].lower().replace(' ','') !='specificdesignationinthatcompany') or (j['Company Name'].strip() and j['Company Name'].lower().replace(' ','') !='nameofcompany'):
                        if j['Company Name'].strip():                                        
                            doc.paragraphs[i+2].add_run(j['Company Name'].strip()).bold=True
                        else:
                            doc.paragraphs[i+2].add_run("Company Name mentioned" + '\n').bold=True
                        if j['Duration'].strip():                                          
                            doc.paragraphs[i+2].add_run('                        '+ j['Duration'].strip() + '\n').bold = True
                        else:
                            doc.paragraphs[i+2].add_run("Duration not mentioned" + '\n').bold=True
                        if j['Designation'].strip():    
                            doc.paragraphs[i+2].add_run(j['Designation'].strip() + '\n\n').bold = True  
                        else:
                            doc.paragraphs[i+2].add_run("Designation not mentioned"+ '\n\n').bold=True
                        
                        if j['Responsibilities'] and j['Responsibilities'][0].lower().replace(' ','') != 'responsibility1':
                            for k in j['Responsibilities']:
                                if k.strip():
                                    doc.paragraphs[i+2].add_run('  • ' + k.strip() + '\n').bold = False
                            doc.paragraphs[i+2].add_run('\n')
        except:
            pass
    
    doc.save(save_path)
    print("Process has Completed...")
