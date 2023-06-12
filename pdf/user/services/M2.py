import os
import re
import openai
import docx
import docx2txt
import traceback
import json
from docx.shared import Pt
import PyPDF2 
from .keys import api_key
from docx2python import docx2python


def month_parser(date):
    date = date.strip()
    date = date.replace('january','Jan').replace('January','Jan').replace('February', 'Feb').replace('february', 'Feb').replace('March', 'Mar').replace('march', 'Mar').replace('April', 'Apr').replace('april', 'Apr').replace('may', 'May').replace('June', 'Jun').replace('june', 'Jun').replace('July', 'Jul').replace('july', 'Jul').replace('August', 'Aug').replace('august', 'Aug').replace('September', 'Sep').replace('september', 'Sep').replace('October', 'Oct').replace('october', 'Oct').replace('November', 'Nov').replace('november', 'Nov').replace('December', 'Dec').replace('december', 'Dec').replace('01','Jan').replace('02', 'Feb').replace('03', 'Mar').replace('04', 'Apr').replace('05', 'May').replace('06', 'Jun').replace('07', 'Jul').replace('08', 'Aug').replace('09', 'Sep').replace('10', 'Oct').replace('11', 'Nov').replace('12', 'Dec').replace('1','Jan').replace('2', 'Feb').replace('3', 'Mar').replace('4', 'Apr').replace('5', 'May').replace('6', 'Jun').replace('7', 'Jul').replace('8', 'Aug').replace('9', 'Sep')
    return date

def year_parser(date):
    date = date.strip()
    if len(date) == 2:
        if int(date) <= 35:
            date = "20" + date
        else:
            date = "19" + date
    return date

def get_completion(prompt, model="gpt-3.5-turbo"):
    messages = [{"role": "user", "content": prompt}]
    response = openai.ChatCompletion.create(
        model=model,
        messages=messages,
        temperature=0, # this is the degree of randomness of the model's output
    )
    return response.choices[0].message["content"]

def m2_partnership_converter(path,pathout,path_save):
    formatted = pathout
    un_formatted = path
    formated_text = docx2txt.process(formatted)
    
    
    try:
        with open(path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            unformated_text = ""
            for i in range (len(pdf_reader.pages)):
                first_page = pdf_reader.pages[i]
                unformated_text += first_page.extract_text() + " "
            print('Its PDF')
    except:
        try:
            unformated_text = docx2txt.process(path)
            print('Its Docx')
        except:
            print('WE DONT SUPPORT THIS TYPE OF FILE')
    
    print("Process has Started...")
    test_text = """
    Extract data from this text:

    \"""" + re.sub('\n+','\n', unformated_text) + """\"

    in following JSON format:
    {
    "Candidate Name" : "value",
    "Work Experience" : [
        {"Company Name" : "Name of company",
        "Company Location" : "Location of company",
        "Designation" : "Designation in that company",
        "Duration" : {"start_month": "month when person started working at company", "start_year": "year when person started working at company", "end_month": "month when person ended working at company", "end_year": "year when person ended working at company"},
        "Responsibilities" : ["Responsibility 1", "Responsibility 2", ...]
        },
        {"Company Name" : "Name of company",
        "Company Location" : "Location of company",
        "Designation" : "Designation in that company",
        "Duration" : {"start_month": "month when person started working at company", "start_year": "year when person started working at company", "end_month": "month when person ended working at company", "end_year": "year when person ended working at company"}
        "Responsibilities" : ["Responsibility 1", "Responsibility 2", ...]
        },
        ...
        ],
    "Education" : [
        {"Institute Name" : "Name Of institute",
        "Degree Name": "Name of degree",
        "Duration" : {"start_month": "month when person started studying at institute", "start_year": "year when person started studying at institute", "end_month": "month when person ended studying at institute", "end_year": "year when person ended studying at institute"}
        },
        {"Institute Name" : "Name Of institute",
        "Degree Name": "Name of degree",
        "Duration" : {"start_month": "month when person started studying at institute", "start_year": "year when person started studying at institute", "end_month": "month when person ended studying at institute", "end_year": "year when person ended studying at institute"}
        },
        ...
        ],
    "Awards" : ["award 1", "award 2", ...],
    "Skills" : ["skill 1", "skill 2", ...],
    "Projects" : ["project 1", "project 2", ...],
    "Languages" : ["language 1", "language 2", ...],
    "Interests" : ["interest 1", "interest 2", ...],
    "Trainings" : ["training 1", "training 2", ...],
    "Hobbies" : ["hobby 1", "hobby 2", ...]
    }

    You must keep the following points in considration while extracting data from text:
        1. Do NOT split, rephrase or summarize list of Responsibilities. Extract each Responsibility as a complete sentence from text.
        2. Make it sure to keep the response in JSON format.
        3. If value not found then leave it empty/blank.
        4. Do not include Mobile number, Email and Home address.
    """

    os.environ["OPEN_API_KEY"] = api_key
    openai.api_key = api_key
    result = get_completion(test_text)
    print('\n\n\n')
    print(result)
    dc = dict(json.loads(re.sub(r'\[\"\"\]',r'[]',re.sub(r'\"[Un]nknown\"|\"[Nn]one\"|\"[Nn]ull\"|\"[Nn]ot [Mm]entioned\"',r'""',re.sub(r',[ \n]*\]',r']',re.sub(r',[ \n]*\}',r'}',result.replace('...','')))))))
    print('\n\n\n')
    print(dc)
    # Open the existing document
    doc = docx.Document(formatted)

    for table in doc.tables:
#         print(table)
        for row in table.rows:
            for cell in row.cells:
                for tb in cell.tables:
                    for i,r in enumerate(tb.rows):
                        try:
                            if r.cells[0].text.strip(' :\n').lower() == 'name':
                                tb.cell(i,0).text = ""
                                if dc['Candidate Name'] and dc['Candidate Name'].lower().replace(' ','') != "value":
                                    para = tb.cell(i,0).add_paragraph()
                                    para.style.font.name = "Century Gothic"
                                    run = para.add_run(dc['Candidate Name'].strip() + '\n')
                                    run.font.size = Pt(22)
                                    run.bold = True
                        except:
                            pass
                            
                        try:
                            if r.cells[0].text.strip(' :\n').lower() == 'e education':
                                para = tb.cell(i+1,0).add_paragraph()
                                para.style.font.name = "Century Gothic"
                                for j in dc['Education']:
                                    if j['Degree Name'] and j['Degree Name'].lower().replace(' ','') != "nameofdegree":
                                        try:
                                            if j['Institute Name'].strip():
                                                para.add_run(j['Institute Name'].strip() + '\n').bold = True
                                            else:
                                                para.add_run('Institute Name not mentioned\n').bold = True
                                        except:
                                            para.add_run('Institute Name not mentioned\n').bold = True
                                        try:
                                            duration = ""
                                            if j['Duration']['start_year'].strip() and (len(j['Duration']['start_year'].strip()) == 2 or len(j['Duration']['start_year'].strip())==4):
                                                if j['Duration']['start_month'].strip() and len(j['Duration']['start_year']) < 20:
                                                    duration += month_parser(j['Duration']['start_month']) + " "

                                                duration += year_parser(j['Duration']['start_year'])

                                            if j['Duration']['end_year'].strip() and (len(j['Duration']['end_year'].strip())==2 or len(j['Duration']['end_year'].strip())==4):
                                                if duration:
                                                        duration += ' - '
                                                if j['Duration']['end_month'].strip() and len(j['Duration']['end_month']) < 20:
                                                    duration += month_parser(j['Duration']['end_month']) + " "

                                                duration += year_parser(j['Duration']['end_year'])
                                            if duration:
                                                para.add_run(duration + '\n').bold = True
                                            else:
                                                para.add_run('Duration not mentioned\n').bold = True
                                        except:
                                            para.add_run('Duration not mentioned\n').bold = True
                                        try:
                                            if j['Degree Name'].strip():
                                                para.add_run(j['Degree Name'].strip() + '\n\n').bold = False
                                            else:
                                                para.add_run('Degree Name not mentioned\n\n').bold = False
                                        except:
                                            para.add_run('Degree Name not mentioned\n\n').bold = False                           
                        except:
                            pass 
                        
                        try:
                            if r.cells[0].text.strip(' :\n').lower() == 'e experience':
                                para = tb.cell(i+1,0).add_paragraph()
                                para.style.font.name = "Century Gothic"
                                for j in dc['Work Experience']:
                                    if (j['Designation'] and j['Designation'].lower().replace(' ','') != "designationinthatcompany") or (j['Company Name'] and j['Company Name'].lower().replace(' ','') != "nameofcompany"):
                                        try:
                                            if j['Company Name'].strip():
                                                para.add_run(j['Company Name'].strip()).bold = True
                                            else:
                                                para.add_run('Name not mentioned').bold = True
                                        except:
                                            para.add_run('Name not mentioned').bold = True
                                        try:
                                            if j['Company Location'].strip():
                                                para.add_run( ', ' + j['Company Location\n'].strip()).bold = False
                                            else:
                                                para.add_run('\n').bold = False
                                        except:
                                            para.add_run('\n').bold = False

                                        try:
                                            if j['Designation'].strip():
                                                para.add_run(j['Designation'].strip()).bold = True
                                            else:
                                                para.add_run('Designation not mentioned').bold = True
                                        except:
                                            para.add_run('Designation not mentioned').bold = True
                                        try:
                                            duration = ""
                                            if j['Duration']['start_year'].strip() and (len(j['Duration']['start_year'].strip())==2 or len(j['Duration']['start_year'].strip())==4):
                                                if j['Duration']['start_month'].strip() and len(j['Duration']['start_year']) < 20:
                                                    duration += month_parser(j['Duration']['start_month']) + " "

                                                duration += year_parser(j['Duration']['start_year'])

                                            if j['Duration']['end_year'].strip() and (len(j['Duration']['end_year'].strip())==2 or len(j['Duration']['end_year'].strip())==4):
                                                if duration:
                                                        duration += ' - '
                                                if j['Duration']['end_month'].strip() and len(j['Duration']['end_month']) < 20:
                                                    duration += month_parser(j['Duration']['end_month']) + " "

                                                duration += year_parser(j['Duration']['end_year'])

                                            if duration:
                                                para.add_run(' | ' + duration + '\n').bold = True
                                            else:
                                                para.add_run(' | Duration not mentioned\n').bold = True

                                        except:
                                            para.add_run('Duration not mentioned\n').bold = True


                                        try:
                                            if j['Responsibilities'] and j['Responsibilities'][0].lower().replace(' ','') != 'responsibility1':
                                                for k in j['Responsibilities']:
                                                    para.add_run('    • ' + k.strip() + '\n').bold = False
                                        except:
                                            pass
                                        para.add_run('\n\n').bold = False
                        except:
                            pass 
                        
                        try:
                            if r.cells[0].text.strip(' :\n').lower() == 's skills':
                                if dc['Skills'] and dc['Skills'][0].lower().replace(' ','') != "skill1":
                                    para = tb.cell(i+1,0).add_paragraph()
                                    para.style.font.name = "Century Gothic"
                                    for j in dc['Skills']:
                                        if j.strip():
                                            para.add_run('    • ' + j.strip() + '\n').bold = False
                        except:
                            pass
                            
                        try:
                            if r.cells[0].text.strip(' :\n').lower() == 'h hobbies':
                                if dc['Hobbies'] and dc['Hobbies'][0].lower().replace(' ','') != "hobby1":
                                    para = tb.cell(i+1,0).add_paragraph()
                                    para.style.font.name = "Century Gothic"
                                    for j in dc['Hobbies']:
                                        if j.strip():
                                            para.add_run('    • ' + j.strip() + '\n').bold = False
                        except:
                            pass                         
                        
                        try:
                            if r.cells[0].text.strip(' :\n').lower() == 'k key projects':
                                if dc['Projects'] and dc['Projects'][0].lower().replace(' ','') != "project1":
                                    para = tb.cell(i+1,0).add_paragraph()
                                    para.style.font.name = "Century Gothic"
                                    for j in dc['Projects']:
                                        if j.strip():
                                            para.add_run(j.strip() + '\n').bold = False
                        except:
                            pass 
                        
                        try:
                            if r.cells[0].text.strip(' :\n').lower() == 'a awards':
                                if dc['Awards'] and dc['Awards'][0].lower().replace(' ','') != "award1":
                                    para = tb.cell(i+1,0).add_paragraph()
                                    para.style.font.name = "Century Gothic"
                                    for j in dc['Awards']:
                                        if j.strip():
                                            para.add_run(j.strip() + '\n').bold = False
                        except:
                            pass 
                        
                        try:
                            if r.cells[0].text.strip(' :\n').lower() == 'l languages':
                                if dc['Languages'] and dc['Languages'][0].lower().replace(' ','') != "language1":
                                    para = tb.cell(i+1,0).add_paragraph()
                                    para.style.font.name = "Century Gothic"
                                    for j in dc['Languages']:
                                        if j.strip():
                                            para.add_run('    • ' + j.strip() + '\n').bold = False
                        except:
                            pass 
                        
                        try:
                            if r.cells[0].text.strip(' :\n').lower() == 'i interests':
                                if dc['Interests'] and dc['Interests'][0].lower().replace(' ','') != "interest1":
                                    para = tb.cell(i+1,0).add_paragraph()
                                    para.style.font.name = "Century Gothic"
                                    for j in dc['Interests']:
                                        if j.strip():
                                            para.add_run('    • ' + j.strip() + '\n').bold = False
                        except:
                            pass 
                        
                        try:
                            if r.cells[0].text.strip(' :\n').lower() == 't trainings':
                                if dc['Trainings'] and dc['Trainings'][0].lower().replace(' ','') != "training1":
                                    para = tb.cell(i+1,0).add_paragraph()
                                    para.style.font.name = "Century Gothic"
                                    for j in dc['Trainings']:
                                        if j.strip():
                                            para.add_run(j.strip() + '\n').bold = False
                        except:
                            pass                  



    # Save the updated document as a new file
    doc.save(path_save)
    print("Process has Completed...")
    
