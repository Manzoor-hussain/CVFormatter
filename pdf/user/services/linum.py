import os
import openai
import docx
import docx2txt
from .keys import api_key
from pprint import pprint
import json
import re
import textwrap
import PyPDF2
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

import openai

openai.api_key = api_key

def get_completion(prompt, model="gpt-3.5-turbo"):
    messages = [{"role": "user", "content": prompt}]
    response = openai.ChatCompletion.create(
        model=model,
        messages=messages,
        temperature=0, # this is the degree of randomness of the model's output
    )
    return response.choices[0].message["content"]
    
    
def linum_converter(path,formatted_path,save_path):

    # path to formatted file
    formatted = formatted_path
    
    try:
        with open(path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            unformated_text = ""
            for i in range (len(pdf_reader.pages)):
                first_page = pdf_reader.pages[i]
                unformated_text += first_page.extract_text() + " "
            print('Its PDF')
    except:
        try:
            unformated_text = docx2txt.process(path)
            print('Its Docx')
        except:
            print('WE DONT SUPPORT THIS TYPE OF FILE')
    
    formatted_text = docx2txt.process(formatted)
   
    print("----------------------------------------------------------------")
    print("                          Unformatted Text                            ")
    print("----------------------------------------------------------------")
    print(unformated_text)
    
    print("Process has started...")

    test_text = """
    Extract data from this text:

    \"""" + re.sub('\n+','\n', unformated_text) + """\"

    in following JSON format:
    {
    "Candidate Name" : "value",
    "Position Applied For" : "value",
    "Notice Period" : "value",
    "Current Location" : "value",
    "Current Salary" : "value",
    "Expected Salary" : "value",
    "Marital Status" : "value",
    "Nationality" : "value",
    "Date of Birth" : "value",
    "Summary" : "value",
    "Work Experience" : [
        {"Company Name" : "Name of company",
        "Company Location" : "Location of company",
        "Duration" : "Working Duration in Company",
        "Designation" : "Specific designation in that Company",
        "Responsibilities" : ["Responsibility 1", "Responsibility 2", ...]
        },
        {"Company Name" : "Name of company",
        "Company Location" : "Location of company",
        "Duration" : "Working Duration in Company",
        "Designation" : "Specific designation in that Company",
        "Responsibilities" : ["Responsibility 1", "Responsibility 2", ...]
        },
        ...
        ],
    "Education" : [
        {"Institute Name" : "Name Of institute",
        "Degree Name": "Name of degree",
        "Duration" : "Studying duration in institute",
        },
        {"Institute Name" : "Name Of institute",
        "Degree Name": "Name of degree",
        "Duration" : "Studying duration in institute",
        },
        ...
        ],
    "Achievements" : ["achievement 1", "achievement 2", ...],
    "Qualifications" : ["qualification 1", "qualification 2", ...],
    "Skills" : ["skill 1", "skill 2", ...],
    "Attributes" : ["attribute 1", "attribute 2", ...],
    "Languages" : ["language 1", "language 2", ...],
    "Interests" : ["interest 1", "interest 2", ...]
    }

    You must keep the following points in considration while extracting data from text:
        1. Do NOT split, rephrase or summarize list of Responsibilities. Extract each Responsibility as a complete sentence from text.
        2. Make it sure to keep the response in JSON format.
        3. If value not found then leave it empty/blank.
        4. Do not include Mobile number, Email and Home address.
        5. Summary/Personal Statement should be as it is. Do not change or rephrase it.
    """


    
    result = get_completion(test_text)
    
        
    print("----------------------------------------------------------------")
    print("                          Result                            ")
    print("----------------------------------------------------------------")
    print(result)
    
    dc = dict(json.loads(re.sub(r'\[\"\"\]',r'[]',re.sub(r'\"[Un]nknown\"|\"[Nn]one\"|\"[Nn]ull\"|\"[Nn]ot [Mm]entioned\"',r'""',re.sub(r',[ \n]*\]',r']',re.sub(r',[ \n]*\}',r'}',result.replace('...','')))))))
    
    print("----------------------------------------------------------------")
    print("                          Dictionary                            ")
    print("----------------------------------------------------------------")
    print(dc)
    
    # Open the existing document
    doc = docx.Document(formatted)
    
    # Get the first paragraph
    for i,p in enumerate(doc.paragraphs):
        
#         doc.paragraphs[i].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        
        if p.text.strip(' :\n').lower() == 'summary':
            try:
                if dc['Summary'].strip() and dc['Summary'].lower().replace(' ','') != 'value':
                    doc.paragraphs[i+2].add_run(dc['Summary']).bold = False
                    doc.paragraphs[i+2].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            except:
                pass
        if p.text.strip(' :\n').lower() == 'skills':
            try:
                if dc['Skills'][0].lower().replace(' ','') != 'skill1':
                    for j in dc['Skills']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('    • ' + j.strip() + '\n').bold = False
            except:
                pass
        if p.text.strip(' :\n').lower() == 'attributes':
            try:
                if dc['Attributes'][0].lower().replace(' ','') != 'attribute1':
                    for j in dc['Attributes']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('    • ' + j.strip() + '\n').bold = False
            except:
                pass
        if p.text.strip(' :\n').lower() == 'achievements':
            try:
                if dc['Achievements'][0].lower().replace(' ','') != 'achievement1':
                    for j in dc['Achievements']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('    • ' + j.strip() + '\n').bold = False
            except:
                pass        
        if p.text.strip(' :\n').lower() == 'languages':
            try:
                if dc['Languages'][0].lower().replace(' ','') != 'language1':
                    for j in dc['Languages']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('    • ' + j.strip() + '\n').bold = False
            except:
                pass     
        if p.text.strip(' :\n').lower() == 'interests':
            try:
                if dc['Interests'][0].lower().replace(' ','') != 'interest1':
                    for j in dc['Interests']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('    • ' + j.strip() + '\n').bold = False
            except:
                pass        
        if p.text.strip(' :\n').lower() == 'work experience':
            try:
                for j in dc['Work Experience']:
                    try:
                        if (j['Company Name'].strip() and j['Company Name'].lower().replace(' ','') != 'nameofcompany') or (j['Designation'].strip() and j['Designation'].lower().replace(' ','') != 'specificdesignationinthatcompany'):
                            try:
                                if (j['Duration'].strip() and j['Duration'].lower().replace(' ','') != 'workingdurationincompany'):
                                    doc.paragraphs[i+2].add_run(j['Duration'].strip() + '\n').bold = False
                                else:
                                    doc.paragraphs[i+2].add_run('Duration not mentioned\n').bold = False

                            except:
                                doc.paragraphs[i+2].add_run('Duration not mentioned\n').bold = False
                            try:
                                if (j['Company Name'].strip() and j['Company Name'].lower().replace(' ','') != 'nameofcompany'):
                                    doc.paragraphs[i+2].add_run(j['Company Name'].strip() + '\n').bold = True
                                else:
                                    doc.paragraphs[i+2].add_run('Company not mentioned\n').bold = True
                            except:
                                doc.paragraphs[i+2].add_run('Company not mentioned\n').bold = True
                            try:
                                if (j['Designation'].strip() and j['Designation'].lower().replace(' ','') != 'specificdesignationinthatcompany'):
                                    doc.paragraphs[i+2].add_run(j['Designation'].strip() + '\n\n').bold = False
                                else:
                                    doc.paragraphs[i+2].add_run('Designation not mentioned\n\n').bold = False
                            except:
                                doc.paragraphs[i+2].add_run('Designation not mentioned\n\n').bold = False

                            try:
                                if j["Responsibilities"][0].strip() and j["Responsibilities"][0].lower().replace(' ','') != "responsibility1":
                                    doc.paragraphs[i+2].add_run('Duties:' + '\n').bold = False
                                    for k in j['Responsibilities']:
                                        if k.strip():
                                            doc.paragraphs[i+2].add_run('    • ' + k.strip() + '\n').bold = False
                                    doc.paragraphs[i+2].add_run('\n').bold = False
                            except:
                                pass
                            doc.paragraphs[i+2].add_run('\n').bold = False
                    except:
                        pass
            except:
                pass

        if p.text.strip(' :\n').lower() == 'education':
            try:
                for j in dc['Education']:
                    try:
                        if j['Degree Name'].strip() and j['Degree Name'].lower().replace(' ','') != 'nameofdegree':
                            
                            try:
                                if j['Duration'].strip() and j['Duration'].lower().replace(' ','') != 'studyingdurationininstitute':
                                    doc.paragraphs[i+2].add_run(j['Duration'].strip() + '\n').bold = False
                                else:
                                    doc.paragraphs[i+2].add_run('Duration not mentioned\n').bold = False
                            except:
                                doc.paragraphs[i+2].add_run('Duration not mentioned\n').bold = False

                            doc.paragraphs[i+2].add_run(j['Degree Name'].strip() + '\n').bold = True

                            try:
                                if j['Institute Name'].strip() and j['Institute Name'].lower().replace(' ','') != 'nameofinstitute':
                                    doc.paragraphs[i+2].add_run(j['Institute Name'].strip() + '\n\n').bold = False
                                else:
                                    doc.paragraphs[i+2].add_run('Institute Name not mentioned\n\n').bold = False
                            except:
                                doc.paragraphs[i+2].add_run('Institute Name not mentioned\n\n').bold = False
                    except:
                        pass
            except:
                pass



    for table in doc.tables:
        for row in table.rows:
            for i,cell in enumerate(row.cells):
                try:
                    if cell.text.strip(' :\n').lower() == 'position applied for':
                        if dc['Position Applied For'].lower().replace(' ','') != 'value':
                            row.cells[i+1].text = dc['Position Applied For']
                except:
                    pass
                try:
                    if cell.text.strip(' :\n').lower() == 'current location':
                        if dc['Current Location'].lower().replace(' ','') != 'value':
                            row.cells[i+1].text = dc['Current Location']
                except:
                    pass
                try:
                    if cell.text.strip(' :\n').lower() == 'qualifications':
                        if dc['Qualifications'][0].lower().strip() != 'qualification1':
                            for j in dc['Qualifications']:
                                if j.sprip():
                                    row.cells[i+1].text = row.cells[i+1].text + j + '\n'
                except:
                    pass                
                try:
                    if cell.text.strip(' :\n').lower() == 'current salary':
                        if dc['Current Salary'].lower().replace(' ','') != 'value':
                            row.cells[i+1].text = dc['Current Salary']
                except:
                    pass                
                try:
                    if cell.text.strip(' :\n').lower() == 'expected salary':
                        if dc['Expected Salary'].lower().replace(' ','') != 'value':
                            row.cells[i+1].text = dc['Expected Salary']
                except:
                    pass                
                try:
                    if cell.text.strip(' :\n').lower() == 'availability / notice period':
                        if dc['Notice Period'].lower().replace(' ','') != 'value':
                            row.cells[i+1].text = dc['Notice Period']
                except:
                    pass
                try:
                    if cell.text.strip(' :\n').lower() == 'personal details':
                        try:
                            if dc['Marital Status'].lower().replace(' ','') != 'value':
                                row.cells[i+1].text = re.sub('MARITAL STATUS:','MARITAL STATUS: ' + dc['Marital Status'] ,row.cells[i+1].text)
                        except:
                            pass
                        try:
                            if dc['Nationality'].lower().replace(' ','') != 'value':
                                row.cells[i+1].text = re.sub('NATIONALITY:','NATIONALITY: ' + dc['Nationality'] ,row.cells[i+1].text)
                        except:
                            pass
                        try:
                            if dc['Date of Birth'].lower().replace(' ','') != 'value':
                                row.cells[i+1].text = re.sub('D.O.B:','D.O.B: ' + dc['Date of Birth'] ,row.cells[i+1].text)
                        except:
                            pass
                except:
                    pass

    # Save the updated document as a new file
    doc.save(save_path)

    print("Conversion Completed...")
