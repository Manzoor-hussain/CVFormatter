import os
import openai
import docx
import docx2txt
from .keys import api_key
from pprint import pprint
import json
import re
import textwrap
import PyPDF2
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def get_completion(prompt, model="gpt-3.5-turbo"):
    messages = [{"role": "user", "content": prompt}]
    response = openai.ChatCompletion.create(
        model=model,
        messages=messages,
        temperature=0, # this is the degree of randomness of the model's output
    )
    return response.choices[0].message["content"]


# Functions to check whether the unformatted file is a docx or pdf
def read_text_from_docx(file_path):
    doc = docx.Document(file_path)
    text = [paragraph.text for paragraph in doc.paragraphs]
    return '\n'.join(text)

def read_text_from_pdf(file_path):
    with open(file_path, 'rb') as file:
        pdf_reader = PyPDF2.PdfReader(file)
        text = []
        for page in pdf_reader.pages:
            text.append(page.extract_text())
        return '\n'.join(text)


def clarus_converter(path_in, path_out, path_save):
    
    formatted= path_out
    
    
    if path_in.endswith('.docx'):
        unformatted_text = read_text_from_docx(path_in)
    elif path_in.endswith('.pdf'):
        unformatted_text = read_text_from_pdf(path_in)
    else:
        error = 'Format not supported.'
        print(error)
    
    formatted_text = docx2txt.process(formatted)
    
    print("----------------------------------------------------------------")
    print("                          Unformatted Text                            ")
    print("----------------------------------------------------------------")
    print(unformatted_text)
    
    
    print("Process has started...")
    
    # Prompt
    os.environ["OPEN_API_KEY"] = api_key
    openai.api_key = api_key

    test_text = """

    Extract data from this text:
    \"""" + unformatted_text + """\"
    in following JSON format:
    {
    "Name" : "value",
    "Personal Statement" : "value",
    "Education" : [
        {"Institute Name" : "Name Of institute",
        "Degree Name": "Name of degree",
        "Duration" : "Studying duration in institute",
        },
        {"Institute Name" : "Name Of institute",
        "Degree Name": "Name of degree",
        "Duration" : "Studying duration in institute",
        },
        ...
        ],
    "Qualification" : ["Qualification1", "Qualification2", ...],
    "Employment Summary" : [
        {"Company Name" : "Name of company",
        "Job Title" : "Title of job",
        "Duration" : "Working Duration in Company",
        "Responsibilities" : ["Responsibility 1", "Responsibility 2", ...],
        },
        {"Company Name" : "Name of company",
        "Job Title" : "Title of job",
        "Duration" : "Working Duration in Company",
        "Responsibilities" : ["Responsibility 1", "Responsibility 2", ...],
        },
        ...
        ],
    "Other Experience" : ["Other Experience1", "Other Experience2", ...],
    "Projects and Exhibitions" : ["Projects and Exhibitions1", "Projects and Exhibitions2", ...],
    "Voluntary Experience/Work" : ["Voluntary Experience/Work1", "Voluntary Experience/Work2", ...],
    "Skills" : ["Skills1", "Skills2", ...],
    "Languages" : ["Language1", "Language2", ...],
    "Leadership" : ["Leadership1", "Leadership2", ...],
    "Interests" : ["interest1", "interest2", ...]
    }
    
    You must keep the following points in considration while extracting data from text:
        1. Do NOT split, rephrase or summarize list of Responsibilities. Extract each Responsibility as a complete sentence from text.
        2. Make it sure to keep the response in JSON format.
        3. If value not found then leave it empty/blank.
        4. Do not include Mobile number, Email and Home address.
        5. Summary/Personal Statement should be as it is. Do not change or rephrase it.
    """
    # Prompt result
    result = get_completion(test_text)
    
    print("----------------------------------------------------------------")
    print("                          Result                            ")
    print("----------------------------------------------------------------")
    print(result)
    
    dc = dict(json.loads(re.sub(r'\[\"\"\]',r'[]',re.sub(r'\"[Un]nknown\"|\"[Nn]one\"|\"[Nn]ull\"|\"[Nn]ot [Mm]entioned\"',r'""',re.sub(r',[ \n]*\]',r']',re.sub(r',[ \n]*\}',r'}',result.replace('...','')))))))

    print("----------------------------------------------------------------")
    print("                          Dictionary                            ")
    print("----------------------------------------------------------------")
    print(dc)
    
    
    doc = docx.Document(formatted)
    font_size = 16
    for i,p in enumerate(doc.paragraphs):


        if p.text.strip(' :\n').lower() == 'name':
            try:
                if dc['Name'].strip() and dc['Name'].lower().replace(' ','') != 'value':
                    name_paragraph = doc.paragraphs[i]
                    name_paragraph.text = str(dc['Name'])
                    name_paragraph.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
                    name_paragraph.runs[0].font.size = Pt(font_size)
            except:
                pass


        if p.text.strip(' :\n').lower() == 'personal statement':
            try:
                if dc['Personal Statement'].strip() and dc['Personal Statement'].lower().replace(' ','') != 'value':
                    doc.paragraphs[i+2].text = str(dc['Personal Statement'])
                    doc.paragraphs[i+2].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            except:
                pass


        if p.text.strip(' :\n').lower() == 'education':
            try:
                for j in dc['Education']:
                    institute_name = ""
                    duration = ""
                    degree_name = ""
                    if j['Degree Name'].strip() and j['Degree Name'].lower().replace(' ','') != 'nameofdegree':
                        degree_name = j['Degree Name'].strip()
                        if j['Institute Name'].strip() and j['Institute Name'].lower().replace(' ','') != 'nameofinstitute':
                            institute_name = j['Institute Name'].strip()
                        else:
                            institute_name = "Institute not mentioned"
                        if j['Duration'].strip() and j['Duration'].lower().replace(' ','') != 'studyingdurationininstitute':
                            duration = j['Duration'].strip()
                        else:
                            duration = "Duration not mentioned"
                            
                        doc.paragraphs[i+2].add_run(institute_name + ' ').bold = True
                        doc.paragraphs[i+2].add_run('(' + duration + ')' + '\n').bold = True
                        doc.paragraphs[i+2].add_run(degree_name + '\n\n').bold = False
            except:
                pass


        if p.text.strip(' :\n').lower() == 'qualification':
            try:
                if dc['Qualifications'][0].lower().replace(' ','') != 'qualification1':
                    for j in dc['Qualifications']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
            except:
                pass


        if p.text.strip(' :\n').lower() == 'employment summary':
            try:
                for j in dc['Employment Summary']:
                    try:
                        company_name = ""
                        duration = ""
                        job_title = ""

                        if (j['Company Name'].strip() and j['Company Name'].lower().replace(' ','') != 'nameofcompany') or (j['Job Title'].strip() and j['Job Title'].lower().replace(' ','') != 'titleofjob'):
                            if (j['Company Name'].strip() and j['Company Name'].lower().replace(' ','') != 'nameofcompany'):
                                company_name = j['Company Name'].strip()
                            else:
                                company_name = "Company not mentioned"

                            if (j['Duration'].strip() and j['Duration'].lower().replace(' ','') != 'workingdurationincompany'):
                                duration = j['Duration'].strip()
                            else:
                                duration = "Duration not mentioned"

                            if (j['Job Title'].strip() and j['Job Title'].lower().replace(' ','') != 'titleofjob'):
                                job_title = j['Job Title'].strip()
                            else:
                                job_title = "Job not mentioned"


                            doc.paragraphs[i+2].add_run(company_name + ' ').bold = True
                            doc.paragraphs[i+2].add_run('(' + duration + ')' + '\n').bold = True
                            doc.paragraphs[i+2].add_run(job_title + '\n\n').bold = False
                            try:
                                if j["Responsibilities"][0].lower().replace(' ','') != "responsibility1":
                                    for k in j['Responsibilities']:
                                        if k.strip():
                                            doc.paragraphs[i+2].add_run('  • ' + k.strip() + '\n').bold = False
                                    doc.paragraphs[i+2].add_run("\n\n")
                            except:
                                pass
                    except:
                        pass
            except:
                pass


        if p.text.strip(' :\n').lower() == 'other experience':
            try:
                if dc['Other Experience'][0].lower().replace(' ','') != 'otherexperience1':
                    for j in dc['Other Experience']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
            except:
                pass


        if p.text.strip(' :\n').lower() == 'projects and exhibitions':
            try:
                if dc['Projects and Exhibitions'][0].lower().replace(' ','') != 'projectsandexhibitions1':
                    for j in dc['Projects and Exhibitions']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
            except:
                pass

        if p.text.strip(' :\n').lower() == 'voluntary experience/work':
            try:
                if dc['Voluntary Experience/Work'][0].lower().replace(' ','') != 'voluntaryexperience/work1':
                    for j in dc['Voluntary Experience/Work']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
            except:
                pass

        if p.text.strip(' :\n').lower() == 'skills':
            try:
                if dc['Skills'][0].lower().replace(' ','') != 'skills1':
                    for j in dc['Skills']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
            except:
                pass

        if p.text.strip(' :\n').lower() == 'languages':
            try:
                if dc['Languages'][0].lower().replace(' ','') != 'language1':
                    for j in dc['Languages']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
            except:
                pass

        if p.text.strip(' :\n').lower() == 'leadership':
            try:
                if dc['Leadership'][0].lower().replace(' ','') != 'leadership1':
                    for j in dc['Leadership']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
            except:
                pass

        if p.text.strip(' :\n').lower() == 'interests':
            try:
                if dc['Interests'][0].lower().replace(' ','') != 'interest1':
                    for j in dc['Interests']:
                        if j.strip():
                            doc.paragraphs[i+1].add_run('\n  • ' + j.strip()).bold = False
            except:
                pass


    doc.save(path_save)
    print("Conversion has completed !!")
    
