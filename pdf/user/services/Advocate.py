import os
import openai
import docx
import docx2txt
import re
import json
from .keys import api_key
from docx.enum.text import WD_UNDERLINE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
import PyPDF2
import pdfplumber
import traceback


def get_completion(prompt, model="gpt-3.5-turbo"):
    messages = [{"role": "user", "content": prompt}]
    response = openai.ChatCompletion.create(
        model=model,
        messages=messages,
        temperature=0, # this is the degree of randomness of the model's output
    )
    return response.choices[0].message["content"]

def advocate_group_converter(path, pathout, path_save):
    formatted= pathout

    try:
        with open(path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            unformated_text = ""
            for i in range (len(pdf_reader.pages)):
                first_page = pdf_reader.pages[i]
                unformated_text += first_page.extract_text() + " "
            print('Its PDF')
    except:
        try:
            unformated_text = docx2txt.process(path)
            print('Its Docx')
        except:
            print('WE DONT SUPPORT THIS TYPE OF FILE')
        

    os.environ["OPEN_API_KEY"] = api_key
    openai.api_key = api_key

    
    print ("Process has Started...")
    test_text = """

    Ectract data from this text:

    \"""" + unformated_text + """\"

    in following JSON format:
    {
    "Name":"value"
    "Summary" : "value",

    "Experience" : [
        {"Company Name" : "Name of company",
         "Duration" : "Working Duration in Company",
         "Designation" : "Specific designation in that Company",
         "Responsibilities" : ["Responsibility1", "Responsibility2", ...],
        },

        {"Company Name" : "Name of company",
         "Duration" : "Working Duration in Company",
         "Designation" : "Specific designation in that Company",
         "Responsibilities" : ["Responsibility1", "Responsibility2", ...],
        },
        ...
        ]
    "Education" : [
        {"Institute Name":"Name of that institute",
         "Institute location":"Location of that institute",
         "Degree":"Name of that degree",
         "Duration":""Studying duration in institute",
        },
        {"Institute Name":"Name of that institute",
         "Institute location":"Location of that institute",
         "Degree":"Name of that degree",
         "Duration":"Studying duration in institute.",
        },
    ...],
    "Training" : ["training1", "training2", ...],
    "Skills" : ["skill1", "skill2", ...],
    "Qualification" : ["qualification1", "qualification2", ...],
    "Languages" : ["language1", "language2", ...],
    "Interests" : ["interest1", "interest2", ...],
    }

   You must keep the following points in considration while extracting data from text:
    1. Do NOT split, rephrase or summarize list of Responsibilities. Extract each Responsibility as a complete sentence from text.
    2. Make it sure to keep the response in JSON format.
    3. If value not found then leave it empty/blank.
    4. Do not include Mobile number, Email and Home address.
    5. Summary/Personal Statement should be complete without being rephrased.
    """


    result = get_completion(test_text)

    dc = dict(json.loads(re.sub(r'\[\"\"\]',r'[]',re.sub(r'\"[Un]nknown\"|\"[Nn]one\"|\"[Nn]ull\"|\"[Nn]ot [Mm]entioned\"',r'""',re.sub(r',[ \n]*\]',r']',re.sub(r',[ \n]*\}',r'}',result.replace('...','')))))))
#     print(dc)


    doc = docx.Document(formatted)

    for i,p in enumerate(doc.paragraphs):
        try:
            if p.text.strip().lower() == 'name':
                doc.paragraphs[i].text = ""
                if dc['Name'].lower().replace(' ','')!='value':
                    run = doc.paragraphs[i].add_run(dc['Name'].strip().title())
                    run.bold = True
                    run.font.size = Pt(16.5)
        except:
            pass
        try:
            if p.text.strip(' :\n').lower() == 'summary':
                if dc['Summary'].lower().replace(' ','') != 'value':
                    doc.paragraphs[i+2].add_run(dc['Summary'].strip()).bold = False
                    doc.paragraphs[i+2].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        except:
            pass

        if p.text.strip(' :\n').lower() == 'education':
            for j in dc['Education']:
                degree=j['Degree'].strip()
                institute=j['Institute Name'].strip()
                duration=j['Duration'].strip()
                if degree and degree.lower().replace(' ','') != "nameofthatdegree":
                    if institute:
                        doc.paragraphs[i+2].add_run(institute).bold=False
                    else:
                        doc.paragraphs[i+2].add_run("Institute not mentioned").bold=False
                        
                    doc.paragraphs[i+2].add_run(" - "+degree + "\n").bold=False
                    
                    if duration:
                        doc.paragraphs[i+2].add_run(duration + "\n\n").bold=False
                    else:
                        doc.paragraphs[i+2].add_run("Duration not mentioned" + "\n\n").bold=False

        try:
            if p.text.strip(' :\n').lower() == 'qualification':
                if dc['Qualification'][0].lower().replace(' ','') != 'qualification1':
                    for j in dc['Qualification']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'languages':
                if dc['Languages'][0].lower().replace(' ','') != 'languages1':
                    for j in dc['Languages']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'interests':
                if dc['Interests'][0].lower().replace(' ','') != 'interests1':
                    for j in dc['Interests']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'training':
                if dc['Training'][0].lower().replace(' ','') != 'training1':
                    for j in dc['Training']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'skills':
                if dc['Skills'][0].lower().replace(' ','') != 'skills1':
                    for j in dc['Skills']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'experience':
                for j in dc['Experience']:
                    company_name = j['Company Name'].strip()
                    duration = j['Duration'].strip()
                    designation = j['Designation'].strip()
                    if designation and designation.lower().replace(' ','') !='specificdesignationinthatcompany' or (company_name and company_name.lower().replace(' ','') !='nameofcompany'):                        
                        if designation:
                            doc.paragraphs[i+2].add_run(designation + '\n').bold = True
                        else:                       
                            doc.paragraphs[i+2].add_run("Designation not mentioned").bold=False

                        if company_name:
                            doc.paragraphs[i+2].add_run(company_name+'\n').bold=True
                        else:
                            doc.paragraphs[i+2].add_run("Company Name not mentioned").bold=False

                        if duration:
                            doc.paragraphs[i+2].add_run(duration + "\n\n").bold=True
                        else:
                            doc.paragraphs[i+2].add_run("Duration not mentioned" + "\n\n").bold=False   

                        if j['Responsibilities'] and j['Responsibilities'][0].lower().replace(' ','') != 'responsibility1':
                            for k in j['Responsibilities']:
                                doc.paragraphs[i+2].add_run('  • ' + k.strip() + '\n').bold = False
                            doc.paragraphs[i+2].add_run('\n\n')
        except:
            pass

    doc.save(path_save)
    print("Process has Completed...")
