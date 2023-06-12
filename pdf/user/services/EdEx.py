import os
import openai
import docx
import docx2txt
import re
import json
import PyPDF2
from docx.shared import Pt
from docx.enum.text import WD_UNDERLINE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from .keys import api_key


def get_completion(prompt, model="gpt-3.5-turbo"):
    messages = [{"role": "user", "content": prompt}]
    response = openai.ChatCompletion.create(
        model=model,
        messages=messages,
        temperature=0, # this is the degree of randomness of the model's output
    )
    return response.choices[0].message["content"]


def edex_converter(path, pathout, path_save):
    formatted= pathout
    un_formatted = path
    formated_text = docx2txt.process(formatted)

    try:
        with open(path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            unformated_text = ""
            for i in range (len(pdf_reader.pages)):
                first_page = pdf_reader.pages[i]
                unformated_text += first_page.extract_text() + " "
            print('Its PDF')
    except:
        try:
            unformated_text = docx2txt.process(path)
            print('Its Docx')
        except:
            print('WE DONT SUPPORT THIS TYPE OF FILE')

    
    print("----------------------------------------------------------------")
    print("                          Unformatted Text                            ")
    print("----------------------------------------------------------------")
    print(unformated_text)            
            
            
    os.environ["OPEN_API_KEY"] = api_key
    openai.api_key = api_key

    print("Process has Started...")

    test_text = """

    Ectract data from this text:

    \"""" + unformated_text + """\"

    in following JSON format:
    {
    "Name" : "value",
    "Profile" : "value",

    "Education" : [
        {"Institute Name" : "Name Of institute",
        "Duration" : "Studying duration in institute",
        "Degree Name": "Name of degree",
        "Details" : ["Detail 1", "Detail 2", ...]
        },
        {"Institute Name" : "Name Of institute",
        "Duration" : "Studying duration in institute",
        "Degree Name": "Name of degree",
        "Details" : ["Detail 1", "Detail 2", ...]
        },
        ...
    ],
    "It Literacy" : ["literacy1", "literacy2", ...],
    "Certificates" : ["certificate1", "certificate2", ...],
    "Projects" : ["project1", "project2", ...],
    "Professional Qualifications" : ["qualification1", "qualification2", ...],
    "Softwares" : ["software1", "software2", ...],
    "Languages" : ["language1", "language2", ...],
    "Interests" : ["interest1", "interest2", ...],
    "Trainings" : ["training1", "training2", ...],
    "Skills" : ["skill1", "skill2", ...],
    "Work Experience" : [
        {"Company Name" : "Name of company",
        "Duration" :  "Working Duration in Company",
        "Designation" : "Specific designation in that Company",
        "Responsibilities" : ["Responsibility 1", "Responsibility 2", ...],
        },
        {"Company Name" : "Name of company",
        "Duration" :  "Working Duration in Company",
        "Designation" : "Specific designation in that Company",
        "Responsibilities" : ["Responsibility 1", "Responsibility 2", ...],
        },
        ...
        ]
    }

    You must keep the following points in considration while extracting data from text:
        1. Do NOT split, rephrase or summarize list of Responsibilities. Extract each Responsibility as a complete sentence from text.
        2. Make it sure to keep the response in JSON format.
        3. If value not found then leave it empty/blank.
        4. Do not include Mobile number, Email and Home address.
        5. Summary/Personal Statement should be as it is. Do not change or rephrase it.
    """


    result = get_completion(test_text)
    
    print("----------------------------------------------------------------")
    print("                          Result                            ")
    print("----------------------------------------------------------------")
    print(result)
    
    dc = dict(json.loads(re.sub(r'\[\"\"\]',r'[]',re.sub(r'\"[Un]nknown\"|\"[Nn]one\"|\"[Nn]ull\"|\"[Nn]ot [Mm]entioned\"',r'""',re.sub(r',[ \n]*\]',r']',re.sub(r',[ \n]*\}',r'}',result.replace('...','')))))))
    
    print("----------------------------------------------------------------")
    print("                          Dictionary                            ")
    print("----------------------------------------------------------------")
    print(dc)

    doc = docx.Document(formatted)

    for i,p in enumerate(doc.paragraphs):

        try:
            if p.text.strip().lower() == 'name:':
                if dc['Name'] and dc['Name'][0].lower().strip() != 'value':
                    doc.paragraphs[i].text = ""
                    run = doc.paragraphs[i].add_run(dc['Name'].strip().title())
                    run.bold = True
                    run.font.size = Pt(14)
        except:
               pass
        try:
            if p.text.strip(' :\n').lower() == 'profile':
                if dc['Profile'] and dc['Profile'][0].lower().strip() != 'value':
                    doc.paragraphs[i+2].add_run(dc['Profile'].strip()).bold = False
                    doc.paragraphs[i+2].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        except:
            pass

        try:        
            if p.text.strip(' :\n').lower() == 'education':
                for j in dc['Education']:
                    institute_name = j['Institute Name'].strip()
                    duration = j['Duration'].strip()
                    degree_name = j['Degree Name'].strip()
                    
                    if j['Degree Name'].strip() and j['Degree Name'].lower().replace(' ','') != 'nameofdegree': 
                        if j['Institute Name'].strip() and j['Institute Name'].lower().replace(' ','') != 'nameofinstitute':
                            doc.paragraphs[i+2].add_run(institute_name + ' ').bold = True
                        if j['Duration'].strip() and j['Duration'].lower().replace(' ','') != 'studyingdurationininstitute':
                            doc.paragraphs[i+2].add_run('(' + duration + ')' + '\n').bold = True
                        else:
                            doc.paragraphs[i+2].add_run('(' + "Not mentioned" + ')' + '\n').bold = True
                            
                        doc.paragraphs[i+2].add_run(degree_name + '\n\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'it literacy':
                if dc['It Literacy'][0] and dc['It Literacy'][0].lower().strip() != 'literacy1':
                    for j in dc['It Literacy']:
                        doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'certificates':
                if dc['Cerificates'][0] and dc['Cerificates'][0].lower().strip() != 'certificate1':
                    for j in dc['Certificates']:
                        doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'projects':
                if dc['Projects'][0] and dc['Projects'][0].lower().strip() != 'project1':
                    for j in dc['Projects']:
                        doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass
        

        try:
            if p.text.strip(' :\n').lower() == 'professional qualifications':
                if dc['Professional Qualifications'][0] and dc['Professional Qualifications'][0].lower().strip() != 'qualification1':
                    for j in dc['Professional Qualifications']:
                        doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'softwares':
                if dc['Softwares'][0] and dc['Softwares'][0].lower().strip() != 'software1':     
                    for j in dc['Softwares']:
                        doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'languages':
                if dc['Languages'][0] and dc['Languages'][0].lower().strip() != 'language1':
                    for j in dc['Languages']:
                        doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'interests':
                if dc['Interests'][0] and dc['Interests'][0].lower().strip() != 'interest1':
                    for j in dc['Interests']:
                        doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'trainings':
                if dc['Trainings'][0] and dc['Trainings'][0].lower().strip() != 'training1':
                    for j in dc['Trainings']:
                        doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'skills':
                if dc['Skills'][0] and dc['Skills'][0].lower().strip() != 'skill1':
                    for j in dc['Skills']:
                        doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'work experience':
                for j in dc['Work Experience']:
                    company_name = j['Company Name'].strip()
                    duration = j['Duration'].strip()
                    job_title = j['Designation'].strip()
                    
                    if (j['Company Name'] and j['Company Name'].lower().replace(' ','') != 'nameofcompany') or (j['Designation'] and j['Designation'].lower().replace(' ','') != 'specificdesignationinthatcompany'):
                        if j['Company Name'] and j['Company Name'].lower().replace(' ','') != 'nameofcompany':  
                            doc.paragraphs[i+2].add_run(company_name + ' ').bold = True
                        if j['Duration'] and j['Duration'].lower().replace(' ','') != 'workingdurationincompany':    
                            doc.paragraphs[i+2].add_run('(' + duration + ')' + '\n').bold = True
                        if j['Designation'] and j['Designation'].lower().replace(' ','') != 'specificdesignationinthatcompany':
                            doc.paragraphs[i+2].add_run(job_title + '\n\n').bold = False

                        if j["Responsibilities"] and j["Responsibilities"][0].lower().replace(' ','') != "responsibility1":
                            for k in j['Responsibilities']:
                                doc.paragraphs[i+2].add_run('  • ' + k.strip() + '\n')
                            doc.paragraphs[i+2].add_run('\n')
        except:
            pass

    doc.save(path_save)
    print("Process has Completed...")
