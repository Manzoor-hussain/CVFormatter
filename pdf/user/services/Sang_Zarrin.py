import os
import openai
import docx
import re
import json
from pprint import pprint
import docx2txt
import PyPDF2
from docx.enum.text import WD_UNDERLINE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from .keys import api_key



def get_completion(prompt, model="gpt-3.5-turbo"):
    messages = [{"role": "user", "content": prompt}]
    response = openai.ChatCompletion.create(
        model=model,
        messages=messages,
        temperature=0, # this is the degree of randomness of the model's output
    )
    return response.choices[0].message["content"]


def sang_zarrin_converter(path, path_out, path_save):
    
    formatted= path_out
    
    
    try:
        with open(path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            unformated_text = ""
            for i in range (len(pdf_reader.pages)):
                first_page = pdf_reader.pages[i]
                unformated_text += first_page.extract_text() + " "
            print('Its PDF')
    except:
        try:
            unformated_text = docx2txt.process(path)
            print('Its Docx')
        except:
            print('WE DONT SUPPORT THIS TYPE OF FILE')
    
    formatted_text = docx2txt.process(formatted)
    
    print("----------------------------------------------------------------")
    print("                          Unformatted Text                            ")
    print("----------------------------------------------------------------")
    print(unformatted_text)
    
    print("Process has started...")

    
    # Prompt
    os.environ["OPEN_API_KEY"] = api_key
    openai.api_key = api_key
    
    
    test_text = """

    Ectract data from this text:

    \"""" + unformatted_text + """\"

    in following JSON format:
    {
    "Name" : "value",
    "Summary" : "value",
    "Experience" : [
        {"Designation" : "The specific designation or position on which he works in this company",
        "Company Name" : "Name of company",
        "Location" : "Location of company",
        "Duration" : "Working Duration in Company",
        "Responsibilities" : ["Responsibility 1", "Responsibility 2", ...]
        },
        {"Designation" : "The specific designation or position on which he works in this company",
        "Company Name" : "Name of company",
        "Location" : "Location of company",
        "Duration" : "Working Duration in Company",
        "Responsibilities" : ["Responsibility 1", "Responsibility 2", ...]
        },
        ...
        ],
    "Education" : [
        {"Institute Name" : "Name Of institute",
        "Degree Name": "Name of degree",
        "Institute Location" : "Location of Institute",
        "Duration" : "Studying duration in institute",
        },
        {"Institute Name" : "Name Of institute",
        "Degree Name": "Name of degree",
        "Institute Location" : "Location of Institute",
        "Duration" : "Studying duration in institute",
        },
        ...
        ],
    "Trainings" : ["trainings1", "trainings2", ...],
    "Computer Skills" : ["computer skill1", "computer skill2", ...],
    "Skills" : ["skill1", "skill2", ...],
    "Qualifications" : ["qualification1", "qualification2", ...],
    "Languages" : ["language1", "language2", ...],
    "Interests" : ["interest1", "interest2", ...]
    }
    
    You must keep the following points in considration while extracting data from text:
        1. Do NOT split, rephrase or summarize list of Responsibilities. Extract each Responsibility as a complete sentence from text.
        2. Make it sure to keep the response in JSON format.
        3. If value not found then leave it empty/blank.
        4. Do not include Mobile number, Email and Home address.
        5. Summary/Personal Statement should be as it is. Do not change or rephrase it.

    """


    # Prompt result
    result = get_completion(test_text)
    
    print("----------------------------------------------------------------")
    print("                          Result                            ")
    print("----------------------------------------------------------------")
    print(result)
    
    dc = dict(json.loads(re.sub(r'\[\"\"\]',r'[]',re.sub(r'\"[Un]nknown\"|\"[Nn]one\"|\"[Nn]ull\"|\"[Nn]ot [Mm]entioned\"',r'""',re.sub(r',[ \n]*\]',r']',re.sub(r',[ \n]*\}',r'}',result.replace('...','')))))))
   
    print("----------------------------------------------------------------")
    print("                          Dictionary                            ")
    print("----------------------------------------------------------------")
    print(dc)
    
    doc = docx.Document(formatted)

    for i,p in enumerate(doc.paragraphs):
        try:
            if p.text.strip(' :\n').lower() == 'summary':
                if dc['Summary'] and dc['Summary'].lower().replace(' ','') != 'value':
                    doc.paragraphs[i+2].add_run(dc['Summary'].strip()).bold = False
                    doc.paragraphs[i+2].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'experience':
                for j in dc['Experience']:
                    company_name = j['Company Name'].strip()
                    duration = j['Duration'].strip()
                    designation = j['Designation'].strip()
                    location = j['Location'].strip()
                    
                    if (j['Company Name'] and j['Company Name'].lower().replace(' ','') != 'name of company') or (j['Designation'] and j['Designation'].lower().replace(' ','') != 'designation'):
                        
                        if j['Designation'] and j['Designation'].lower().replace(' ','') != 'designation':
                            doc.paragraphs[i+2].add_run(designation + '\n').bold = False
                        if j['Company Name'] and j['Company Name'].lower().replace(' ','') != 'name of company':  
                            doc.paragraphs[i+2].add_run(company_name + '\n').bold = True
                        if j['Location'] and j['Location'].lower().replace(' ','') != 'location of company':  
                            doc.paragraphs[i+2].add_run(location + '\n').bold = True    
                        if j['Duration'] and j['Duration'].lower().replace(' ','') != 'working duration in company':    
                            doc.paragraphs[i+2].add_run('(' + duration + ')' + '\n').bold = True
                            
                        doc.paragraphs[i+2].add_run('\n')
                        if j["Responsibilities"] and j["Responsibilities"][0].lower().replace(' ','') != "responsibility1":
                            for k in j['Responsibilities']:
                                doc.paragraphs[i+2].add_run('  • ' + k.strip() + '\n').bold = False
                            doc.paragraphs[i+2].add_run('\n\n')
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'education':
                for j in dc['Education']:
                    institute_name = j['Institute Name'].strip()
                    duration = j['Duration'].strip()
                    degree_name = j['Degree Name'].strip()
                    
                    if j['Degree Name'].strip() and j['Degree Name'].lower().replace(' ','') != 'name of degree': 
                        if j['Institute Name'].strip() and j['Institute Name'].lower().replace(' ','') != 'name of institute':
                            doc.paragraphs[i+2].add_run(institute_name + ' ').bold = True
                        if j['Duration'].strip() and j['Duration'].lower().replace(' ','') != 'studying duration in institute':
                            doc.paragraphs[i+2].add_run('(' + duration + ')' + '\n').bold = True
                        else:
                            doc.paragraphs[i+2].add_run('(' + "Not mentioned" + ')' + '\n').bold = True
                            
                        doc.paragraphs[i+2].add_run(degree_name + '\n\n').bold = False
        except:
            pass

        try: 
            if p.text.strip(' :\n').lower() == 'trainings':
                if dc['Trainings'][0] and dc['Trainings'][0].lower().strip() != 'training1':
                    for j in dc['Trainings']:
                        doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'computer skills':
                if dc['Computer Skills'][0] and dc['Computer Skills'][0].lower().strip() != 'computer skill1':
                    for j in dc['Computer Skills']:
                        doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'skills':
                if dc['Skills'][0] and dc['Skills'][0].lower().strip() != 'skill1':
                    for j in dc['Skills']:
                        doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'qualifications':
                if dc['Qualifications'][0] and dc['Qualifications'][0].lower().strip() != 'qualification1':
                    for j in dc['Qualifications']:
                        doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'languages':
                if dc['Languages'][0] and dc['Languages'][0].lower().strip() != 'language1':
                    for j in dc['Languages']:
                        doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'interests':
                if dc['Interests'][0] and dc['Interests'][0].lower().strip() != 'interest1':
                    for j in dc['Interests']:
                        doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

    doc.save(path_save)
    print("Process has Completed...")
    
    
