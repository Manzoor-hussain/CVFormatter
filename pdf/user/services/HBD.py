import os
import openai
import docx
import docx2txt
from .keys import api_key
from pprint import pprint
import json
import re
import textwrap
import PyPDF2
import pdfplumber
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def get_completion(prompt, model="gpt-3.5-turbo"):
    messages = [{"role": "user", "content": prompt}]
    response = openai.ChatCompletion.create(
        model=model,
        messages=messages,
        temperature=0, # this is the degree of randomness of the model's output
    )
    return response.choices[0].message["content"]
    
def hbd_converter(path, path_out, path_save):
    
    formatted = path_out
    
    try:
        with open(path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            unformated_text = ""
            for i in range (len(pdf_reader.pages)):
                first_page = pdf_reader.pages[i]
                unformated_text += first_page.extract_text() + " "
            print('Its PDF')
    except:
        try:
            unformated_text = docx2txt.process(path)
            print('Its Docx')
        except:
            print('WE DONT SUPPORT THIS TYPE OF FILE')

    # formatted document
    formatted_text = docx2txt.process(formatted)
    
    print("----------------------------------------------------------------")
    print("                          Unformatted Text                            ")
    print("----------------------------------------------------------------")
    print(unformated_text)
    
    print("Process has started...")
    
    os.environ["OPEN_API_KEY"] = api_key
    openai.api_key = api_key
    
    test_text = """

    Ectract data from this text:

    \"""" + unformated_text + """\"

    in following JSON format:
    {
    "Name" : "value",
    "Profile" : "value",

    "Education" : [
        {"Institute Name" : "Name Of institute",
        "Degree Name": "Name of degree",
        "Duration" : "Studying duration in institute.",
        "Study Details" : ["Study Detail1", "Study Detail2", ...],
        },
        {"Institute Name" : "Name Of institute",
        "Degree Name": "Name of degree",
        "Duration" : "Studying duration in institute.",
        "Study Details" : ["Study Detail1", "Study Detail2", ...],
        },
        ...
        ],
    "Certificates" : ["certificate1", "certificate2", ...],
    "Achievements" : ["achievement1", "achievement2", ...],
    "Qualifications" : ["qualification1", "qualification2", ...],
    "Computer Skills" : ["computer skill1", "computer skill2", ...],
    "Expertise" : ["expertise1", "expertise2", ...],
    "Languages" : ["language1", "language2", ...],
    "Interests" : ["interest1", "interest2", ...],
    "Trainings" : ["training1", "training2", ...],
    "Skills" : ["skill1", "skill2", ...],
    "Work Experience" : [
        {"Company Name" : "Name of company",
        "Job Title" : "Title of job",
        "Duration" : "Working Duration in Company in (Mon YYYY - Mon YYYY) Format.",
        "Duties" : ["Duty 1", "Duty 2", ...],
        },
        {"Company Name" : "Name of company",
        "Job Title" : "Title of job",
        "Duration" : "Working Duration in Company in (Mon YYYY - Mon YYYY) Format.",
        "Duties" : ["Duty 1", "Duty 2", ...],
        },
        ...
        ],
    }

    You must keep the following points in considration while extracting data from text:
        1. Do NOT split, rephrase or summarize list of Responsibilities. Extract each Responsibility as a complete sentence from text.
        2. Make it sure to keep the response in JSON format.
        3. If value not found then leave it empty/blank.
        4. Do not include Mobile number, Email and Home address. 
        5. Summary/Personal Statement should be as it is. Do not change or rephrase it.
    """


    result = get_completion(test_text)
    
    print("----------------------------------------------------------------")
    print("                          Result                            ")
    print("----------------------------------------------------------------")
    print(result)
    
    dc = dict(json.loads(re.sub(r'\[\"\"\]',r'[]',re.sub(r'\"[Un]nknown\"|\"[Nn]one\"|\"[Nn]ull\"|\"[Nn]ot [Mm]entioned\"',r'""',re.sub(r',[ \n]*\]',r']',re.sub(r',[ \n]*\}',r'}',result.replace('...','')))))))
    
    print("----------------------------------------------------------------")
    print("                          Dictionary                            ")
    print("----------------------------------------------------------------")
    print(dc)
    
    doc = docx.Document(formatted)

    for i,p in enumerate(doc.paragraphs):

        try:        
            if p.text.strip().lower() == 'name':
                if dc['Name'] and dc['Name'].lower().replace(' ','') != 'value':
                    doc.paragraphs[i].text = ""
                    run = doc.paragraphs[i+1].add_run(dc['Name'].strip().title())
                    run.bold = True
                    run.font.size = Pt(16)

        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'profile':
                if dc['Profile'] and dc['Profile'].lower().replace(' ','') != 'value':
                    doc.paragraphs[i+2].add_run(dc['Profile'].strip()).bold = False
#                     doc.paragraphs[i+2].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        except:
            pass
        
        try:
            if p.text.strip(' :\n').lower() == 'education':
                for j in dc['Education']:
                    institute_name = j['Institute Name'].strip()
                    duration = j['Duration'].strip()
                    degree_name = j['Degree Name'].strip()

                    if j['Degree Name'].strip() and j['Degree Name'].lower().replace(' ','') != 'nameofdegree':
                        if j['Duration'].strip() and j['Duration'].lower().replace(' ','') != 'studyingdurationininstitute':
                            duration_run = doc.paragraphs[i+2].add_run(duration + '\n')
#                             duration_run.font.italic = True
                        else:
                            duration_run = doc.paragraphs[i+2].add_run("Not mentioned" + '\n').bold = True
#                             duration_run.font.italic = True

                        if j['Institute Name'].strip() and j['Institute Name'].lower().replace(' ','') != 'nameofinstitute':
                            doc.paragraphs[i+2].add_run(institute_name + '\n').bold = True

                        if j['Degree Name'].strip() and j['Degree Name'].lower().replace(' ','') != 'nameofdegree':
                            degree_run = doc.paragraphs[i+2].add_run(degree_name + '\n\n')
#                             degree_run.font.underline = True
                        else:
                            degree_run = doc.paragraphs[i+2].add_run("Not mentioned" + '\n\n')
#                             degree_run.font.underline = True


        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'expertise':
                if dc['Expertise'] and dc['Expertise'][0].lower().strip() != 'expertise1':
                    for j in dc['Expertise']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'certificates':
                if dc['Certificates'] and dc['Certificates'][0].lower().strip() != 'certificate1':
                    for j in dc['Certificates']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'achievements':
                if dc['Achievements'] and dc['Achievements'][0].lower().strip() != 'achievement1':
                    for j in dc['Achievements']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'qualifications':
                if dc['Qualifications'] and dc['Qualifications'][0].lower().strip() != 'qualification1':
                    for j in dc['Qualifications']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'computer skills':
                if dc['Computer Skills'] and dc['Computer Skills'][0].lower().strip() != 'computerskill1':
                    for j in dc['Computer Skills']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'softwares':
                if dc['Softwares'] and dc['Softwares'][0].lower().strip() != 'software1':
                    for j in dc['Softwares']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'languages':
                if dc['Languages'] and dc['Languages'][0].lower().strip() != 'language1':
                    for j in dc['Languages']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'interests':
                if dc['Interests'] and dc['Interests'][0].lower().strip() != 'interest1':
                    for j in dc['Interests']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'trainings':
                if dc['Trainings'] and dc['Trainings'][0].lower().strip() != 'training1':
                    for j in dc['Trainings']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'skills':
                if dc['Skills'] and dc['Skills'][0].lower().strip() != 'skill1':
                    for j in dc['Skills']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'work experience':
                for j in dc['Work Experience']:
                    company_name = j['Company Name'].strip()
                    duration = j['Duration'].strip()
                    job_title = j['Job Title'].strip()

                    if (j['Company Name'] and j['Company Name'].lower().replace(' ','') != 'nameofcompany') or (j['Job Title'] and j['Job Title'].lower().replace(' ','') != 'titleofjob'):
                        if j['Duration'] and j['Duration'].lower().replace(' ','') != 'workingdurationincompany':
                            duration_run = doc.paragraphs[i+2].add_run(duration + '\n')
                            duration_run.bold = True
                        
                        if j['Company Name'] and j['Company Name'].lower().replace(' ','') != 'nameofcompany':
                            doc.paragraphs[i+2].add_run(company_name + '\n').bold = True
                        
                        if j['Job Title'] and j['Job Title'].lower().replace(' ','') != 'titleofjob':
                            job_run = doc.paragraphs[i+2].add_run(job_title + '\n\n')
                            job_run.font.bold = True

                        if j["Duties"] and j["Duties"][0].lower().replace(' ','') != "duty1":
#                             doc.paragraphs[i+2].add_run('Duties:' + '\n')
                            for k in j['Duties']:
                                if k.strip():
                                    doc.paragraphs[i+2].add_run('  • ' + k.strip() + '\n')
                            doc.paragraphs[i+2].add_run('\n')
        except:
            pass


    doc.save(path_save)
    print("Conversion completed !!")    
    
