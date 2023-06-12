import os
import openai
import docx
import docx2txt
import re
import json
from .keys import api_key
from docx.enum.text import WD_UNDERLINE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
import PyPDF2
import pdfplumber


def get_completion(prompt, model="gpt-3.5-turbo"):
    messages = [{"role": "user", "content": prompt}]
    response = openai.ChatCompletion.create(
        model=model,
        messages=messages,
        temperature=0, # this is the degree of randomness of the model's output
    )
    return response.choices[0].message["content"]

def adway_converter(path ,formatted_path,save_path):
    formatted= formatted_path
    try:
        with open(path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            unformated_text = ""
            for i in range (len(pdf_reader.pages)):
                first_page = pdf_reader.pages[i]
                unformated_text += first_page.extract_text() + " "
            print('Its PDF')
    except:
        try:
            unformated_text = docx2txt.process(path)
            print('Its Docx')
        except:
            print('WE DONT SUPPORT THIS TYPE OF FILE')
        

    os.environ["OPEN_API_KEY"] = api_key
    openai.api_key = api_key
    # llm=OpenAI(temperature=0, max_tokens=1500,openai_api_key=api_key)
    fields_labels = "Name, PROFILE, EDUCATION, IT LITERACY, CERTIFICATES, PROJECTS, PROFESSIONAL QUALIFICATIONS, SOFTWARES, languages, Interests, TRAININGS, skills, WORK EXPERIENCE"

    
    print ("Process has Started...")
    test_text = """

    Ectract data from this text:

    \"""" + unformated_text + """\"

    in following JSON format:
{
"Name":"candidate name",
"Right to work":"value",
"Resides":"location of candidiate",
"Summary" : "value",

"Employment History" : [
    {"Duration" : "Working duration in company",
     "Company Name" : "Name of company",
     "Designation" : "Specific designation in that Company",
     "Responsibilities" : ["Responsibility 1", "Responsibility 2", ...],
    },
    
    {"Duration" : "Working duration in company",
     "Company Name" : "Name of company",
     "Designation" : "Specific designation in that Company",
     "Responsibilities" : ["Responsibility 1", "Responsibility 2", ...],
    },
    ...
    ]
"Education" : [
    {"Duration":Studying duration in institute.",
     "Institute Name":"Name of that institute",
     "Location":"Location of that institute ",
     "Degree":"Name of that degree",
    },
    
    {"Duration":Studying duration in institute.",
     "Institute Name":"Name of that institute",
     "Location":"Location of that institute ",
     "Degree":"Name of that degree",
    },
...],
"Skills" : ["skill1", "skill2", ...],
"Qualifications" : ["qualifications1", "qualifications2", ...],
"Certifications":["certifications1","certifications2",...],
"Languages":["languages1","languages2",...],
"Interests":["interests1","interests2",...],

}

   You must keep the following points in considration while extracting data from text:
     1. Do NOT split, rephrase or summarize list of Responsibilities. Extract each Responsibility as a complete sentence from text.
     2. Make it sure to keep the response in JSON format.
     3. If value not found then leave it empty/blank.
     4. Do not include Mobile number, Email and Home address.
     5. Summary/Personal Statement should be complete without being rephrased.
    """


    result = get_completion(test_text)

    dc = dict(json.loads(re.sub(r'\[\"\"\]',r'[]',re.sub(r'\"[Un]nknown\"|\"[Nn]one\"|\"[Nn]ull\"|\"[Nn]ot [Mm]entioned\"',r'""',re.sub(r',[ \n]*\]',r']',re.sub(r',[ \n]*\}',r'}',result.replace('...','')))))))
#     print(dc)
    doc = docx.Document(formatted)

    for i,p in enumerate(doc.paragraphs):
        try:
            if p.text.strip().lower() == 'name':
                if dc['Name'].lower().strip() != 'candidate name':
                    run = doc.paragraphs[i].add_run(dc['Name'].strip().title())
                    run.bold = True
    #                 run.font.size = Pt(16.5)
        except:
            pass
    #     try:
    #         if p.text.strip().lower() == 'location':
    #                 doc.paragraphs[i].text = ""
    #                 run = doc.paragraphs[i].add_run(dc['Location'].strip().title())
    #                 run.bold = True
    # #                 run.font.size = Pt(16.5)
    #     except:
    #         pass
        try:
            if p.text.strip(' :\n').lower() == 'summary':               
                if dc['Summary'].lower().replace(' ','') != 'value':
                    doc.paragraphs[i+2].add_run(dc['Summary'].strip()).bold = False
                    doc.paragraphs[i+2].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        except:
            pass
        try:
            if p.text.strip(' :\n').lower() == 'resides':
                if dc['Resides'].lower().replace(' ','') != 'locationofcandidiate':
                    run=doc.paragraphs[i].add_run("\t"+dc['Resides'].strip())
                    run.bold=True

        except:
            pass

        if p.text.strip(' :\n').lower() == 'education':
            for j in dc['Education']:
                if j['Degree'].strip() and j['Degree'].lower().replace(' ','') != "nameofthatdegree":
                    if j['Duration'].strip(): 
                        run2=doc.paragraphs[i+2].add_run(j['Duration'].strip()).bold=True
                    else:
                        run2=doc.paragraphs[i+2].add_run("Duration not mentioned").bold=True                       
                    if j['Institute Name'].strip():
                        run1 = doc.paragraphs[i+2].add_run("\t\t\t"+j['Institute Name'].strip()+"\n")
                        run1.bold=True
                    else:
                        run1=doc.paragraphs[i+2].add_run("Institute Name not mentioned"+"\n").bold=False
                    if j['Degree'].strip():
                        run4 = doc.paragraphs[i+2].add_run("\t\t\t"+j['Degree'].strip()+"\n\n")
                        run4.bold=True
                    else:
                        run4=doc.paragraphs[i+2].add_run("Degree not mentioned"+"\n\n")                       
        try:
            if p.text.strip(' :\n').lower() == 'certifications':
                if dc['Certifications'] and dc['Certifications'][0].lower().replace(' ','') != 'certifications1':
                    for j in dc['Certifications']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'languages':
                if dc['Languages'][0].lower().replace(' ','') != 'language1':
                    for j in dc['Languages']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'interests':
                if dc['Interests'][0].lower().replace(' ','') != 'interests1':
                    for j in dc['Interests']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'training':
                if dc['Training'][0].lower().replace(' ','') != 'training1':
                    for j in dc['Training']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'skills':
                if dc['Skills'][0].lower().replace(' ','') != 'skill1':
                    for j in dc['Skills']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'employment history':
                for j in dc['Employment History']:
                    if j['Designation'].strip() and j['Designation'].lower().replace(' ','') !='specificdesignationinthatcompany' or (j['Company Name'].strip() and j['Company Name'].lower().replace(' ','') !='nameofcompany'):
                        if j['Duration'].strip():                                              
                            run2=doc.paragraphs[i+2].add_run(j['Duration'].strip()).bold=True
                        else:
                            run2=doc.paragraphs[i+2].add_run("Duration not mentioned")                           
                        if j['Company Name'].strip():
                            run1=doc.paragraphs[i+2].add_run("\t\t\t"+j['Company Name'].strip()+"\n").bold=True
                        else:
                            run1=doc.paragraphs[i+2].add_run("Company Name not mentioned"+"\n")
                        if j['Designation'].strip():                                                        
                            run3=doc.paragraphs[i+2].add_run("Position"+"\t\t\t"+j['Designation'].strip() + "\n\n")
                            run3.bold=True
                        else:
                            run3=doc.paragraphs[i+2].add_run("Designation not mentioned"+"\n\n")
                        if j['Responsibilities'] and j['Responsibilities'][0].lower().replace(' ','') != 'responsibility1':   
                            doc.paragraphs[i+2].add_run('Responsibilities:' + '\n').bold = True
                            for k in j['Responsibilities']:
                                if k.strip():
                                    doc.paragraphs[i+2].add_run('  • ' + k.strip() + '\n').bold = False
                            doc.paragraphs[i+2].add_run('\n\n')
        except:
            pass

    doc.save(save_path)
    print("\n")
    print("---------------------------------------------------------------------------------------------------------------------")
    print("\n")
    print("Process has Completed...")
