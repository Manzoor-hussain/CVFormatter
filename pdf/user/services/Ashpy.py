import os
import openai
import docx
import docx2txt
import re
import json
import PyPDF2
from docx.enum.text import WD_UNDERLINE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from .keys import api_key

def get_completion(prompt, model="gpt-3.5-turbo"):
    messages = [{"role": "user", "content": prompt}]
    response = openai.ChatCompletion.create(
        model=model,
        messages=messages,
        temperature=0, # this is the degree of randomness of the model's output
    )
    return response.choices[0].message["content"]

def ashbys_converter(path, pathout, path_save):
    
    formatted= pathout
    un_formatted = path
    formated_text = docx2txt.process(formatted)

    try:
        with open(path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            unformated_text = ""
            for i in range (len(pdf_reader.pages)):
                first_page = pdf_reader.pages[i]
                unformated_text += first_page.extract_text() + " "
            print('Its PDF')
    except:
        try:
            unformated_text = docx2txt.process(path)
            print('Its Docx')
        except:
            print('WE DONT SUPPORT THIS TYPE OF FILE')

    os.environ["OPEN_API_KEY"] = api_key
    openai.api_key = api_key

    print("Process has Started...")
    
    test_text = """

    Ectract data from this text:

    \"""" + unformated_text + """\"

    in following JSON format:
    {
    "Profile" : "value",

    "Education" : [
        {"Duration" : "Studying duration in institute",
        "Institute Name" : "Name Of institute",
        "Degree Name": "Name of degree",
        },
        {"Duration" : "Studying duration in institute",
        "Institute Name" : "Name Of institute",
        "Degree Name": "Name of degree",
        },
        ...
        ],
    "Work History" : [
        {"Duration" : "Working duration in specific company",
        "Company Name" : "Name of company",
        "Designation" : "Specific designation in that Company",
        "Responsibilities" : ["Responsibility1", "Responsibility2", ...],
        },
        {"Duration" : "Working duration in specific company",
        "Company Name" : "Name of company",
        "Designation" : "Specific designation in that Company",
        "Responsibilities" : ["Responsibility1", "Responsibility2", ...],
        },
        ...
        ],

    "Trainings" : ["training1", "training2", ...],
    "Projects" : ["project1", "project2", ...],
    "Skills" : ["skill1", "skill2", ...],
    "Languages" : ["language1", "language2", ...],
    "Personal Skills" : ["personal skill1", "personal skill2", ...],
    "Interests" : ["interest1", "interest2", ...]
    }
    
    You must keep the following points in considration while extracting data from text:
        1. Do NOT split, rephrase or summarize list of Responsibilities. Extract each Responsibility as a complete sentence from text.
        2. Make it sure to keep the response in JSON format.
        3. If value not found then leave it empty/blank.
        4. Do not include Mobile number, Email and Home address.
        5. Summary/Personal Statement should be complete without being rephrased.

    """

    result = get_completion(test_text)
    dc = dict(json.loads(re.sub(r'\[\"\"\]',r'[]',re.sub(r'\"[Un]nknown\"|\"[Nn]one\"|\"[Nn]ull\"|\"[Nn]ot [Mm]entioned\"',r'""',re.sub(r',[ \n]*\]',r']',re.sub(r',[ \n]*\}',r'}',result.replace('...','')))))))  
    doc = docx.Document(formatted)

    
    for i,p in enumerate(doc.paragraphs):
        try:
            if p.text.strip(' :\n').lower() == 'profile':
                doc.paragraphs[i].text = ""
                if dc['Profile'].lower().replace(' ','') != 'value':
                    doc.paragraphs[i].add_run(dc['Profile'].strip()).bold = False
                    doc.paragraphs[i].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        except:
            pass
        
        try:        
            if p.text.strip(' :\n').lower() == 'education':
                doc.paragraphs[i].text = ""
                for j in dc['Education']:
                    if j['Degree Name'].strip() and j['Degree Name'].lower().replace(' ','') != "nameofthatdegree":
                        if j["Duration"].strip():
                            doc.paragraphs[i].add_run(j["Duration"].strip() + '\n').bold = False
                        else:
                            doc.paragraphs[i].add_run("Duration not mentioned"+"\n")
                        if j["Institute Name"].strip():                                                       
                            doc.paragraphs[i].add_run(j["Institute Name"].strip() + '\n').bold = False
                        else:
                            doc.paragraphs[i].add_run("Institute Name not mentioned"+"\n")   
                            
                        doc.paragraphs[i].add_run(j['Degree Name'].strip() + '\n\n').bold = True
        except:
            pass
        
        try:
            if p.text.strip(' :\n').lower() == 'projects':
                doc.paragraphs[i].text = ""
                if dc['Projects'][0].lower().replace(' ','') != 'project1':
                    for j in dc['Projects']:
                        if j.strip():
                            doc.paragraphs[i].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass
        

        try:
            if p.text.strip(' :\n').lower() == 'personal skill':
                doc.paragraphs[i].text = ""
                if dc['Personal Skills'][0].lower().replace(' ','') != 'personalskill1':
                    for j in dc['Personal Skills']:
                        if j.strip():
                            doc.paragraphs[i].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass


        try:
            if p.text.strip(' :\n').lower() == 'languages':
                doc.paragraphs[i].text = ""
                if dc['Languages'][0].lower().replace(' ','') != 'languages1':
                    for j in dc['Languages']:
                        if j.strip():
                            doc.paragraphs[i].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'interests':
                doc.paragraphs[i].text = ""
                if dc['Interests'][0].lower().replace(' ','') != 'interests1':
                    for j in dc['Interests']:
                        if j.strip():
                            doc.paragraphs[i].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'training':
                doc.paragraphs[i].text = ""
                if dc['Trainings'][0].lower().replace(' ','') != 'training1':
                    for j in dc['Trainings']:
                        if j.strip():
                            doc.paragraphs[i].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'skills':
                doc.paragraphs[i].text = ""
                if dc['Skills'][0].lower().replace(' ','') != 'skills1':
                    for j in dc['Skills']:
                        if j.strip():
                            doc.paragraphs[i].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'work history':
                doc.paragraphs[i].text = ""
                for j in dc['Work History']:
                    if j['Designation'].strip() and j['Designation'].lower().replace(' ','') !='specificdesignationinthatcompany' or (j['Company Name'].strip() and j['Company Name'].lower().replace(' ','') !='nameofcompany'):
                        if j['Duration'].strip():
                            doc.paragraphs[i].add_run(j['Duration'].strip() + '\n').bold = True
                        else:
                            doc.paragraphs[i].add_run("Duration not mentioned"+"\n")
                        if j['Company Name'].strip():
                            doc.paragraphs[i].add_run(j['Company Name'].strip() + '\n').bold = True
                        else:
                            doc.paragraphs[i].add_run("Company Name not mentioned"+"\n")
                        if j['Designation'].strip():
                            doc.paragraphs[i].add_run(j['Designation'].strip() + '\n\n').bold = True
                        else:
                            doc.paragraphs[i].add_run("Designation not mentioned"+"\n\n")

                        if j['Responsibilities'] and j['Responsibilities'][0].lower().replace(' ','') != 'responsibility1':       
                            for k in j['Responsibilities']:
                                doc.paragraphs[i].add_run('\t' + '  • ' + k.strip() + '\n').bold = False
                            doc.paragraphs[i].add_run('\n')
        except:
            pass

    doc.save(path_save)
    print("Process has Completed...")
   
