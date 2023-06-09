import os
import openai
import docx
import docx2txt
from keys import api_key
from .keys import api_key
import json
import re
import textwrap
import PyPDF2
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def read_text_from_docx(file_path):
    doc = docx.Document(file_path)
    text = [paragraph.text for paragraph in doc.paragraphs]
    return '\n'.join(text)

def read_text_from_pdf(file_path):
    with open(file_path, 'rb') as file:
        pdf_reader = PyPDF2.PdfReader(file)
        text = []
        for page in pdf_reader.pages:
            text.append(page.extract_text())
        return '\n'.join(text)


def get_completion(prompt, model="gpt-3.5-turbo"):
    messages = [{"role": "user", "content": prompt}]
    response = openai.ChatCompletion.create(
        model=model,
        messages=messages,
    )
    return response.choices[0].message["content"]


def cw_executive_converter(path, pathout, path_save):
    formatted= pathout
    file_path = path  

    if file_path.endswith('.docx'):
        unformated_text = read_text_from_docx(file_path)
    elif file_path.endswith('.pdf'):
        unformated_text = read_text_from_pdf(file_path)
    else:
        print('Unsupported file format')


    os.environ["OPEN_API_KEY"] = api_key
    openai.api_key = api_key

    print ("Process has Started...")
    test_text = """

    Extract data from this text:

    \"""" + unformated_text + """\"

    in following JSON format:
    {
    "Name" : "value",
    "Professional Profile" : "value",
    "Career History" : [
        {"Duration" : "Working Duration in Company",
        "Job Title" : "Title of job",
        "Company Name" : "Name of Company",
        "Responsibilities" : ["Responsiblility1, Responsibility2", ...],
        },
        {"Duration" : "Working Duration in Company",
        "Job Title" : "Title of job",
        "Company Name" : "Name of Company",
        "Responsibilities" : ["Responsiblility1, Responsibility2", ...],
        },
        ...
        ],
    "Education" : [
        {"Duration" : "Studying duration in institute",
        "Institute Name" : "Name Of institute",
        "Degree Name": "Name of degree"
        },
        {"Duration" : "Studying duration in institute",
        "Institute Name" : "Name Of institute",
        "Degree Name": "Name of degree"
        },
        ...
        ],
    "Achievements" : ["Achievement1","Achievement2",...],
    "Qualifications" : ["Qualification1", "Qualification2", ...],
    "Skills" : ["Skill1", "Skill2", ...],
    "Languages" : ["Language1", "Language2", ...],
    "Interests" : ["Interest1", "Interest2", ...],
    }

    You must keep the following points in considration while extracting data from text:
        1. Do NOT split, rephrase or summarize list of Responsibilities. Extract each Responsibility as a complete sentence from text.
        2. Make it sure to keep the response in JSON format.
        3. If value not found then leave it empty/blank.
        4. Do not include Mobile number, Email and Home address.
        5. Summary/Personal Statement should be complete without being rephrased.
        
        """

    result = get_completion(test_text)




    dc = dict(json.loads(re.sub(r'\[\"\"\]',r'[]',re.sub(r'\"[Un]nknown\"|\"[Nn]one\"|\"[Nn]ull\"|\"[Nn]ot [Mm]entioned\"',r'""',re.sub(r',[ \n]*\]',r']',re.sub(r',[ \n]*\}',r'}',result.replace('...','')))))))

    doc = docx.Document(formatted)


    def change_font_size(paragraph, size):
        for run in paragraph.runs:
            run.font.size = size

    for i, p in enumerate(doc.paragraphs):
        if p.text.strip(' :\n').lower() == 'name':
            try:
                name = doc.paragraphs[i]
                if dc['Name'].lower().replace(' ','') != 'value':
                    name.text = dc['Name']
                    name.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
                for run in name.runs:
                    run.bold = True
                    run.font.size = docx.shared.Pt(16) 
            except:
                pass

        if p.text.strip(' :\n').lower() == 'professional profile':
            try:
                summary = doc.paragraphs[i+1] 
                if dc['Professional Profile'][0].lower().replace(' ','') != 'professionalprofile1':
                    summary.text = str(dc['Professional Profile'])
                    doc.paragraphs[i+1].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    change_font_size(summary, docx.shared.Pt(12)) 
            except:
                pass

        if p.text.strip(' :\n').lower() == 'career history':
            try:
                for j in dc['Career History']:
                    if j['Job Title'].strip() and j['Job Title'].lower().replace(' ','') !='titleofjob' or (j['Company Name'].strip() and j['Company Name'].lower().replace(' ','') !='nameofcompany'):
                        if j['Duration'].strip():
                            doc.paragraphs[i+1].add_run(j['Duration'].strip() + '\n').bold = True
                        else:
                            doc.paragraphs[i+1].add_run("Duration not mentioned"+ '\n').bold = True                          
                        if j['Job Title'].strip():
                            doc.paragraphs[i+1].add_run(j['Job Title'].strip() + '\n').bold = True
                        else:
                            doc.paragraphs[i+1].add_run("Job Title not mentioned"+ '\n').bold = True                          
                        if j['Company Name'].strip():
                            doc.paragraphs[i+1].add_run(j['Company Name'].strip() + '\n\n').bold = True
                        else:
                            doc.paragraphs[i+1].add_run("Company Name not mentioned"+ '\n\n').bold = True                          
                        if j['Responsibilities'] and j['Responsibilities'][0].lower().replace(' ','') != 'responsibility1':
                            for k in j['Responsibilities']:
                                doc.paragraphs[i+1].add_run('•   ' + k.strip() + '\n').bold = False
                            doc.paragraphs[i+1].add_run('\n')
            except:
                pass

        if p.text.strip(' :\n').lower() == 'education':
            try:
                for j in dc['Education']:
                    if j['Degree Name'].strip() and j['Degree Name'].lower().replace(' ','') != "nameofthatdegree":
                        if j['Duration'].strip():
                            duration = j['Duration'].strip()
                            doc.paragraphs[i+1].add_run(duration + '\n').bold = True
                        else:
                            doc.paragraphs[i+1].add_run("Duration not mentioned"+'\n')
                            
                        if j['Degree Name'].strip():
                            degree_name = j['Degree Name'].strip()
                            doc.paragraphs[i+1].add_run(degree_name + '\n').bold = True
                        else:
                            doc.paragraphs[i+1].add_run("Degree Name not mentioned"+'\n')

                        if j['Institute Name'].strip():
                            institute_name = j['Institute Name'].strip()
                            doc.paragraphs[i+1].add_run(institute_name + '\n').bold = True
                            doc.paragraphs[i+1].add_run('\n')
                        change_font_size(doc.paragraphs[i+1], docx.shared.Pt(12))
#                         else:
#                             doc.paragraphs[i+1].add_run("Institute Name not mentioned"+'\n')                            
            except:
                pass

        if p.text.strip(' :\n').lower() == 'achievements':
            try:
                if dc['Achievements'][0].lower().replace(' ','') != 'achievements1':
                    for j in dc['Achievements']:
                        if j.strip():
                            doc.paragraphs[i+1].add_run('•   ' + j.strip() + '\n')
                            change_font_size(doc.paragraphs[i+1], docx.shared.Pt(12)) 
            except:
                pass

        if p.text.strip(' :\n').lower() == 'qualifications':
            try:
                if dc['Qualifications'][0].lower().replace(' ','') != 'qualifications1':
                    for j in dc['Qualifications']:
                        if j.strip():
                            doc.paragraphs[i+1].add_run('•   ' + j.strip() + '\n')
                            change_font_size(doc.paragraphs[i+1], docx.shared.Pt(12)) 
            except:
                pass

        if p.text.strip(' :\n').lower() == 'skills':
            try:
                if dc['Skills'][0].lower().replace(' ','') != 'skills1':
                    for j in dc['Skills']:
                        if j.strip():
                            doc.paragraphs[i+1].add_run('•   ' + j.strip() + '\n')
                            change_font_size(doc.paragraphs[i+1], docx.shared.Pt(12)) 
            except:
                pass

        if p.text.strip(' :\n').lower() == 'languages':
            try:
                if dc['Languages'][0].lower().replace(' ','') != 'languages1':
                    for j in dc['Languages']:
                        if j.strip():
                            doc.paragraphs[i+1].add_run('•   ' + j.strip() + '\n')
                            change_font_size(doc.paragraphs[i+1], docx.shared.Pt(12)) 
            except:
                pass

        if p.text.strip(' :\n').lower() == 'interests':
            try:
                if dc['Interests'][0].lower().replace(' ','') != 'interests1':
                    for j in dc['Interests']:
                        if j.strip():
                            doc.paragraphs[i+1].add_run('•   ' + j.strip() + '\n')
                            change_font_size(doc.paragraphs[i+1], docx.shared.Pt(12)) 
            except:
                pass


    doc.save(path_save)
    print("Process has Completed...")
