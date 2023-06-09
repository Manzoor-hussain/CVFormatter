import os
import openai
import docx
import docx2txt
from .keys import api_key
from pprint import pprint
import json
import re
import textwrap
import PyPDF2
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def get_completion(prompt, model="gpt-3.5-turbo"):
    messages = [{"role": "user", "content": prompt}]
    response = openai.ChatCompletion.create(
        model=model,
        messages=messages,
        temperature=0, # this is the degree of randomness of the model's output
    )
    return response.choices[0].message["content"]


# Functions to check whether the unformatted file is a docx or pdf
def read_text_from_docx(file_path):
    doc = docx.Document(file_path)
    text = [paragraph.text for paragraph in doc.paragraphs]
    return '\n'.join(text)

def read_text_from_pdf(file_path):
    with open(file_path, 'rb') as file:
        pdf_reader = PyPDF2.PdfReader(file)
        text = []
        for page in pdf_reader.pages:
            text.append(page.extract_text())
        return '\n'.join(text)
    

def scale_genesis_converter(path_in, path_out, path_save):
    
    formatted = path_out
    
    # unformatted document
    if path_in.endswith('.docx'):
        unformatted_text = read_text_from_docx(path_in)
    elif path_in.endswith('.pdf'):
        unformatted_text = read_text_from_pdf(path_in)
    else:
        error = 'Format not supported.'
        print(error)

    # formatted document
    formatted_text = docx2txt.process(formatted)
    
    
    print("----------------------------------------------------------------")
    print("                          Unformatted Text                            ")
    print("----------------------------------------------------------------")
    print(unformatted_text)
    
    print("Process has started...")
    
    
    openai.api_key = api_key
    test_text = """

    Extract data from this text:

    \"""" + unformatted_text + """\"

    in following JSON format:
    {
    "Name" : "value",
    "Position" : "value",
    "Availability" : "value",
    "Summary of skills" : ["Key Skill1", "Key Skill2", ...],
    "Salary Expectations/Rate" : "value",
    "Location" : "value",
    "Summary" : "value",

    "Work Experience" : [
        {"Company Name" : "Name of company",
        "Job Title" : "Title of job",
        "Duration" : "Working Duration in Company",
        "Responsibilities" : ["Responsibility 1", "Responsibility 2", ...],
        },
        {"Company Name" : "Name of company",
        "Job Title" : "Title of job",
        "Duration" : "Working Duration in Company",
        "Responsibilities" : ["Responsibility 1", "Responsibility 2", ...],
        },
        ...
        ],

    "Education" : [
            {"Institute Name" : "Name Of institute",
            "Degree Name": "Name of degree",
            "Duration" : "Studying duration in institute",
            },
            {"Institute Name" : "Name Of institute",
            "Degree Name": "Name of degree",
            "Duration" : "Studying duration in institute",
            },
            ...
            ],
    "Professional Qualifications" : ["Professional Qualification1", "Professional Qualification2", ...],
    "Skills" : ["Skill1", "Skill2", ...],
    "Languages" : ["Language1", "Language2", ...],
    "Interests" : ["Interest1", "Interest2", ...],

    }
    You must keep the following points in considration while extracting data from text:
        1. Do NOT split, rephrase or summarize list of Responsibilities. Extract each Responsibility as a complete sentence from text.
        2. Make it sure to keep the response in JSON format.
        3. If value not found then leave it empty/blank.
        4. Do not include Mobile number, Email and Home address.
        5. Summary/Personal Statement should be as it is. Do not change or rephrase it.
        
    """
    result = get_completion(test_text)
    
    print("----------------------------------------------------------------")
    print("                          Result                            ")
    print("----------------------------------------------------------------")
    print(result)
    
    dc = dict(json.loads(re.sub(r'\[\"\"\]',r'[]',re.sub(r'\"[Un]nknown\"|\"[Nn]one\"|\"[Nn]ull\"|\"[Nn]ot [Mm]entioned\"',r'""',re.sub(r',[ \n]*\]',r']',re.sub(r',[ \n]*\}',r'}',result.replace('...','')))))))
    
    print("----------------------------------------------------------------")
    print("                          Dictionary                            ")
    print("----------------------------------------------------------------")
    print(dc)

    
    doc = docx.Document(formatted)


    for table in doc.tables:
            for row in table.rows:
                for i,cell in enumerate(row.cells):

    #                 doc.paragraphs[i].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                    try:
                        if cell.text.strip(' :\n').lower() == 'name':
                            if dc['Name'].strip() and dc['Name'].lower().replace(' ','') != 'value':
                                row.cells[i+1].text = dc['Name']
                    except:
                        pass

                    try:
                        if cell.text.strip(' :\n').lower() == 'position':
                            if dc['Position'].strip() and dc['Position'].lower().replace(' ','') != 'value':
                                row.cells[i+1].text = dc['Position']
                    except:
                        pass

                    try:
                        if cell.text.strip(' :\n').lower() == 'availability':
                            if dc['Availability'].strip() and dc['Availability'].lower().replace(' ','') != 'value':
                                row.cells[i+1].text = dc['Availability']
                    except:
                        pass

                    try:
                        if cell.text.strip(' :\n').lower() == 'summary of skills':
                            if dc['Summary of skills'][0].strip() and dc['Summary of skills'][0].lower().replace(' ','') != 'keyskill1':
                                for j in dc['Summary of skills']:
                                    if j.strip():
                                        run = row.cells[i+1].paragraphs[0].add_run('  • ' + j.strip() + '\n')
                    except:
                        pass

                    try:
                        if cell.text.strip(' :\n').lower() == 'salary expectations/rate':
                            if dc['Salary Expectations/Rate'].strip() and dc['Salary Expectations/Rate'].lower().replace(' ','') != 'value':
                                row.cells[i+1].text = dc['Salary Expectations/Rate']
                    except:
                        pass

                    try:
                        if cell.text.strip(' :\n').lower() == 'location':
                            if dc['Location'].strip() and dc['Location'].lower().replace(' ','') != 'value':
                                row.cells[i+1].text = dc['Location']
                    except:
                        pass


    for i,p in enumerate(doc.paragraphs):

        if p.text.strip(' :\n').lower() == 'summary':
            if dc['Summary'].strip() and dc['Summary'].lower().replace(' ','') != 'value':
                doc.paragraphs[i+2].add_run(dc['Summary'])

        if p.text.strip(' :\n').lower() == 'work experience':
            try:
                for j in dc['Work Experience']:
                    try:
                        company_name = ""
                        duration = ""
                        job_title = ""
                        if (j['Company Name'].strip() and j['Company Name'].lower().replace(' ','') != 'nameofcompany') or (j['Job Title'].strip() and j['Job Title'].lower().replace(' ','') != 'titleofjob'):                    
                            if (j['Company Name'].strip() and j['Company Name'].lower().replace(' ','') != 'nameofcompany'):
                                company_name = j['Company Name'].strip()
                            if (j['Duration'].strip() and j['Duration'].lower().replace(' ','') != 'workingdurationincompany'):
                                duration = j['Duration'].strip()
                            if (j['Job Title'].strip() and j['Job Title'].lower().replace(' ','') != 'titleofjob'):
                                job_title = j['Job Title'].strip()

                            if company_name:
                                doc.paragraphs[i+2].add_run(company_name + ' ').bold = True
                            else:
                                doc.paragraphs[i+2].add_run("Company not mentioned" + ' ').bold = True

                            if duration:
                                doc.paragraphs[i+2].add_run('(' + duration + ')' + '\n').bold = True
                            else:
                                doc.paragraphs[i+2].add_run('(' + "Duration not mentioned" + ')' + '\n').bold = True
                            if job_title:
                                doc.paragraphs[i+2].add_run(job_title + '\n\n').bold = False
                            else:
                                doc.paragraphs[i+2].add_run("Job Title not mentioned" + '\n\n').bold = False
                                
                            try:
                                if j["Responsibilities"][0].lower().replace(' ','') != "responsibility1":
                                    for k in j['Responsibilities']:
                                        if k.strip():
                                            doc.paragraphs[i+2].add_run('  - ' + k.strip() + '\n')
                                    doc.paragraphs[i+2].add_run("\n\n")
                            except:
                                pass
                    except:
                        pass
            except:
                pass


        if p.text.strip(' :\n').lower() == 'education':
            for j in dc['Education']:
                try:
                    institute_name = ""
                    duration = ""
                    degree_name = ""
                    if j['Degree Name'].strip() and j['Degree Name'].lower().replace(' ','') != 'nameofdegree':
                        degree_name = j['Degree Name'].strip()
                        if j['Institute Name'].strip() and j['Institute Name'].lower().replace(' ','') != 'nameofinstitute':
                            institute_name = j['Institute Name'].strip()
                        if j['Duration'].strip() and j['Duration'].lower().replace(' ','') != 'studyingdurationininstitute':
                            duration = j['Duration'].strip()

                        if duration:
                            doc.paragraphs[i+2].add_run(institute_name + ' ').bold = True
                        else:
                            doc.paragraphs[i+2].add_run("Institute not mentioned ").bold = True
                        if duration:
                            doc.paragraphs[i+2].add_run('(' + duration + ')' + '\n').bold = True
                        else:
                            doc.paragraphs[i+2].add_run('(' + "Duration not mentioned" + ')' + '\n').bold = True

                        doc.paragraphs[i+2].add_run(degree_name + '\n\n').bold = False
                except:
                    pass


        if p.text.strip(' :\n').lower() == 'professional qualifications':
            try:
                if dc['Professional Qualifications'][0].lower().replace(' ','') != 'professionalqualification1':
                    for j in dc['Professional Qualifications']:
                        if j.strip():
                            language_run = doc.paragraphs[i+2].add_run('  - ' + j.strip() + '\n')

            except:
                pass
            
        if p.text.strip(' :\n').lower() == 'skills':
            try:
                if dc['Skills'][0].lower().replace(' ','') != 'skill1':
                    for j in dc['Skills']:
                        if j.strip():
                            language_run = doc.paragraphs[i+2].add_run('  - ' + j.strip() + '\n')

            except:
                pass

        if p.text.strip(' :\n').lower() == 'languages':
            try:
                if dc['Languages'][0].lower().replace(' ','') != 'language1':
                    for j in dc['Languages']:
                        if j.strip():
                            language_run = doc.paragraphs[i+2].add_run('  - ' + j.strip() + '\n')

            except:
                pass


        if p.text.strip(' :\n').lower() == 'interests':
            try:
                if dc['Interests'][0].lower().replace(' ','') != 'interest1':
                    for j in dc['Interests']:
                        if j.strip():
                            doc.paragraphs[i+1].add_run("\n")
                            doc.paragraphs[i+1].add_run('  - ' + j.strip())
            except:
                pass

    doc.save(path_save)
    print("Conversion completed !!")
    
