import os
import openai
import docx
import docx2txt
import re
import json
import PyPDF2
from docx.enum.text import WD_UNDERLINE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.shared import RGBColor
from docx import Document
from .keys import api_key

def get_completion(prompt, model="gpt-3.5-turbo"):
    messages = [{"role": "user", "content": prompt}]
    response = openai.ChatCompletion.create(
        model=model,
        messages=messages,
        temperature=0, # this is the degree of randomness of the model's output
    )
    return response.choices[0].message["content"]

def timber_seed_format_converter(path, pathout, path_save):
    formatted = pathout
    formated_text = docx2txt.process(formatted)
    un_formatted = path

    try:
        with open(path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            unformated_text = ""
            for i in range (len(pdf_reader.pages)):
                first_page = pdf_reader.pages[i]
                unformated_text += first_page.extract_text() + " "
            print('Its PDF')
    except:
        try:
            unformated_text = docx2txt.process(path)
            print('Its Docx')
        except:
            print('WE DONT SUPPORT THIS TYPE OF FILE')


    os.environ["OPEN_API_KEY"] = api_key
    openai.api_key = api_key
    
    print("Process has Started...")
    test_text = """
    Ectract data from this text:
    \"""" + unformated_text + """\"
    in following JSON format:
    {"Name" : "value",
    "Summary" : "value",
    "Education" : [
        {"Institute Name" : "Name Of institute",
        "Location" : "Location of Institute",
        "Duration" : "Duration for specific degree in this specific institute",
        "Degree Name": "Name of degree"
        },
        {"Institute Name" : "Name Of institute",
        "Location" : "Location of Institute",
        "Duration" : "Duration for specific degree in this specific institute",
        "Degree Name": "Name of degree"
        },
        ...
        ],
    "Previous Employment" : [
        {"Company Name" : "Name of company",
        "Company Location" : "Location of Company",
        "Duration" : "Working Duration in Company",
        "Designation" : "Specific designation in that Company",
        "Responsibilities" : ["Responsibility 1", "Responsibility 2", ...],
        },
        {"Company Name" : "Name of company",
        "Company Location" : "Location of Company",
        "Duration" : "Working Duration in Company",
        "Designation" : "Specific designation in that Company",
        "Responsibilities" : ["Responsibility 1", "Responsibility 2", ...],
        },
        ...
        ],
    "Skills" : ["skill1", "skill2", ...],
    "Leadership and Management" : ["Leadership1", "Leadership2", ...],
    "Other Training" : ["Training1", "Training2", ...],
    "Languages" : ["language1", "language2", ...],
    "Interests" : ["interest1", "interest2", ...],
    }
    You must keep the following points in considration while extracting data from text:
        1. Do NOT split, rephrase or summarize list of Responsibilities. Extract each Responsibility as a complete sentence from text.
        2. Make it sure to keep the response in JSON format.
        3. If value not found then leave it empty/blank.
        4. Do not include Mobile number, Email and Home address.
        5. Summary/Personal Statement should be complete without being rephrased.

    """


    result = get_completion(test_text)

    dc = dict(json.loads(re.sub(r'\[\"\"\]',r'[]',re.sub(r'\"[Un]nknown\"|\"[Nn]one\"|\"[Nn]ull\"|\"[Nn]ot [Mm]entioned\"',r'""',re.sub(r',[ \n]*\]',r']',re.sub(r',[ \n]*\}',r'}',result.replace('...','')))))))
#     print(dc)

    doc = docx.Document(formatted)

    for i,p in enumerate(doc.paragraphs):

        try:
            if p.text.strip().lower() == 'name':
                doc.paragraphs[i].text = ""
                if dc['Name'].lower().replace(' ','') != 'value':
                    run = doc.paragraphs[i].add_run(dc['Name'].strip().title())
                    run.bold = True
                    run.font.size = Pt(26.5)
                    run.font.color.rgb = RGBColor(255, 0, 0)  # Red color
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'summary':
                doc.paragraphs[i].text = ""
                if dc['Summary'].lower().replace(' ','') != 'value':
                    doc.paragraphs[i].add_run(dc['Summary'].strip()).bold = False
                    doc.paragraphs[i].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        except:
            pass

        try:        
            if p.text.strip(' :\n').lower() == 'education':
                for j in dc['Education']:
                    if j['Degree Name'].strip() and j['Degree Name'].lower().replace(' ','') != "nameofdegree":
                        if j["Institute Name"].strip():
                            doc.paragraphs[i+2].add_run(j["Institute Name"].strip() + ", " + j['Location'] + "\n").font.underline = True
                        else:
                            doc.paragraphs[i+2].add_run("Institute Name not mentioned").bold=False
                        if j['Duration'].strip():                            
                            doc.paragraphs[i+2].add_run(j['Duration'].strip() + '\n').bold = False
                        else:
                            doc.paragraphs[i+2].add_run("Duration not mentioned").bold=False
                        if j['Degree Name'].strip():                                                       
                            doc.paragraphs[i+2].add_run(j['Degree Name'].strip() + '\n\n').bold = False
                        else:
                            doc.paragraphs[i+2].add_run("Degree not mentioned").bold=False

                            
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'leadership and management':
                if dc['Leadership and Management'][0].lower().replace(' ','') != 'leadership1':
                    for j in dc['Leadership and Management']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass



        try:
            if p.text.strip(' :\n').lower() == 'languages':
                if dc['Languages'][0].lower().replace(' ','') != 'languages1':
                    for j in dc['Languages']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'interests':
                if dc['Interests'][0].lower().replace(' ','') != 'interests1':
                    for j in dc['Interests']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'other training':
                if dc['Other Training'][0].lower().replace(' ','') != 'othertraining1':
                    for j in dc['Other Training']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'skills':
                if dc['Skills'][0].lower().replace(' ','')!= 'skills1':
                    for j in dc['Skills']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass


        try:
            if p.text.strip(' :\n').lower() == 'previous employment':
                for j in dc['Previous Employment']:
                    if j['Designation'].strip() and j['Designation'].lower().replace(' ','') !='specificdesignationinthatcompany' or (j['Company Name'].strip() and j['Company Name'].lower().replace(' ','') !='nameofcompany'):
                        if j['Company Name'].strip():                          
                            doc.paragraphs[i+2].add_run(j['Company Name'].strip()).font.underline=True
                        else:
                            doc.paragraphs[i+2].add_run("Company Name not mentioned").bold=False
                        if j['Company Location'].strip():                            
                            doc.paragraphs[i+2].add_run(' ' + j['Company Location'] + '\n').font.underline = True
                        else:
                            doc.paragraphs[i+2].add_run("Company Location not mentioned"+'\n').bold=False
                        if j['Duration'].strip():                           
                            doc.paragraphs[i+2].add_run(j['Duration'].strip() + '\n').bold = False
                        else:
                            doc.paragraphs[i+2].add_run("Duration not mentioned"+'\n').bold=False
                        if j['Designation'].strip():
                            doc.paragraphs[i+2].add_run(j['Designation'].strip() + '\n\n').bold = False
                        else:
                            doc.paragraphs[i+2].add_run("Designation not mentioned"+'\n\n').bold=False
                            
                        if j['Responsibilities'] and j['Responsibilities'][0].lower().replace(' ','') != 'responsibility1':
                            for k in j['Responsibilities']:
                                if k.strip():
                                    doc.paragraphs[i+2].add_run('  • ' + k.strip() + '\n').bold = False
                            doc.paragraphs[i+2].add_run('\n\n')
        except:
            pass

    doc.save(path_save)
    print("Process has Completed...")
