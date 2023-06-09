import os
import openai
import docx
import docx2txt
import re
import json
from .keys import api_key
from docx.enum.text import WD_UNDERLINE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
import PyPDF2


def get_completion(prompt, model="gpt-3.5-turbo"):
    messages = [{"role": "user", "content": prompt}]
    response = openai.ChatCompletion.create(
        model=model,
        messages=messages,
        temperature=0, # this is the degree of randomness of the model's output
    )
    return response.choices[0].message["content"]

def jler_converter(path,formatted_path,savepath):
    formatted= formatted_path
#     un_formatted=os.getcwd() + path
    
#     doc = docx.Document(un_formatted)
#     formated_text = docx2txt.process(formatted)
#     unformated_text = docx2txt.process(un_formatted)
    def read_text_from_docx(path):
        doc = docx.Document(path)
        text = [paragraph.text for paragraph in doc.paragraphs]
        return '\n'.join(text)

    def read_text_from_pdf(path):
        with open(path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = []
            for page in pdf_reader.pages:
                text.append(page.extract_text())
            return '\n'.join(text)

    if path.endswith('.docx'):
        unformated_text = read_text_from_docx(path)
    elif path.endswith('.pdf'):
        unformated_text = read_text_from_pdf(path)
    else:
        unformated_text = 'Unsupported file format'
        

    os.environ["OPEN_API_KEY"] = api_key
    openai.api_key = api_key
    # llm=OpenAI(temperature=0, max_tokens=1500,openai_api_key=api_key)
    fields_labels = "Name, PROFILE, EDUCATION, IT LITERACY, CERTIFICATES, PROJECTS, PROFESSIONAL QUALIFICATIONS, SOFTWARES, languages, Interests, TRAININGS, skills, WORK EXPERIENCE"

    
    print ("Process has Started...")
    test_text = """

    Ectract data from this text:

    \"""" + unformated_text + """\"

    in following JSON format:
    {

    "Name":"value"
    "Career summary" :"summary",

    "Employment History" : [
        {"Duration" : "Working Duration in Company",
         "Designation":"Specific designation in that Company",
         "Company Name" :"Name of company",
         "Location":"Country",
         "Responsibilities" : ["Responsibility 1", "Responsibility 2", ...],

        },
        {"Duration" : "Working Duration in Company",
         "Designation":"Specific designation in that Company",
         "Company Name" :"Name of company",
         "Location":"Country",
         "Responsibilities" : ["Responsibility 1", "Responsibility 2", ...],
        },

    "Education" : [
        {"Degree" : "Name of degree",
         "Duration":"Studying duration in institute",
         "Institute Name":"Name Of institute",
         "Location":"location of that institute",
        },
        {"Degree" : "Name of degree",
         "Duration":"Studying duration in institute",
         "Institute Name":"Name Of institute",
         "Location":"location of that institute",
        },
        ...
        ],
    "Trainings" : ["training1", "training2", ...],
    "Skills" : ["skill1", "skill2", ...],
    "Interests" : ["interest1", "interest2", ...],
    "Languages" : ["language1", "language2", ...],

        ...
        ]
    }

     You must keep the following points in considration while extracting data from text:
      1. Do NOT split, rephrase or summarize list of Responsibilities. Extract each Responsibility as a complete sentence from text.
      2. Make it sure to keep the response in JSON format.
      3. If value not found then leave it empty/blank.
      4. Do not include Mobile number, Email and Home address
      5. Summary/Personal Statement should be complete without being rephrased.
    """


    result = get_completion(test_text)

    dc = dict(json.loads(re.sub(r'\[\"\"\]',r'[]',re.sub(r'\"[Un]nknown\"|\"[Nn]one\"|\"[Nn]ull\"|\"[Nn]ot [Mm]entioned\"',r'""',re.sub(r',[ \n]*\]',r']',re.sub(r',[ \n]*\}',r'}',result.replace('...','')))))))
#     print(dc)
    doc = docx.Document(formatted)

    for i,p in enumerate(doc.paragraphs):
        try:
            if p.text.strip().lower() == 'name':
                doc.paragraphs[i].text = ""
                if dc['Name'].lower().replace(' ','') != 'value':
                    run = doc.paragraphs[i].add_run(dc['Name'].strip().title())
                    run.bold = True
                    run.font.size = Pt(16.5)
        except:
            pass
        try:
            if p.text.strip(' :\n').lower() == 'career summary.':
                if dc['Career summary'].lower().replace(' ','') != 'summary':
                    doc.paragraphs[i+2].add_run(dc['Career summary'].strip()).bold = False
                    doc.paragraphs[i+2].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        except:
            pass

        try:        
            if p.text.strip(' :\n').lower() == 'education':
                for j in dc['Education']:
                    if j['Degree'].strip() and j['Degree'].lower().replace(' ','') != "nameofdegree":
                        if j['Duration'].strip(): 
                            doc.paragraphs[i+2].add_run(j['Duration'].strip()).bold=True
                        else:
                            doc.paragraphs[i+2].add_run("Duration not mentioned").bold=False  
    
                        if j['Degree'].strip():
                            doc.paragraphs[i+2].add_run("                   "+j['Degree'].strip() +'\n').bold =True
                        else:
                            doc.paragraphs[i+2].add_run("Degree not mentioned"+'\n').bold=False  
                        if j["Institute Name"].strip():    
                            doc.paragraphs[i+2].add_run(j["Institute Name"].strip()).bold=False
                        else:
                            doc.paragraphs[i+2].add_run("Institute Name not mentioned"+'\n').bold=False  
                            
                        if j["Location"].strip():    
                            doc.paragraphs[i+2].add_run(' , ' + j["Location"].strip() + '\n\n').bold= False
                        else:
                            doc.paragraphs[i+2].add_run(' , ' +"Location not mentioned"+'\n\n').bold=False  

        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'languages':
                if dc['Languages'][0].lower().replace(' ','') != 'languages1':
                    for j in dc['Languages']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'interests':
                if dc['Interests'][0].lower().replace(' ','') != 'interests1':
                    for j in dc['Interests']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'training':
                if dc['Trainings'][0].lower().replace(' ','') != 'training1':
                    for j in dc['Trainings']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'skills':
                if dc['Skills'][0].lower().replace(' ','') != 'skills1':
                    for j in dc['Skills']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:

            if p.text.strip(' :\n').lower() == 'employment history':
                for j in dc['Employment History']:
                    if (j['Designation'].strip() and j['Designation'].lower().replace(' ','') !='specificdesignationinthatcompany') or (j['Company Name'].strip() and j['Company Name'].lower().replace(' ','') !='nameofcompany'):
                        if j['Duration'].strip():                                                       
                            doc.paragraphs[i+2].add_run(j['Duration'].strip()).bold=True
                        else:
                            doc.paragraphs[i+2].add_run("Duration not mentioned").bold=True
                        if j['Designation'].strip():    
                            doc.paragraphs[i+2].add_run("                  "+j['Designation'].strip()+ '\n').bold=True
                        else:
                            doc.paragraphs[i+2].add_run("Designation not mentioned"+"\n").bold=True
                        if j['Company Name'].strip():     
                            doc.paragraphs[i+2].add_run(j['Company Name'].strip()).bold=True
                        else:
                            doc.paragraphs[i+2].add_run("Company Name not mentioned").bold=True                           
                        if j['Location'].strip():    
                            doc.paragraphs[i+2].add_run(' – ' + j['Location'].strip() + '\n\n').bold =True
                        else:
                            doc.paragraphs[i+2].add_run("Location not mentioned"+'\n\n').bold=True
                            
                        if j['Responsibilities'] and j['Responsibilities'][0].lower().replace(' ','') != 'responsibility1':
                            for k in j['Responsibilities']:
                                doc.paragraphs[i+2].add_run('  • ' + k.strip() + '\n').bold = False
                            doc.paragraphs[i+2].add_run('\n\n')
        except:
            pass
    
    doc.save(savepath) 
    print("Process has Completed...")
