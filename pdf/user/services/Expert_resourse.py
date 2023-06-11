import os
import openai
import docx
import docx2txt
import PyPDF2
import re
import json
from docx import Document
from .keys import api_key


def get_completion(prompt, model="gpt-3.5-turbo"):
    messages = [{"role": "user", "content": prompt}]
    response = openai.ChatCompletion.create(
        model=model,
        messages=messages,
        temperature=0, # this is the degree of randomness of the model's output
    )
    return response.choices[0].message["content"]

def expert_resource_converter(path, pathout, path_save):
    formatted= pathout
    un_formatted = path

    formated_text = docx2txt.process(formatted)

    try:
        with open(path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            unformated_text = ""
            for i in range (len(pdf_reader.pages)):
                first_page = pdf_reader.pages[i]
                unformated_text += first_page.extract_text() + " "
            print('Its PDF')
    except:
        try:
            unformated_text = docx2txt.process(path)
            print('Its Docx')
        except:
            print('WE DONT SUPPORT THIS TYPE OF FILE')


    os.environ["OPEN_API_KEY"] = api_key
    openai.api_key = api_key
    
    print("Process has Started...")

    test_text = """

    Ectract data from this text:

    \"""" + unformated_text + """\"

    in following JSON format:
    {
    "Name" : "value",
    "Summary" : "value",
    "Summary Overview" : "value",
    "Volunteering" : "value",
    "Education" : [
        {"Institute" : "Name Of institute",
        "Duration": "Studying duration in institute",
        "Location" : "Location of institute",
        "Degree title" : "Name of degree completed from that institute",
        },
        {"Institute" : "Name Of institute",
        "Duration": "Studying duration in institute",
        "Location" : "Location of institute",
        "Degree title" : "Name of degree completed from that institute",
        },
        ...
        ],
    "Achievements" : ["achievement1", "achievement2", ...],
    "Certifications" : ["certification1", "certification2", ...],
    "Awards" : ["award1", "award2", ...],
    "Publications" : ["publication1", "publication2", ...],
    "Tools" : ["tool1", "tool2", ...],
    "Relevant Certifications and experiences" : ["relevant Certification1", "relevant Certification2", ...],
    "Soft Skills" : ["soft skill1", "soft skill2", ...],
    "Technical Skills" : ["technical skill1", "technical skill2", ...],
    "Languages" : ["language1", "language2", ...],
    "Interests" : ["interest1", "interest2", ...],
    "Trainings" : ["training1", "training2", ...],
    "Skills" : ["skill1", "skill2", ...],
    "Experience" : [
        {"Role" : "Specific role in that Company",
        "Name of Company" : "Company Name here",
        "Client" : "value",
        "Duration" : "Time period in whic he the person has worked in that company",
        "Responsibilities" : ["Responsibility 1", "Responsibility 2", ...],
        },
        {"Role" : "Specific role in that Company",
        "Name of Company" : "Company Name here",
        "Client" : "value",
        "Duration" : "Time period in whic he the person has worked in that company",
        "Responsibilities" : ["Responsibility 1", "Responsibility 2", ...],
        },
        ...
        ]
    }

    You must keep the following points in considration while extracting data from text:
        1. Do NOT split, rephrase or summarize list of Responsibilities. Extract each Responsibility as a complete sentence from text.
        2. Make it sure to keep the response in JSON format.
        3. If value not found then leave it empty/blank.
        4. Do not include Mobile number, Email and Home address. 
        5. Summary/Personal Statement should be complete without being rephrased.

    """


    result = get_completion(test_text)


    dc = dict(json.loads(re.sub(r'\[\"\"\]',r'[]',re.sub(r'\"[Un]nknown\"|\"[Nn]one\"|\"[Nn]ull\"|\"[Nn]ot [Mm]entioned\"',r'""',re.sub(r',[ \n]*\]',r']',re.sub(r',[ \n]*\}',r'}',result.replace('...','')))))))

    doc = docx.Document(formatted)

    for i,p in enumerate(doc.paragraphs):  

        try:
            if p.text.strip(' :\n').lower() == 'name':
                if dc['Name'].lower().replace(' ','') != 'value':
                    doc.paragraphs[i].add_run("\t" + dc['Name'].strip()).bold = False

        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'volunteering':
                if dc['Volunteering'].lower().replace(' ','') != 'value':
                    doc.paragraphs[i+2].add_run(dc['Volunteering'].strip()).bold = False

        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'summary overview':
                if dc['Summary Overview'].lower().replace(' ','') != 'value':
                    doc.paragraphs[i+2].add_run(dc['Summary Overview'].strip()).bold = False
                    doc.paragraphs[i+2].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'summary':
                if dc['Summary'].lower().replace(' ','') != 'value':
                    doc.paragraphs[i+2].add_run(dc['Summary'].strip()).bold = False
                    doc.paragraphs[i+2].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        except:
            pass

        try:     
            if p.text.strip(' :\n').lower() == 'education':
                for j in dc['Education']:
                    if j['Degree title'].strip() and j['Degree title'].lower().replace(' ','') != "nameofdegreecompletedfromthatinstitute":
                        if j['Duration'].strip():                                                  
                            doc.paragraphs[i+2].add_run(j['Duration'].strip() + '\n').bold = True
                        else:
                            doc.paragraphs[i].add_run("Duration not mentioned"+"\n")
                        if j['Institute'].strip():    
                            doc.paragraphs[i+2].add_run(j['Institute'].strip()).bold=False
                        else:
                            doc.paragraphs[i].add_run("Institute name not mentioned")
                        if j["Location"].strip():    
                            doc.paragraphs[i+2].add_run(", " + j["Location"].strip() + '\n').bold = False
                        else:
                            doc.paragraphs[i].add_run("Location not mentioned")
  
                        doc.paragraphs[i+2].add_run(j['Degree title'].strip() + '\n\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'achievements':
                if dc['Achievements'][0].lower().replace(' ','') != 'achievement1':
                    for j in dc['Achievements']:
                        doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'awards':
                if dc['Awards'][0].lower().replace(' ','') != 'award1':
                    for j in dc['Awards']:
                        doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'relevant certifications and experiences':
                if dc['Relevant Certifications and experiences'][0].lower().replace(' ','') != 'relevant Certification1':
                    for j in dc['Relevant Certifications and experiences']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'publications':
                if dc['Publications'][0].lower().replace(' ','') != 'publication1':
                    for j in dc['Publications']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'certifications':
                if dc['Certifications'][0].lower().replace(' ','') != 'certification1':
                    for j in dc['Certifications']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'soft skills':
                if dc['Soft Skills'][0].lower().replace(' ','') != 'softskill1':
                    for j in dc['Soft Skills']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'technical skills':
                if dc['Technical Skills'][0].lower().replace(' ','') != 'technicalskill1':
                    for j in dc['Technical Skills']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'languages':
                if dc['Languages'][0].lower().replace(' ','') != 'language1':
                    for j in dc['Languages']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'interests':
                if dc['Interests'][0].lower().replace(' ','') != 'interest1':
                    for j in dc['Interests']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'trainings':
                if dc['Trainings'][0].lower().replace(' ','') != 'training1':
                    for j in dc['Trainings']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        try:
            if p.text.strip(' :\n').lower() == 'skills':
                if dc['Skills'][0].lower().replace(' ','') != 'skill1':
                    for j in dc['Skills']:
                        if j.strip():
                            doc.paragraphs[i+2].add_run('  • ' + j.strip() + '\n').bold = False
        except:
            pass

        if p.text.strip(' :\n').lower() == 'employment summary':
            for j in dc['Experience']:
                if j['Role'].strip() and j['Role'].lower().replace(' ','') !='specificroleinthatcompany' or (j['Name of Company'].strip() and j['Name of Company'].lower().replace(' ','') !='companyNamehere'):                    
                    if j['Duration'].strip():
                        doc.paragraphs[i+2].add_run(j['Duration'].strip() + '\n').bold = True
                    else:
                        doc.paragraphs[i+2].add_run("Duration not mentioned"+"\n")
                    if j['Role'].strip(): 
                        doc.paragraphs[i+2].add_run("Role:" + "\t\t" + j['Role'].strip() + '\n').bold = False
                    else:
                        doc.paragraphs[i+2].add_run("Role not mentioned"+"\n")
                    if j['Client'].strip():                      
                        doc.paragraphs[i+2].add_run("Client:" + "\t\t" + j['Client'].strip() + '\n').bold = False
                    else:
                        doc.paragraphs[i+2].add_run("Client not mentioned"+"\n")
                    if j['Name of Company'].strip(): 
                        doc.paragraphs[i+2].add_run("Company:" + "\t" + j['Name of Company'].strip() + '\n\n').bold = False
                    else:                        
                        doc.paragraphs[i+2].add_run("Company not mentioned"+"\n\n")
                    if j['Responsibilities'] and j['Responsibilities'][0].lower().replace(' ','') != 'responsibility1':          
                        for k in j['Responsibilities']:
                            doc.paragraphs[i+2].add_run('  • ' + k.strip() + '\n').bold = False
                        doc.paragraphs[i+2].add_run('\n')

    doc.save(path_save)
    print("Process has Completed...")
